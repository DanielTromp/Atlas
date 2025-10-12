from __future__ import annotations

import asyncio
import csv
import html
import json
import os
import re
import secrets
import shutil
import sys
import time
import uuid
import warnings
from collections import Counter
from collections.abc import Collection, Iterable, Mapping, Sequence
from contextlib import contextmanager
from dataclasses import dataclass
from datetime import UTC, datetime, timedelta
from io import BytesIO, StringIO
from pathlib import Path
from threading import Lock
from typing import Annotated, Any, Literal
from zoneinfo import ZoneInfo

import duckdb
import numpy as np
import pandas as pd
import requests
from dotenv import dotenv_values
from fastapi import Body, FastAPI, HTTPException, Query, Request, Response
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import (
    FileResponse,
    HTMLResponse,
    JSONResponse,
    PlainTextResponse,
    RedirectResponse,
    StreamingResponse,
)
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from sqlalchemy import select
from sqlalchemy.orm import Session
from starlette.middleware.base import BaseHTTPMiddleware
from starlette.middleware.sessions import SessionMiddleware

from enreach_tools import backup_sync
from enreach_tools.application.chat_agents import AgentRuntime, AgentRuntimeError
from enreach_tools.application.dto import (
    AdminEnvResponseDTO,
    BackupInfoDTO,
    EnvSettingDTO,
    SuggestionCommentDTO,
    SuggestionCommentResponseDTO,
    SuggestionItemDTO,
    SuggestionListDTO,
    meta_to_dto,
    suggestion_to_dto,
    suggestions_to_dto,
)
from enreach_tools.application.role_defaults import DEFAULT_ROLE_DEFINITIONS
from enreach_tools.application.security import hash_password, verify_password
from enreach_tools.application.services import create_vcenter_service
from enreach_tools.db import get_sessionmaker, init_database
from enreach_tools.db.models import (
    ChatMessage,
    ChatSession,
    GlobalAPIKey,
    RolePermission,
    User,
    UserAPIKey,
)
from enreach_tools.domain.integrations.commvault import (
    CommvaultJob,
    CommvaultJobList,
    CommvaultPlan,
    CommvaultStoragePool,
)
from enreach_tools.env import load_env, project_root
from enreach_tools.infrastructure.external import (
    ZabbixAuthError,
    ZabbixClient,
    ZabbixClientConfig,
    ZabbixConfigError,
    ZabbixError,
)
from enreach_tools.infrastructure.external.commvault_client import (
    CommvaultClient,
    CommvaultClientConfig,
    CommvaultError,
)
from enreach_tools.infrastructure.logging import get_logger, logging_context, setup_logging
from enreach_tools.infrastructure.metrics import get_metrics_snapshot, snapshot_to_prometheus
from enreach_tools.infrastructure.tracing import init_tracing, tracing_enabled
from enreach_tools.interfaces.api import bootstrap_api
from enreach_tools.interfaces.api.dependencies import (
    CurrentUserDep,
    DbSessionDep,
    OptionalUserDep,
)
from enreach_tools.interfaces.api.middleware import ObservabilityMiddleware
from enreach_tools.interfaces.api.routes import tools as tools_router
from enreach_tools.interfaces.api.schemas import ToolDefinition

try:
    from openai import OpenAI  # type: ignore
except Exception:  # optional dependency
    OpenAI = None  # type: ignore

try:  # optional dependency; urllib3 may not always be available
    from urllib3 import disable_warnings as _disable_urllib3_warnings
    from urllib3.exceptions import InsecureRequestWarning as _InsecureRequestWarning
except Exception:  # pragma: no cover - urllib3 optional
    _disable_urllib3_warnings = None  # type: ignore[assignment]
    _InsecureRequestWarning = None  # type: ignore[assignment]

load_env()

AMS_TZ = ZoneInfo("Europe/Amsterdam")

logger = get_logger(__name__)
if tracing_enabled():
    init_tracing("enreach-api")

# Ensure the application database is migrated to the latest revision so
# authentication state can be loaded lazily by request handlers.
try:
    init_database()
except Exception as exc:  # pragma: no cover - surfaces during boot
    logger.exception("Failed to initialise database during API startup", extra={"error": str(exc)})
    raise
else:
    # Alembic can override logging handlers; ensure our configuration is still active.
    setup_logging()
    logger.disabled = False

SessionLocal = get_sessionmaker()


@dataclass(slots=True)
class ChatProviderResult:
    text: str
    usage: dict[str, int] | None = None
    metadata: dict[str, Any] | None = None


class _TaskLogger:
    """Mutable container passed into task_logging for success metadata."""

    __slots__ = ("_success_extra",)

    def __init__(self) -> None:
        self._success_extra: dict[str, Any] = {}

    def add_success(self, **kwargs: Any) -> None:
        for key, value in kwargs.items():
            if value is not None:
                self._success_extra[key] = value

    @property
    def success_extra(self) -> dict[str, Any]:
        return self._success_extra


@contextmanager
def task_logging(task: str, **context: Any):
    """Log lifecycle events around long-running UI-triggered tasks."""

    start = time.perf_counter()
    tracker = _TaskLogger()
    safe_context = {k: v for k, v in context.items() if v not in (None, "")}

    with logging_context(task=task, **safe_context):
        logger.info("Task started", extra={"event": "task_started"})
        try:
            yield tracker
        except Exception:
            duration_ms = int((time.perf_counter() - start) * 1000)
            with logging_context(duration_ms=duration_ms, **tracker.success_extra):
                logger.exception("Task failed", extra={"event": "task_failed"})
            raise
        else:
            duration_ms = int((time.perf_counter() - start) * 1000)
            with logging_context(duration_ms=duration_ms, **tracker.success_extra):
                logger.info("Task completed", extra={"event": "task_completed"})


_USAGE_FIELD_ALIASES: dict[str, str] = {
    "prompt_tokens": "prompt_tokens",
    "completion_tokens": "completion_tokens",
    "total_tokens": "total_tokens",
    "input_tokens": "prompt_tokens",
    "output_tokens": "completion_tokens",
    "usage_tokens": "total_tokens",
    "input_token_count": "prompt_tokens",
    "output_token_count": "completion_tokens",
    "total_token_count": "total_tokens",
    "promptTokenCount": "prompt_tokens",
    "candidatesTokenCount": "completion_tokens",
    "totalTokens": "total_tokens",
}


def _normalise_usage(raw: Any) -> dict[str, int] | None:
    if raw is None:
        return None
    usage: dict[str, int] = {}
    for source, target in _USAGE_FIELD_ALIASES.items():
        value = None
        if isinstance(raw, dict):
            value = raw.get(source)
        else:
            value = getattr(raw, source, None)
        if isinstance(value, int | float):
            usage[target] = int(value)
    return usage or None


def _env_flag(name: str, default: bool = False) -> bool:
    value = os.getenv(name)
    if value is None:
        return default
    return value.strip().lower() in {"1", "true", "yes", "on"}


METRICS_ENABLED = _env_flag("ENREACH_METRICS_ENABLED")
METRICS_TOKEN = os.getenv("ENREACH_METRICS_TOKEN", "").strip()
METRICS_MEDIA_TYPE = "text/plain; version=0.0.4; charset=utf-8"


def ensure_default_role_permissions() -> None:
    with SessionLocal() as db:
        existing = {
            record.role: record
            for record in db.execute(select(RolePermission)).scalars()
        }
        changed = False
        for role, spec in DEFAULT_ROLE_DEFINITIONS.items():
            label = (spec.get("label") or role).strip() or role
            description = spec.get("description") or None
            default_perms = sorted({str(p).strip() for p in spec.get("permissions", []) if str(p).strip()})
            record = existing.get(role)
            if record is None:
                record = RolePermission(
                    role=role,
                    label=label,
                    description=description,
                    permissions=default_perms,
                )
                db.add(record)
                changed = True
            else:
                updated = False
                if not record.label:
                    record.label = label
                    updated = True
                if description and not record.description:
                    record.description = description
                    updated = True
                current = set(record.permissions or [])
                desired = current.union(default_perms)
                if desired != current:
                    record.permissions = sorted(desired)
                    updated = True
                if updated:
                    db.add(record)
                    changed = True
        if changed:
            db.commit()


def ensure_default_admin() -> None:
    with SessionLocal() as db:
        existing = db.execute(select(User.id).limit(1)).scalar_one_or_none()
        if existing:
            return

        username = os.getenv("ENREACH_DEFAULT_ADMIN_USERNAME", "admin").strip().lower() or "admin"
        seed_password = os.getenv("ENREACH_DEFAULT_ADMIN_PASSWORD", "").strip() or UI_PASSWORD

        if not seed_password:
            logger.warning(
                "No users exist and ENREACH_DEFAULT_ADMIN_PASSWORD is not set; set it (or ENREACH_UI_PASSWORD) to bootstrap the first login.",
                extra={"username": username},
            )
            return

        user = User(
            username=username,
            display_name="Administrator",
            role="admin",
            password_hash=hash_password(seed_password),
            is_active=True,
        )
        db.add(user)
        db.commit()
        logger.info(
            "Created default admin user",
            extra={"username": username},
        )


ensure_default_role_permissions()
ensure_default_admin()


def get_user_by_username(db: Session, username: str) -> User | None:
    uname = (username or "").strip().lower()
    if not uname:
        return None
    stmt = select(User).where(User.username == uname)
    return db.execute(stmt).scalar_one_or_none()


def authenticate_user(db: Session, username: str, password: str) -> User | None:
    user = get_user_by_username(db, username)
    if not user or not user.is_active:
        return None
    if verify_password(password, user.password_hash):
        return user
    return None


def _ensure_admin(user: User) -> User:
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    return user


def _iso(dt: datetime | None) -> str | None:
    if not dt:
        return None
    if isinstance(dt, datetime):
        return dt.isoformat()
    try:
        return str(dt)
    except Exception:
        return None


def _serialize_user(user: User) -> dict[str, Any]:
    return {
        "id": user.id,
        "username": user.username,
        "display_name": user.display_name,
        "email": user.email,
        "role": user.role,
        "is_active": user.is_active,
        "created_at": _iso(user.created_at),
        "updated_at": _iso(user.updated_at),
    }


def _serialize_user_api_key(record: UserAPIKey) -> dict[str, Any]:
    return {
        "id": record.id,
        "provider": record.provider,
        "label": record.label,
        "created_at": _iso(record.created_at),
        "updated_at": _iso(record.updated_at),
    }


def _serialize_global_api_key(record: GlobalAPIKey) -> dict[str, Any]:
    return {
        "id": record.id,
        "provider": record.provider,
        "label": record.label,
        "created_at": _iso(record.created_at),
        "updated_at": _iso(record.updated_at),
    }


def _get_user_api_key(db: Session, user_id: str, provider: str) -> UserAPIKey | None:
    stmt = select(UserAPIKey).where(
        UserAPIKey.user_id == user_id,
        UserAPIKey.provider == provider,
    )
    return db.execute(stmt).scalar_one_or_none()


def _get_global_api_key(db: Session, provider: str) -> GlobalAPIKey | None:
    stmt = select(GlobalAPIKey).where(GlobalAPIKey.provider == provider)
    return db.execute(stmt).scalar_one_or_none()


app = FastAPI(title="Enreach Tools API", version="0.1.0")
app.add_middleware(
    ObservabilityMiddleware,
    metrics_enabled=METRICS_ENABLED,
)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["GET", "POST"],
    allow_headers=["*"],
)


app.include_router(bootstrap_api())

# Include monitoring routes
try:
    from enreach_tools.interfaces.api.routes.monitoring import router as monitoring_router
    app.include_router(monitoring_router)
except ImportError:
    logger.warning("Monitoring routes not available - optional dependencies missing")


def _require_metrics_token(request: Request) -> None:
    if not METRICS_TOKEN:
        return
    auth = request.headers.get("Authorization", "")
    if not auth.lower().startswith("bearer "):
        raise HTTPException(status_code=401, detail="Unauthorized")
    token = auth[7:].strip()
    if not secrets.compare_digest(token, METRICS_TOKEN):
        raise HTTPException(status_code=401, detail="Unauthorized")


if METRICS_ENABLED:

    @app.get("/metrics", response_class=PlainTextResponse)
    def metrics_endpoint(request: Request) -> PlainTextResponse:
        _require_metrics_token(request)
        payload = snapshot_to_prometheus(get_metrics_snapshot())
        return PlainTextResponse(payload, media_type=METRICS_MEDIA_TYPE)


# Reduce aggressive caching for the static UI during development to avoid stale assets
@app.middleware("http")
async def no_cache_static(request, call_next):
    response = await call_next(request)
    try:
        path = request.url.path
        if path.startswith("/app"):
            response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
            response.headers["Pragma"] = "no-cache"
            response.headers["Expires"] = "0"
    except Exception:
        pass
    return response


# ---------------------------
# Auth configuration
# ---------------------------

API_TOKEN = os.getenv("ENREACH_API_TOKEN", "").strip()
UI_PASSWORD = os.getenv("ENREACH_UI_PASSWORD", "").strip()  # legacy fallback
UI_SECRET = os.getenv("ENREACH_UI_SECRET", "").strip() or secrets.token_hex(32)
SESSION_COOKIE_NAME = "enreach_ui"
SESSION_USER_KEY = "user_id"


class AuthMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: Request, call_next):
        path = request.url.path or "/"

        request.state.user = None
        request.state.permissions = frozenset()
        user_id = request.session.get(SESSION_USER_KEY) if hasattr(request, "session") else None
        if user_id:
            with SessionLocal() as db:
                user = db.get(User, user_id)
                if user and user.is_active:
                    perms_record = db.get(RolePermission, user.role)
                    perms = frozenset((perms_record.permissions or []) if perms_record else [])
                    request.state.user = user
                    request.state.permissions = perms
                    try:
                        user.permissions = perms
                    except Exception:
                        pass
                else:
                    request.session.pop(SESSION_USER_KEY, None)
                    request.state.user = None
                    request.state.permissions = frozenset()

        # Public endpoints
        if path in ("/favicon.ico", "/health"):
            return await call_next(request)

        if path == "/":
            if not _has_ui_session(request):
                return RedirectResponse(url="/auth/login")
            return await call_next(request)
        if path.startswith("/app"):
            if not _has_ui_session(request):
                next_url = path
                if request.url.query:
                    next_url += f"?{request.url.query}"
                return RedirectResponse(url=f"/auth/login?next={next_url}")
            return await call_next(request)

        # Auth endpoints are always allowed
        if path.startswith("/auth/"):
            return await call_next(request)

        # API protection via Bearer token; allow UI session as alternative
        if API_TOKEN and _is_api_path(path):
            if _has_bearer_token(request) or _has_ui_session(request):
                return await call_next(request)
            # Unauthorized
            return JSONResponse({"detail": "Unauthorized"}, status_code=401, headers={"WWW-Authenticate": "Bearer"})

        return await call_next(request)


def _has_bearer_token(request: Request) -> bool:
    if not API_TOKEN:
        return False
    auth = request.headers.get("Authorization", "")
    if not auth.lower().startswith("bearer "):
        return False
    token = auth[7:].strip()
    return secrets.compare_digest(token, API_TOKEN)


def _has_ui_session(request: Request) -> bool:
    try:
        return bool(request.session.get(SESSION_USER_KEY))
    except Exception:
        return False


def require_permission(request: Request, permission: str) -> None:
    user = getattr(request.state, "user", None)
    if user is None:
        return
    if getattr(user, "role", "") == "admin":
        return
    permissions = getattr(request.state, "permissions", frozenset())
    if permission not in permissions:
        raise HTTPException(status_code=403, detail="Forbidden: missing permission")


def _is_api_path(path: str) -> bool:
    # Treat everything except static/frontend and auth endpoints as API
    if path.startswith("/app") or path.startswith("/auth"):
        return False
    if path in ("/", "/favicon.ico"):
        return False
    if path == "/metrics":
        return False
    return True


app.add_middleware(AuthMiddleware)

# Session support for UI auth (cookie-based). Added AFTER AuthMiddleware so
# that Session is applied first (outermost), making request.session available.
app.add_middleware(
    SessionMiddleware,
    secret_key=UI_SECRET,
    session_cookie=SESSION_COOKIE_NAME,
    same_site="lax",
)


def _login_html(next_url: str, error: str | None = None) -> HTMLResponse:
    err_html = f'<div class="error">{error}</div>' if error else ""
    logo_svg = ""
    try:
        raw_logo = (Path(__file__).parent / "static" / "enreach.svg").read_text(encoding="utf-8")
        start = raw_logo.find("<svg")
        if start != -1:
            logo_svg = raw_logo[start:]
        else:
            logo_svg = raw_logo
    except OSError:
        logo_svg = ""
    logo_html = logo_svg or '<span class="brand-logo__fallback">Enreach</span>'
    return HTMLResponse(
        f"""
        <!doctype html>
        <html><head>
        <meta charset=\"utf-8\"><meta name=\"viewport\" content=\"width=device-width, initial-scale=1\">
        <title>Login — Enreach Tools</title>
        <style>
        * {{ box-sizing: border-box; }}
        body {{ font-family: -apple-system, system-ui, Segoe UI, Roboto, Ubuntu, Cantarell, 'Helvetica Neue', Arial, 'Noto Sans', 'Apple Color Emoji', 'Segoe UI Emoji', 'Segoe UI Symbol';
               background: radial-gradient(120% 160% at 50% 0%, rgba(156, 75, 255, 0.28) 0%, rgba(53, 20, 112, 0.68) 40%, #0b0424 75%, #050012 100%);
               color: #f5f3ff; display: flex; align-items: center; justify-content: center; min-height: 100vh; margin: 0; padding: 24px; }}
        .box {{ background: rgba(26, 9, 55, 0.86); padding: 32px 30px; border-radius: 16px; width: min(400px, 100%);
                border: 1px solid rgba(155, 100, 255, 0.38); box-shadow: 0 28px 58px rgba(15, 5, 38, 0.55);
                backdrop-filter: blur(12px); display: flex; flex-direction: column; gap: 18px; }}
        .brand {{ display: flex; align-items: center; gap: 12px; }}
        .brand-logo {{ display: flex; align-items: center; justify-content: center; width: 42px; height: 42px; }}
        .brand-logo svg {{ width: 100%; height: auto; filter: brightness(0) invert(1); opacity: 0.92; display: block; }}
        .brand-logo__fallback {{ color: #f5f3ff; font-weight: 700; }}
        .brand h1 {{ font-size: 20px; margin: 0; letter-spacing: 0.28px; font-weight: 680; }}
        label {{ display: block; }}
        input[type=text], input[type=password] {{ width: 100%; padding: 12px 14px; border-radius: 10px; border: 1px solid rgba(170, 111, 255, 0.55);
                               background: rgba(31, 15, 66, 0.94); color: #f5f3ff; font-size: 14px; transition: border-color 160ms ease, box-shadow 160ms ease; }}
        input[type=text]::placeholder, input[type=password]::placeholder {{ color: rgba(212, 198, 255, 0.65); }}
        input[type=text]:focus, input[type=password]:focus {{ outline: none; border-color: #c084fc; box-shadow: 0 0 0 3px rgba(168, 85, 247, 0.45); }}
        button {{ display: block; width: 100%; padding: 12px 14px; border-radius: 10px; border: 1px solid rgba(138, 78, 255, 0.9);
                  background: linear-gradient(180deg, rgba(138, 78, 255, 0.96), rgba(98, 49, 209, 0.92)); color: #f8f5ff;
                  margin-top: 4px; cursor: pointer; font-weight: 600; transition: transform 120ms ease, box-shadow 160ms ease; }}
        button:hover {{ transform: translateY(-1px); box-shadow: 0 18px 46px rgba(98, 58, 178, 0.5); }}
        button:active {{ transform: translateY(0); }}
        .tagline {{ margin: 0; color: #d5ccff; font-size: 13px; }}
        .hint {{ color: #b2a8d9; font-size: 12px; text-align: center; margin-top: 8px; }}
        .error {{ background: rgba(127, 29, 29, 0.85); color: #fecaca; padding: 10px 12px; border-radius: 10px;
                  border: 1px solid rgba(248, 113, 113, 0.65); margin: 4px 0; }}
        </style>
        </head><body>
        <form class=\"box\" method=\"post\" action=\"/auth/login\">
          <div class=\"brand\">
            <div class=\"brand-logo\">{logo_html}</div>
            <h1>Enreach Tools</h1>
          </div>
          <p class=\"tagline\">Sign in to manage NetBox exports and tools.</p>
          {err_html}
          <input type=\"hidden\" name=\"next\" value=\"{next_url}\" />
          <label>
            <input type=\"text\" name=\"username\" placeholder=\"Username\" autofocus required />
          </label>
          <label>
            <input type=\"password\" name=\"password\" placeholder=\"Password\" required />
          </label>
          <button type=\"submit\">Sign in</button>
          <div class=\"hint\">UI access enables API calls from this browser.</div>
        </form>
        </body></html>
        """
    )


@app.get("/auth/login")
def auth_login_form(request: Request, next_url: str | None = None):
    n = next_url or "/app/"
    # If already logged in, hop to target
    if _has_ui_session(request):
        return RedirectResponse(url=n)
    return _login_html(n)


@app.post("/auth/login")
async def auth_login(request: Request):
    content_type = request.headers.get("content-type", "").lower()
    is_json = "application/json" in content_type

    if is_json:
        payload = await request.json()
    else:
        payload = await request.form()

    username = str(payload.get("username") or "").strip()
    password = str(payload.get("password") or "")
    next_url = str(payload.get("next") or "/app/")
    if not next_url.startswith("/"):
        next_url = "/app/"

    with SessionLocal() as db:
        user = authenticate_user(db, username, password)
        if user:
            request.session.clear()
            request.session[SESSION_USER_KEY] = user.id
            request.session["username"] = user.username
            if is_json:
                return {"status": "ok", "next": next_url}
            return RedirectResponse(url=next_url, status_code=303)

    if is_json:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    return _login_html(next_url, error="Invalid username or password")


@app.get("/auth/logout")
def auth_logout(request: Request):
    try:
        request.session.clear()
    except Exception:
        pass
    return RedirectResponse(url="/auth/login")


@app.get("/admin/users")
def _data_dir() -> Path:
    root = project_root()
    raw = os.getenv("NETBOX_DATA_DIR", "data")
    path = Path(raw) if os.path.isabs(raw) else (root / raw)
    path.mkdir(parents=True, exist_ok=True)
    return path


def _csv_path(name: str) -> Path:
    return _data_dir() / name


# Simple export log file (appends)
LOG_PATH = project_root() / "export.log"

COMMVAULT_BACKUPS_JSON = "commvault_backups.json"
COMMVAULT_STORAGE_JSON = "commvault_storage.json"
COMMVAULT_PLANS_JSON = "commvault_plans.json"
COMMVAULT_DEFAULT_PLAN_NAME = "Unassigned"
COMMVAULT_PLAN_TYPE_LABELS: dict[int, str] = {
    2: "Server",
}
_commvault_backups_lock = Lock()
_commvault_storage_lock = Lock()
_commvault_plans_lock = Lock()


def _commvault_client_from_env() -> CommvaultClient:
    base_url = os.getenv("COMMVAULT_BASE_URL", "").strip()
    if not base_url:
        raise RuntimeError("Commvault integration is not configured (COMMVAULT_BASE_URL missing)")

    authtoken = os.getenv("COMMVAULT_API_TOKEN", "").strip() or None
    username = os.getenv("COMMVAULT_EMAIL", "").strip() or None
    password = os.getenv("COMMVAULT_PASSWORD", "")
    if password:
        password = password.strip()
    if not authtoken and not (username and password):
        raise RuntimeError("Commvault credentials are not configured (token or email/password required)")

    verify_tls = _env_flag("COMMVAULT_VERIFY_TLS", True)
    timeout_raw = os.getenv("COMMVAULT_TIMEOUT", "").strip()
    try:
        timeout = float(timeout_raw) if timeout_raw else 30.0
    except ValueError:
        timeout = 30.0

    if not verify_tls:
        _maybe_disable_commvault_tls_warnings()

    config = CommvaultClientConfig(
        base_url=base_url,
        authtoken=authtoken,
        username=username,
        password=password or None,
        verify_tls=verify_tls,
        timeout=timeout,
    )
    return CommvaultClient(config)


def _maybe_disable_commvault_tls_warnings() -> None:
    """Mute urllib3 TLS warnings when verification is intentionally disabled."""

    if _disable_urllib3_warnings and _InsecureRequestWarning:
        _disable_urllib3_warnings(_InsecureRequestWarning)
        return
    warnings.filterwarnings("ignore", category=Warning, module="urllib3")


def _collect_commvault_jobs_for_ui(
    client: CommvaultClient,
    *,
    limit: int,
    offset: int,
    since: datetime | None,
    tracked_job_ids: Collection[int] | None = None,
    latest_start: datetime | None = None,
    cached_recent_count: int | None = None,
    cutoff_start: datetime | None = None,
) -> CommvaultJobList:
    if limit < 0:
        raise ValueError("limit must be zero or positive")

    job_type = None
    page_target = 200 if limit == 0 else max(1, min(500, limit))
    tracked_remaining: set[int] = {int(job_id) for job_id in (tracked_job_ids or [])}
    threshold: datetime | None = since
    if latest_start and (threshold is None or latest_start < threshold):
        threshold = latest_start
    if cutoff_start and (threshold is None or cutoff_start < threshold):
        threshold = cutoff_start

    def _list_jobs_with_retry(*, fetch_limit: int, fetch_offset: int) -> CommvaultJobList:
        attempt_limit = max(1, fetch_limit)
        last_error: CommvaultError | None = None
        while attempt_limit >= 1:
            try:
                return client.list_jobs(
                    limit=attempt_limit,
                    offset=fetch_offset,
                    job_type=job_type,
                )
            except CommvaultError as exc:
                last_error = exc
                if attempt_limit <= 25:
                    break
                attempt_limit = max(25, attempt_limit // 2)
        if last_error is not None:
            raise last_error
        return CommvaultJobList(total_available=0, jobs=())

    try:
        initial = _list_jobs_with_retry(fetch_limit=1, fetch_offset=0)
    except CommvaultError:
        initial = _list_jobs_with_retry(fetch_limit=25, fetch_offset=0)

    total_available = initial.total_available or len(initial.jobs)
    if not total_available:
        return CommvaultJobList(total_available=0, jobs=())

    target_end = max(0, total_available - offset)
    if target_end <= 0:
        return CommvaultJobList(total_available=total_available, jobs=())

    approximate_needed = page_target
    if cached_recent_count is not None:
        approximate_needed = max(approximate_needed, min(total_available, cached_recent_count + page_target))
    if tracked_remaining:
        approximate_needed = max(approximate_needed, min(total_available, len(tracked_remaining) + page_target))
    if limit > 0:
        approximate_needed = max(approximate_needed, min(total_available, limit + page_target))

    target_start = max(0, target_end - approximate_needed)
    jobs: list[CommvaultJob] = []
    seen: set[int] = set()
    current_end = target_end

    while current_end > target_start and (limit == 0 or len(jobs) < limit):
        request_limit = min(page_target, current_end - target_start)
        try:
            response = _list_jobs_with_retry(
                fetch_limit=request_limit,
                fetch_offset=current_end - request_limit,
            )
        except CommvaultError:
            break

        batch = list(response.jobs)
        if not batch:
            break

        for job in reversed(batch):
            try:
                job_id = int(job.job_id)
            except (TypeError, ValueError):
                job_id = None
            if job_id is not None and job_id in seen:
                continue
            jobs.append(job)
            if job_id is not None:
                seen.add(job_id)
                tracked_remaining.discard(job_id)

        current_end -= len(batch)

        if limit > 0 and len(jobs) >= limit:
            break
        if tracked_remaining:
            continue
        if threshold:
            oldest = batch[0].start_time if batch else None
            if oldest and oldest < threshold:
                break
        if len(batch) < request_limit:
            break

    if limit > 0 and len(jobs) > limit:
        jobs = jobs[:limit]
    return CommvaultJobList(total_available=total_available, jobs=tuple(jobs))


def _serialise_commvault_job(job: CommvaultJob) -> dict[str, Any]:
    def iso(dt: datetime | None) -> str | None:
        if not dt:
            return None
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=UTC)
        return dt.astimezone(UTC).isoformat()

    return {
        "job_id": job.job_id,
        "job_type": job.job_type,
        "status": job.status,
        "localized_status": job.localized_status,
        "localized_operation": job.localized_operation,
        "client_name": job.client_name,
        "client_id": job.client_id,
        "destination_client_name": job.destination_client_name,
        "subclient_name": job.subclient_name,
        "backup_set_name": job.backup_set_name,
        "application_name": job.application_name,
        "backup_level_name": job.backup_level_name,
        "plan_name": job.plan_name,
        "client_groups": list(job.client_groups),
        "storage_policy_name": job.storage_policy_name,
        "start_time": iso(job.start_time),
        "end_time": iso(job.end_time),
        "elapsed_seconds": job.elapsed_seconds,
        "size_of_application_bytes": job.size_of_application_bytes,
        "size_on_media_bytes": job.size_on_media_bytes,
        "total_num_files": job.total_num_files,
        "percent_complete": job.percent_complete,
        "percent_savings": job.percent_savings,
        "average_throughput_gb_per_hr": job.average_throughput,
        "retain_until": iso(job.retain_until),
    }


def _serialise_commvault_plan(plan: CommvaultPlan) -> dict[str, Any]:
    return {
        "plan_id": plan.plan_id,
        "name": plan.name,
        "plan_type": plan.plan_type,
        "associated_entities": plan.associated_entities,
        "rpo": plan.rpo,
        "copy_count": plan.copy_count,
        "status": plan.status,
        "tags": list(plan.tags),
        "raw": plan.raw,
    }


def _write_commvault_backups_json(payload: Mapping[str, Any]) -> None:
    path = _data_dir() / COMMVAULT_BACKUPS_JSON
    tmp = path.with_suffix(path.suffix + ".tmp")
    with tmp.open("w", encoding="utf-8") as handle:
        json.dump(payload, handle, indent=2)
    tmp.replace(path)


def _write_commvault_storage_json(payload: Mapping[str, Any]) -> None:
    path = _data_dir() / COMMVAULT_STORAGE_JSON
    tmp = path.with_suffix(path.suffix + ".tmp")
    with tmp.open("w", encoding="utf-8") as handle:
        json.dump(payload, handle, indent=2)
    tmp.replace(path)


def _write_commvault_plans_json(payload: Mapping[str, Any]) -> None:
    path = _data_dir() / COMMVAULT_PLANS_JSON
    tmp = path.with_suffix(path.suffix + ".tmp")
    with tmp.open("w", encoding="utf-8") as handle:
        json.dump(payload, handle, indent=2)
    tmp.replace(path)


def _load_commvault_backups() -> dict[str, Any]:
    path = _data_dir() / COMMVAULT_BACKUPS_JSON
    if not path.exists():
        return {
            "jobs": [],
            "generated_at": None,
            "total_cached": 0,
            "version": 2,
        }
    try:
        with path.open("r", encoding="utf-8") as handle:
            data = json.load(handle)
            if isinstance(data, Mapping):
                jobs = list(data.get("jobs") or [])
                generated = data.get("generated_at")
                total_cached = data.get("total_cached")
                try:
                    total_cached = int(total_cached)
                except (TypeError, ValueError):
                    total_cached = len(jobs)
                return {
                    "jobs": jobs,
                    "generated_at": generated,
                    "total_cached": total_cached,
                    "version": data.get("version", 2),
                }
    except Exception:
        logger.exception("Failed to load cached Commvault backups", extra={"event": "commvault_cache_error"})
    return {
        "jobs": [],
        "generated_at": None,
        "total_cached": 0,
        "version": 2,
    }


def _load_commvault_plans() -> dict[str, Any]:
    path = _data_dir() / COMMVAULT_PLANS_JSON
    if not path.exists():
        return {
            "plans": [],
            "generated_at": None,
            "total_cached": 0,
            "plan_types": [],
            "version": 1,
        }
    try:
        with path.open("r", encoding="utf-8") as handle:
            data = json.load(handle)
            if isinstance(data, Mapping):
                plans = list(data.get("plans") or [])
                plan_types = list(data.get("plan_types") or [])
                return {
                    "plans": plans,
                    "generated_at": data.get("generated_at"),
                    "total_cached": data.get("total_cached", len(plans)),
                    "plan_types": plan_types,
                    "version": data.get("version", 1),
                }
    except Exception:
        logger.exception("Failed to load cached Commvault plans", extra={"event": "commvault_plan_cache_error"})
    return {
        "plans": [],
        "generated_at": None,
        "total_cached": 0,
        "plan_types": [],
        "version": 1,
    }


def _parse_job_datetime(value: Any) -> datetime | None:
    if isinstance(value, datetime):
        return value if value.tzinfo else value.replace(tzinfo=UTC)
    if not value:
        return None
    try:
        parsed = datetime.fromisoformat(str(value))
    except ValueError:
        return None
    if parsed.tzinfo is None:
        parsed = parsed.replace(tzinfo=UTC)
    return parsed


def _filter_commvault_jobs(jobs: Iterable[Mapping[str, Any]], since_hours: int) -> list[dict[str, Any]]:
    if since_hours <= 0:
        return [dict(job) for job in jobs]
    cutoff = datetime.now(tz=UTC) - timedelta(hours=since_hours)
    filtered: list[dict[str, Any]] = []
    for job in jobs:
        start = _parse_job_datetime(job.get("start_time"))
        if start and start >= cutoff:
            filtered.append(dict(job))
    return filtered


_ACTIVE_STATUS_KEYWORDS = ("running", "pending", "waiting", "queued", "active", "suspended", "in progress")

_FAILURE_STATUS_KEYWORDS = ("fail", "error", "denied", "invalid", "timeout", "timed out", "kill")
_SUCCESS_STATUS_KEYWORDS = ("complete", "success", "ok", "done")
_WARNING_STATUS_KEYWORDS = ("warning", "partial", "skipped")


def _commvault_job_digest(job: Mapping[str, Any]) -> str:
    try:
        return json.dumps(job, sort_keys=True, default=str)
    except Exception:
        return repr(sorted(job.items()))


def _commvault_job_is_active(job: Mapping[str, Any]) -> bool:
    status_raw = job.get("localized_status") or job.get("status") or ""
    status = str(status_raw).strip().lower()
    if not status:
        return True
    for keyword in _ACTIVE_STATUS_KEYWORDS:
        if keyword in status:
            return True
    end_time = _parse_job_datetime(job.get("end_time"))
    return end_time is None


def _normalise_commvault_plan_name(value: Any) -> str:
    if isinstance(value, str):
        text = value.strip()
        if text:
            return text
    if value is None:
        return COMMVAULT_DEFAULT_PLAN_NAME
    text = str(value).strip()
    return text or COMMVAULT_DEFAULT_PLAN_NAME


def _normalise_commvault_status(value: Any) -> str:
    if isinstance(value, str):
        text = value.strip()
        if text:
            return text
    if value is None:
        return "Unknown"
    text = str(value).strip()
    return text or "Unknown"


def _commvault_status_failed(status: str) -> bool:
    lowered = status.casefold()
    return any(keyword in lowered for keyword in _FAILURE_STATUS_KEYWORDS)


def _commvault_status_successful(status: str) -> bool:
    if _commvault_status_failed(status):
        return False
    lowered = status.casefold()
    return any(keyword in lowered for keyword in _SUCCESS_STATUS_KEYWORDS)


def _commvault_status_warning(status: str) -> bool:
    if _commvault_status_failed(status) or _commvault_status_successful(status):
        return False
    lowered = status.casefold()
    return any(keyword in lowered for keyword in _WARNING_STATUS_KEYWORDS)


def _summarise_commvault_plans(jobs: Iterable[Mapping[str, Any]]) -> dict[str, Any]:
    plan_records: dict[str, Any] = {}
    total_jobs = 0
    active_jobs = 0
    global_status = Counter()
    global_success = 0
    global_failure = 0
    unique_clients: set[str] = set()
    now = datetime.now(tz=UTC)

    for payload in jobs:
        if not isinstance(payload, Mapping):
            continue

        total_jobs += 1
        plan_display = _normalise_commvault_plan_name(payload.get("plan_name"))
        plan_key = plan_display.casefold()
        bucket = plan_records.setdefault(
            plan_key,
            {
                "plan_name": plan_display,
                "job_count": 0,
                "clients": Counter(),
                "applications": Counter(),
                "subclients": Counter(),
                "status_counts": Counter(),
                "success_count": 0,
                "failed_count": 0,
                "warning_count": 0,
                "active_jobs": 0,
                "retained_jobs": 0,
                "total_application_bytes": 0,
                "total_media_bytes": 0,
                "savings_total": 0.0,
                "savings_samples": 0,
                "latest_start": None,
                "latest_job": None,
            },
        )
        if bucket["plan_name"] == COMMVAULT_DEFAULT_PLAN_NAME and plan_display != COMMVAULT_DEFAULT_PLAN_NAME:
            bucket["plan_name"] = plan_display

        bucket["job_count"] += 1

        status_label = _normalise_commvault_status(payload.get("localized_status") or payload.get("status"))
        bucket["status_counts"][status_label] += 1
        global_status[status_label] += 1
        if _commvault_status_failed(status_label):
            bucket["failed_count"] += 1
            global_failure += 1
        elif _commvault_status_successful(status_label):
            bucket["success_count"] += 1
            global_success += 1
        elif _commvault_status_warning(status_label):
            bucket["warning_count"] += 1

        if _commvault_job_is_active(payload):
            bucket["active_jobs"] += 1
            active_jobs += 1

        client_label = payload.get("client_name") or payload.get("destination_client_name")
        if client_label:
            client_name = str(client_label).strip()
            if client_name:
                bucket["clients"][client_name] += 1
                unique_clients.add(client_name)

        application = payload.get("application_name")
        if application:
            app_name = str(application).strip()
            if app_name:
                bucket["applications"][app_name] += 1

        subclient_label = payload.get("subclient_name") or payload.get("backup_set_name")
        if subclient_label:
            subclient_name = str(subclient_label).strip()
            if subclient_name:
                bucket["subclients"][subclient_name] += 1

        retain_until = _parse_job_datetime(payload.get("retain_until"))
        if retain_until and retain_until > now:
            bucket["retained_jobs"] += 1

        try:
            bucket["total_application_bytes"] += int(payload.get("size_of_application_bytes") or 0)
        except (TypeError, ValueError):
            pass
        try:
            bucket["total_media_bytes"] += int(payload.get("size_on_media_bytes") or 0)
        except (TypeError, ValueError):
            pass

        try:
            savings_value = float(payload.get("percent_savings"))
        except (TypeError, ValueError):
            savings_value = None
        if savings_value is not None:
            bucket["savings_total"] += savings_value
            bucket["savings_samples"] += 1

        start_dt = _parse_job_datetime(payload.get("start_time"))
        if start_dt and (bucket["latest_start"] is None or start_dt > bucket["latest_start"]):
            bucket["latest_start"] = start_dt
            bucket["latest_job"] = {
                "job_id": payload.get("job_id"),
                "job_type": payload.get("job_type"),
                "status": payload.get("status"),
                "localized_status": payload.get("localized_status"),
                "client_name": payload.get("client_name"),
                "destination_client_name": payload.get("destination_client_name"),
                "subclient_name": payload.get("subclient_name"),
                "application_name": payload.get("application_name"),
                "backup_level_name": payload.get("backup_level_name"),
                "plan_name": plan_display,
                "start_time": payload.get("start_time"),
                "end_time": payload.get("end_time"),
                "size_of_application_bytes": payload.get("size_of_application_bytes"),
                "size_on_media_bytes": payload.get("size_on_media_bytes"),
                "percent_savings": payload.get("percent_savings"),
            }

    plans: list[dict[str, Any]] = []
    for bucket in plan_records.values():
        clients_counter: Counter[str] = bucket["clients"]
        applications_counter: Counter[str] = bucket["applications"]
        subclients_counter: Counter[str] = bucket["subclients"]
        status_counter: Counter[str] = bucket["status_counts"]
        savings_count = bucket["savings_samples"] or 0
        average_savings = bucket["savings_total"] / savings_count if savings_count else None
        job_count = bucket["job_count"] or 0

        plans.append(
            {
                "plan_name": bucket["plan_name"],
                "job_count": job_count,
                "client_count": len(clients_counter),
                "clients": [
                    {"name": name, "count": count}
                    for name, count in sorted(
                        clients_counter.items(), key=lambda item: (-item[1], item[0].casefold())
                    )
                ],
                "applications": [
                    {"name": name, "count": count}
                    for name, count in sorted(
                        applications_counter.items(), key=lambda item: (-item[1], item[0].casefold())
                    )
                ],
                "subclients": [
                    {"name": name, "count": count}
                    for name, count in sorted(
                        subclients_counter.items(), key=lambda item: (-item[1], item[0].casefold())
                    )
                ],
                "status_counts": [
                    {"status": name, "count": count}
                    for name, count in sorted(
                        status_counter.items(), key=lambda item: (-item[1], item[0].casefold())
                    )
                ],
                "success_count": bucket["success_count"],
                "failed_count": bucket["failed_count"],
                "warning_count": bucket["warning_count"],
                "active_jobs": bucket["active_jobs"],
                "retained_jobs": bucket["retained_jobs"],
                "total_application_bytes": bucket["total_application_bytes"],
                "total_media_bytes": bucket["total_media_bytes"],
                "average_savings_percent": average_savings,
                "latest_job": bucket["latest_job"],
                "latest_job_start": bucket["latest_start"].isoformat() if bucket["latest_start"] else None,
                "success_rate": (bucket["success_count"] / job_count * 100.0) if job_count else None,
                "failure_rate": (bucket["failed_count"] / job_count * 100.0) if job_count else None,
            }
        )

    plans.sort(key=lambda plan: (-plan["job_count"], plan["plan_name"].casefold()))

    total_success = sum(plan["success_count"] for plan in plans)
    total_failure = sum(plan["failed_count"] for plan in plans)

    return {
        "plans": plans,
        "total_plans": len(plans),
        "total_jobs": total_jobs,
        "active_jobs": active_jobs,
        "unique_clients": len(unique_clients),
        "status_counts": [
            {"status": name, "count": count}
            for name, count in sorted(global_status.items(), key=lambda item: (-item[1], item[0].casefold()))
        ],
        "success_rate": (total_success / total_jobs * 100.0) if total_jobs else None,
        "failure_rate": (total_failure / total_jobs * 100.0) if total_jobs else None,
    }


def _latest_commvault_start_time(jobs: Iterable[Mapping[str, Any]]) -> datetime | None:
    latest: datetime | None = None
    for job in jobs:
        start = _parse_job_datetime(job.get("start_time"))
        if start and (latest is None or start > latest):
            latest = start
    return latest


def _safe_commvault_client_id(value: Any) -> int | None:
    try:
        client_id = int(value)
    except (TypeError, ValueError):
        return None
    return client_id if client_id >= 0 else None


def _normalise_commvault_client_name(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _cached_commvault_clients() -> tuple[list[dict[str, Any]], list[Mapping[str, Any]], str | None]:
    cache = _load_commvault_backups()
    jobs_payload = cache.get("jobs")
    if not isinstance(jobs_payload, list):
        jobs_payload = []

    clients: dict[tuple[int | None, str], dict[str, Any]] = {}
    for raw in jobs_payload:
        if not isinstance(raw, Mapping):
            continue
        client_id = _safe_commvault_client_id(raw.get("client_id"))
        client_name = _normalise_commvault_client_name(raw.get("client_name"))
        dest_name = _normalise_commvault_client_name(raw.get("destination_client_name"))
        primary = client_name or dest_name
        display = dest_name or client_name or primary
        key_name = (primary or display or "").casefold()
        key = (client_id, key_name)
        entry = clients.get(key)
        if not entry:
            label = display or primary or (f"Client {client_id}" if client_id is not None else "Unnamed client")
            entry = {
                "client_id": client_id,
                "name": primary or label,
                "display_name": label,
                "name_variants": set(),
                "job_count": 0,
            }
            clients[key] = entry
        entry["job_count"] += 1
        for candidate in (client_name, dest_name, primary, display):
            if candidate:
                entry["name_variants"].add(candidate.casefold())

    client_list = list(clients.values())
    for entry in client_list:
        entry["name_variants"].add((entry["name"] or "").casefold())
        entry["name_variants"].add((entry["display_name"] or "").casefold())
    client_list.sort(key=lambda item: (-item["job_count"], (item["display_name"] or item["name"] or "").lower()))

    jobs = [job for job in jobs_payload if isinstance(job, Mapping)]
    generated_at = cache.get("generated_at")
    return client_list, jobs, generated_at


def _match_cached_commvault_client(identifier: str, clients: list[dict[str, Any]]) -> dict[str, Any]:
    ident = (identifier or "").strip()
    if not ident:
        raise HTTPException(status_code=400, detail="Client identifier is required")

    if ident.isdigit():
        client_id = int(ident)
        for entry in clients:
            if entry["client_id"] == client_id:
                return entry
        raise HTTPException(status_code=404, detail=f"No cached client with ID {client_id} found")

    needle = ident.casefold()
    exact = [entry for entry in clients if needle in entry["name_variants"]]
    if len(exact) == 1:
        return exact[0]
    if len(exact) > 1:
        names = ", ".join(sorted({entry["display_name"] or entry["name"] or str(entry["client_id"]) for entry in exact}))
        raise HTTPException(status_code=409, detail=f"Ambiguous client '{identifier}'. Matches: {names}")

    matches = [
        entry
        for entry in clients
        if any(needle in variant for variant in entry["name_variants"] if variant)
    ]
    if not matches:
        raise HTTPException(status_code=404, detail=f"No cached client matching '{identifier}' found")
    if len(matches) > 1:
        names = ", ".join(sorted({entry["display_name"] or entry["name"] or str(entry["client_id"]) for entry in matches}))
        raise HTTPException(status_code=409, detail=f"Ambiguous client '{identifier}'. Matches: {names}")
    return matches[0]


def _job_matches_cached_client(job: Mapping[str, Any], client_record: dict[str, Any]) -> bool:
    if client_record.get("client_id") is not None:
        job_id = _safe_commvault_client_id(job.get("client_id"))
        if job_id is not None and job_id == client_record["client_id"]:
            return True
    names = {
        _normalise_commvault_client_name(job.get("client_name")).casefold(),
        _normalise_commvault_client_name(job.get("destination_client_name")).casefold(),
    }
    names.discard("")
    if not names:
        return False
    variants = client_record.get("name_variants") or set()
    return any(name in variants for name in names)


def _build_cached_client_summary(record: dict[str, Any]) -> dict[str, Any]:
    return {
        "client_id": record.get("client_id"),
        "name": record.get("name"),
        "display_name": record.get("display_name"),
        "host_name": None,
        "os_name": None,
        "os_type": None,
        "os_subtype": None,
        "processor_type": None,
        "cpu_count": None,
        "is_media_agent": None,
        "is_virtual": None,
        "is_infrastructure": None,
        "is_commserve": None,
        "readiness_status": None,
        "last_ready_time": None,
        "sla_status_code": None,
        "sla_description": None,
        "agent_applications": [],
        "client_groups": [],
    }


def _build_cached_job_metrics(
    jobs: list[dict[str, Any]],
    *,
    since_hours: int,
    retained_only: bool,
    cache_generated_at: str | None,
) -> dict[str, Any]:
    def _to_int(value: Any) -> int:
        try:
            return int(value)
        except (TypeError, ValueError):
            return 0

    total_app = sum(max(0, _to_int(job.get("size_of_application_bytes"))) for job in jobs)
    total_media = sum(max(0, _to_int(job.get("size_on_media_bytes"))) for job in jobs)
    last_start: datetime | None = None
    for job in jobs:
        start_dt = _parse_job_datetime(job.get("start_time"))
        if start_dt and (last_start is None or start_dt > last_start):
            last_start = start_dt

    fetched_at = _parse_job_datetime(cache_generated_at) if cache_generated_at else None
    if fetched_at is None:
        fetched_at = datetime.now(tz=UTC)

    retain_cutoff = fetched_at if retained_only else None

    return {
        "window_hours": since_hours,
        "job_count": len(jobs),
        "total_application_bytes": total_app,
        "total_media_bytes": total_media,
        "last_job_start": last_start.isoformat() if last_start else None,
        "within_window": bool(jobs),
        "descending": True,
        "retain_cutoff": retain_cutoff.isoformat() if retain_cutoff else None,
        "retain_required": retained_only,
        "fetched_at": fetched_at.isoformat(),
    }


def _slugify_commvault_name(value: str, default: str = "commvault") -> str:
    text = (value or "").strip().lower()
    if not text:
        return default
    safe = re.sub(r"[^a-z0-9]+", "-", text)
    safe = re.sub(r"-+", "-", safe).strip("-")
    return safe or default


def _format_bytes_for_report(value: int | float | None) -> str:
    if not value:
        return "0 B"
    size = float(value)
    if size <= 0:
        return "0 B"
    units = ["B", "KB", "MB", "GB", "TB", "PB"]
    idx = 0
    while size >= 1024 and idx < len(units) - 1:
        size /= 1024
        idx += 1
    precision = 0 if size >= 100 else 1 if size >= 10 else 2
    return f"{size:.{precision}f} {units[idx]}"


def _format_minutes_label(value: int | None) -> str | None:
    if value is None:
        return None
    minutes = int(value)
    if minutes < 0:
        minutes = abs(minutes)
    if minutes == 0:
        return "0 minutes"
    if minutes % 1440 == 0:
        days = minutes // 1440
        return f"{days} day" if days == 1 else f"{days} days"
    if minutes % 60 == 0:
        hours = minutes // 60
        return f"{hours} hour" if hours == 1 else f"{hours} hours"
    return f"{minutes} minute" if minutes == 1 else f"{minutes} minutes"


def _plan_status_from_flag(flag: int | None) -> str | None:
    if flag is None:
        return None
    if flag == 0:
        return "Enabled"
    if flag == 1:
        return "Disabled"
    if flag == 2:
        return "Retired"
    return f"Flag {flag}"


def _label_commvault_plan_type(code: int | None, fallback: str | None = None) -> str | None:
    if code is None:
        return fallback
    label = COMMVAULT_PLAN_TYPE_LABELS.get(code)
    if label:
        return label
    if fallback:
        return fallback
    return str(code)


def _compute_commvault_server_metrics(jobs: list[dict[str, Any]]) -> dict[str, Any]:
    if not jobs:
        return {
            "job_count": 0,
            "plan_breakdown": [],
            "subclient_breakdown": [],
            "policy_breakdown": [],
            "total_application_bytes": 0,
            "total_media_bytes": 0,
            "retained_jobs": 0,
            "average_savings_percent": None,
            "average_reduction_ratio": None,
            "average_reduction_ratio_text": None,
            "savings_bytes": 0,
            "series": {"timeline": []},
            "latest_job": None,
            "latest_job_started_at": None,
            "next_retain_expiry": None,
        }

    plan_counts: Counter[str] = Counter()
    subclient_counts: Counter[str] = Counter()
    policy_counts: Counter[str] = Counter()
    total_application = 0
    total_media = 0
    savings_samples: list[float] = []
    retained_jobs = 0
    next_expiry: datetime | None = None
    latest_start: datetime | None = None
    latest_job: dict[str, Any] | None = None
    timeline: list[dict[str, Any]] = []

    def _as_int(value: Any) -> int:
        try:
            return int(value)
        except (TypeError, ValueError):
            return 0

    def _as_float(value: Any) -> float | None:
        try:
            return float(value)
        except (TypeError, ValueError):
            return None

    for job in jobs:
        plan = job.get("plan_name") or "Unassigned"
        subclient = job.get("subclient_name") or "Unspecified"
        policy = job.get("storage_policy_name") or "Unspecified"
        plan_counts[plan] += 1
        subclient_counts[subclient] += 1
        policy_counts[policy] += 1

        app_size = max(0, _as_int(job.get("size_of_application_bytes")))
        media_size = max(0, _as_int(job.get("size_on_media_bytes")))
        total_application += app_size
        total_media += media_size

        savings = _as_float(job.get("percent_savings"))
        if savings is not None:
            savings_samples.append(savings)

        retain_until = _parse_job_datetime(job.get("retain_until"))
        if retain_until:
            retained_jobs += 1
            if next_expiry is None or retain_until < next_expiry:
                next_expiry = retain_until

        start_dt = _parse_job_datetime(job.get("start_time"))
        if start_dt and (latest_start is None or start_dt > latest_start):
            latest_start = start_dt
            latest_job = job

        timeline.append(
            {
                "start_time": job.get("start_time"),
                "start_timestamp": start_dt.timestamp() if start_dt else None,
                "size_of_application_bytes": app_size,
                "size_on_media_bytes": media_size,
                "percent_savings": savings,
                "plan_name": job.get("plan_name"),
                "status": job.get("localized_status") or job.get("status"),
            }
        )

    timeline.sort(key=lambda item: item.get("start_timestamp") or 0)
    for entry in timeline:
        entry.pop("start_timestamp", None)

    average_savings = None
    if savings_samples:
        average_savings = sum(savings_samples) / len(savings_samples)

    reduction_ratio = None
    reduction_ratio_text = None
    if total_media > 0:
        reduction_ratio = total_application / total_media if total_application else 0
        if reduction_ratio and reduction_ratio > 0:
            reduction_ratio_text = f"{reduction_ratio:.1f}:1"

    savings_bytes = max(0, total_application - total_media)

    def _counts_to_rows(counter: Counter[str]) -> list[dict[str, Any]]:
        return [
            {"label": label, "restore_points": count}
            for label, count in counter.most_common()
        ]

    return {
        "job_count": len(jobs),
        "plan_breakdown": _counts_to_rows(plan_counts),
        "subclient_breakdown": _counts_to_rows(subclient_counts),
        "policy_breakdown": _counts_to_rows(policy_counts),
        "total_application_bytes": total_application,
        "total_media_bytes": total_media,
        "retained_jobs": retained_jobs,
        "average_savings_percent": average_savings,
        "average_reduction_ratio": reduction_ratio,
        "average_reduction_ratio_text": reduction_ratio_text,
        "savings_bytes": savings_bytes,
        "series": {"timeline": timeline},
        "latest_job": latest_job,
        "latest_job_started_at": latest_start.isoformat() if latest_start else None,
        "next_retain_expiry": next_expiry.isoformat() if next_expiry else None,
    }


def _load_commvault_server_data(
    client_identifier: str,
    *,
    job_limit: int,
    since_hours: int,
    retained_only: bool,
    refresh_cache: bool,
) -> tuple[dict[str, Any], dict[str, Any], list[dict[str, Any]], dict[str, Any]]:
    clients, jobs_payload, cache_generated_at = _cached_commvault_clients()
    if not clients:
        raise HTTPException(status_code=503, detail="Commvault cached backups are not available. Run an export first.")

    client_record = _match_cached_commvault_client(client_identifier, clients)

    cutoff = datetime.now(tz=UTC) - timedelta(hours=since_hours) if since_hours > 0 else None
    now_utc = datetime.now(tz=UTC)
    default_order = datetime.min.replace(tzinfo=UTC)

    selected: list[tuple[datetime | None, dict[str, Any]]] = []
    for raw in jobs_payload:
        if not _job_matches_cached_client(raw, client_record):
            continue

        start_dt = _parse_job_datetime(raw.get("start_time"))
        if cutoff and (start_dt is None or start_dt < cutoff):
            continue

        if retained_only:
            retain_dt = _parse_job_datetime(raw.get("retain_until"))
            if retain_dt is None or retain_dt <= now_utc:
                continue

        selected.append((start_dt, dict(raw)))

    selected.sort(key=lambda item: item[0] or default_order, reverse=True)
    jobs = [job for _, job in selected]
    if job_limit > 0:
        jobs = jobs[: job_limit]

    stats = _compute_commvault_server_metrics(jobs)
    metrics_payload = _build_cached_job_metrics(
        jobs,
        since_hours=max(0, since_hours),
        retained_only=retained_only,
        cache_generated_at=cache_generated_at,
    )
    summary_payload = _build_cached_client_summary(client_record)
    return summary_payload, metrics_payload, jobs, stats


def _format_iso_human(value: str | None) -> str:
    if not value:
        return "-"
    dt = _parse_job_datetime(value)
    if not dt:
        return "-"
    local_dt = dt.astimezone(AMS_TZ)
    tz_label = local_dt.tzname() or "CET"
    return f"{local_dt.strftime('%Y-%m-%d %H:%M')} {tz_label}"


def _format_percent(value: float | None) -> str:
    if value is None:
        return "-"
    return f"{value:.1f}%"


COMMVAULT_SERVER_EXPORT_COLUMNS: list[tuple[str, str]] = [
    ("job_id", "Job ID"),
    ("start", "Start"),
    ("status", "Status"),
    ("plan", "Plan"),
    ("app_size", "App Size"),
    ("media_size", "Media Size"),
    ("savings", "Savings"),
]


def _commvault_job_table_rows(jobs: list[dict[str, Any]]) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    for job in jobs:
        job_id = str(job.get("job_id") or "-")
        start_label = _format_iso_human(job.get("start_time"))
        status_label = job.get("localized_status") or job.get("status") or "-"
        plan_label = job.get("plan_name") or "-"
        app_size = _format_bytes_for_report(job.get("size_of_application_bytes"))
        media_size = _format_bytes_for_report(job.get("size_on_media_bytes"))
        savings_value = job.get("percent_savings")
        try:
            savings_float = float(savings_value) if savings_value is not None else None
        except (TypeError, ValueError):
            savings_float = None
        savings_label = _format_percent(savings_float)
        rows.append(
            {
                "job_id": job_id,
                "start": start_label,
                "status": status_label,
                "plan": plan_label,
                "app_size": app_size,
                "media_size": media_size,
                "savings": savings_label,
            }
        )
    return rows


def _render_commvault_server_text_report(
    summary: dict[str, Any],
    stats: dict[str, Any],
    jobs: list[dict[str, Any]],
    metrics_payload: dict[str, Any] | None,
) -> str:
    display_name = summary.get("display_name") or summary.get("name") or f"#{summary.get('client_id')}"
    generated = metrics_payload.get("fetched_at") if metrics_payload else None
    generated_label = _format_iso_human(generated)
    window_hours = metrics_payload.get("window_hours") if metrics_payload else None
    job_rows = _commvault_job_table_rows(jobs)

    lines = [
        f"Commvault server report for {display_name} (ID {summary.get('client_id')})",
        f"Generated: {generated_label}",
    ]
    if window_hours is not None:
        lines.append(f"Window: {window_hours}h | Restore points: {stats.get('job_count', 0)}")

    lines.append(
        f"Application data: {_format_bytes_for_report(stats.get('total_application_bytes'))}"
        f" | Media: {_format_bytes_for_report(stats.get('total_media_bytes'))}"
        f" | Savings: {_format_bytes_for_report(stats.get('savings_bytes'))}"
    )

    reduction_text = stats.get("average_reduction_ratio_text") or "-"
    savings_percent = _format_percent(stats.get("average_savings_percent"))
    lines.append(f"Average data reduction: {savings_percent} (≈ {reduction_text})")

    plan_rows = stats.get("plan_breakdown") or []
    if plan_rows:
        lines.append("Plan restore points:")
        for row in plan_rows:
            label = row.get("label") or row.get("plan") or "Unnamed"
            count = row.get("restore_points") or 0
            lines.append(f"  - {label}: {count}")

    subclient_rows = stats.get("subclient_breakdown") or []
    if subclient_rows:
        lines.append("Subclients:")
        for row in subclient_rows:
            label = row.get("label") or "Unnamed"
            count = row.get("restore_points") or 0
            lines.append(f"  - {label}: {count}")

    lines.append("")
    header = "  ".join(
        [
            f"{title:<18}" if title not in {"Job ID", "Plan"} else f"{title:<10}"
            for _, title in COMMVAULT_SERVER_EXPORT_COLUMNS
        ]
    )
    lines.append(header)
    lines.append("-" * len(header))
    for row in job_rows:
        lines.append(
            "  ".join(
                [
                    f"{row['job_id']:<10}",
                    f"{row['start']:<20}",
                    f"{row['status']:<18}",
                    f"{row['plan']:<18}",
                    f"{row['app_size']:>12}",
                    f"{row['media_size']:>12}",
                    f"{row['savings']:>8}",
                ]
            )
        )

    if not job_rows:
        lines.append("No jobs found for the selected parameters.")

    return "\n".join(lines)


def _render_commvault_server_markdown_report(
    summary: dict[str, Any],
    stats: dict[str, Any],
    jobs: list[dict[str, Any]],
    metrics_payload: dict[str, Any] | None,
) -> str:
    display_name = summary.get("display_name") or summary.get("name") or f"#{summary.get('client_id')}"
    generated = metrics_payload.get("fetched_at") if metrics_payload else None
    generated_label = _format_iso_human(generated)
    window_hours = metrics_payload.get("window_hours") if metrics_payload else None
    job_rows = _commvault_job_table_rows(jobs)

    lines = [f"# Commvault server report — {display_name}", ""]
    lines.append(f"*ID*: `{summary.get('client_id')}`")
    lines.append(f"*Generated*: {generated_label}")
    if window_hours is not None:
        lines.append(f"*Window*: {window_hours}h")
    lines.append(
        f"*Application data*: {_format_bytes_for_report(stats.get('total_application_bytes'))}"
        f" — *Media*: {_format_bytes_for_report(stats.get('total_media_bytes'))}"
        f" — *Savings*: {_format_bytes_for_report(stats.get('savings_bytes'))}"
    )
    reduction_text = stats.get("average_reduction_ratio_text") or "-"
    savings_percent = _format_percent(stats.get("average_savings_percent"))
    lines.append(f"*Average data reduction*: {savings_percent} (≈ {reduction_text})")
    lines.append("")

    plan_rows = stats.get("plan_breakdown") or []
    if plan_rows:
        lines.append("## Restore points by plan")
        for row in plan_rows:
            label = row.get("label") or row.get("plan") or "Unnamed"
            count = row.get("restore_points") or 0
            lines.append(f"- **{label}** — {count} restore point(s)")
        lines.append("")

    lines.append("## Jobs")
    header = " | ".join(title for _, title in COMMVAULT_SERVER_EXPORT_COLUMNS)
    lines.append(f"{header}")
    lines.append(" | ".join(["---"] * len(COMMVAULT_SERVER_EXPORT_COLUMNS)))
    for row in job_rows:
        lines.append(
            " | ".join(
                [
                    row["job_id"],
                    row["start"],
                    row["status"],
                    row["plan"],
                    row["app_size"],
                    row["media_size"],
                    row["savings"],
                ]
            )
        )
    if not job_rows:
        lines.append("No jobs found for the selected parameters.")

    return "\n".join(lines)


def _render_commvault_server_html_report(
    summary: dict[str, Any],
    stats: dict[str, Any],
    jobs: list[dict[str, Any]],
    metrics_payload: dict[str, Any] | None,
) -> str:
    display_name = summary.get("display_name") or summary.get("name") or f"#{summary.get('client_id')}"
    generated = metrics_payload.get("fetched_at") if metrics_payload else None
    generated_label = _format_iso_human(generated)
    window_hours = metrics_payload.get("window_hours") if metrics_payload else None
    job_rows = _commvault_job_table_rows(jobs)

    plan_rows = stats.get("plan_breakdown") or []
    plan_html = "".join(
        f"<li><strong>{html.escape(str(row.get('label') or row.get('plan') or 'Unnamed'))}</strong>: {row.get('restore_points') or 0}</li>"
        for row in plan_rows
    )

    job_rows_html = "".join(
        "<tr>" + "".join(
            f"<td>{html.escape(row[key])}</td>"
            for key, _ in COMMVAULT_SERVER_EXPORT_COLUMNS
        ) + "</tr>"
        for row in job_rows
    )

    if not job_rows_html:
        job_rows_html = "<tr><td colspan=\"7\">No jobs found for the selected parameters.</td></tr>"

    headings = "".join(f"<th>{html.escape(title)}</th>" for _, title in COMMVAULT_SERVER_EXPORT_COLUMNS)

    return f"""
<!doctype html>
<html>
  <head>
    <meta charset=\"utf-8\" />
    <title>Commvault report — {html.escape(display_name)}</title>
    <style>
      body {{ font-family: system-ui, sans-serif; margin: 24px; color: #0f172a; }}
      h1 {{ margin-top: 0; }}
      table {{ border-collapse: collapse; width: 100%; margin-top: 16px; }}
      th, td {{ border: 1px solid #cbd5f5; padding: 8px 10px; text-align: left; font-size: 13px; }}
      th {{ background: #eef2ff; }}
      caption {{ text-align: left; font-weight: 600; margin-bottom: 8px; }}
    </style>
  </head>
  <body>
    <h1>Commvault server report — {html.escape(display_name)}</h1>
    <p><strong>ID:</strong> {html.escape(str(summary.get('client_id')))}<br/>
       <strong>Generated:</strong> {html.escape(generated_label)}<br/>
       <strong>Window:</strong> {html.escape(str(window_hours) + 'h' if window_hours is not None else 'n/a')}</p>
    <p><strong>Application data:</strong> {_format_bytes_for_report(stats.get('total_application_bytes'))}<br/>
       <strong>Media:</strong> {_format_bytes_for_report(stats.get('total_media_bytes'))}<br/>
       <strong>Savings:</strong> {_format_bytes_for_report(stats.get('savings_bytes'))}<br/>
       <strong>Average reduction:</strong> {_format_percent(stats.get('average_savings_percent'))} (≈ {stats.get('average_reduction_ratio_text') or '-'})
    </p>
    {'<h2>Restore points by plan</h2><ul>' + plan_html + '</ul>' if plan_html else ''}
    <table>
      <caption>Commvault jobs</caption>
      <thead><tr>{headings}</tr></thead>
      <tbody>{job_rows_html}</tbody>
    </table>
  </body>
</html>
"""


def _render_commvault_server_docx(
    summary: dict[str, Any],
    stats: dict[str, Any],
    jobs: list[dict[str, Any]],
    metrics_payload: dict[str, Any] | None,
) -> BytesIO:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise HTTPException(status_code=500, detail="python-docx is required for DOCX export") from exc

    display_name = summary.get("display_name") or summary.get("name") or f"#{summary.get('client_id')}"
    generated = metrics_payload.get("fetched_at") if metrics_payload else None
    generated_label = _format_iso_human(generated)
    window_hours = metrics_payload.get("window_hours") if metrics_payload else None
    job_rows = _commvault_job_table_rows(jobs)

    doc = Document()
    doc.add_heading(f"Commvault server report — {display_name}", level=0)
    doc.add_paragraph(f"Client ID: {summary.get('client_id')}")
    doc.add_paragraph(f"Generated: {generated_label}")
    if window_hours is not None:
        doc.add_paragraph(f"Window: {window_hours}h")

    doc.add_paragraph(
        "Application data: "
        + _format_bytes_for_report(stats.get("total_application_bytes"))
        + "; Media: "
        + _format_bytes_for_report(stats.get("total_media_bytes"))
        + "; Savings: "
        + _format_bytes_for_report(stats.get("savings_bytes"))
    )
    doc.add_paragraph(
        "Average data reduction: "
        + _format_percent(stats.get("average_savings_percent"))
        + f" (≈ {stats.get('average_reduction_ratio_text') or '-'})"
    )

    plan_rows = stats.get("plan_breakdown") or []
    if plan_rows:
        doc.add_heading("Restore points by plan", level=1)
        for row in plan_rows:
            label = row.get("label") or row.get("plan") or "Unnamed"
            count = row.get("restore_points") or 0
            doc.add_paragraph(f"{label}: {count}", style="List Bullet")

    doc.add_heading("Jobs", level=1)
    table = doc.add_table(rows=1, cols=len(COMMVAULT_SERVER_EXPORT_COLUMNS))
    header_cells = table.rows[0].cells
    for idx, (_, title) in enumerate(COMMVAULT_SERVER_EXPORT_COLUMNS):
        header_cells[idx].text = title
    for row in job_rows:
        cells = table.add_row().cells
        for idx, (key, _) in enumerate(COMMVAULT_SERVER_EXPORT_COLUMNS):
            cells[idx].text = row.get(key, "")
    if not job_rows:
        row = table.add_row().cells
        row[0].text = "No jobs"

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _render_commvault_server_xlsx(
    summary: dict[str, Any],
    jobs: list[dict[str, Any]],
) -> BytesIO:
    try:
        from openpyxl import Workbook  # type: ignore
    except ImportError as exc:  # pragma: no cover - dependency managed elsewhere
        raise HTTPException(status_code=500, detail="openpyxl is required for XLSX export") from exc

    workbook = Workbook()
    sheet = workbook.active
    display_name = summary.get("display_name") or summary.get("name") or f"#{summary.get('client_id')}"
    sheet.title = display_name[:31]

    headers = [
        "job_id",
        "job_type",
        "status",
        "localized_status",
        "localized_operation",
        "client_name",
        "client_id",
        "destination_client_name",
        "subclient_name",
        "backup_set_name",
        "application_name",
        "backup_level_name",
        "plan_name",
        "client_groups",
        "storage_policy_name",
        "start_time",
        "end_time",
        "elapsed_seconds",
        "size_of_application_bytes",
        "size_on_media_bytes",
        "total_num_files",
        "percent_complete",
        "percent_savings",
        "average_throughput_gb_per_hr",
        "retain_until",
    ]
    sheet.append(headers)

    for job in jobs:
        row = dict(job)
        groups = row.get("client_groups")
        if isinstance(groups, list):
            row["client_groups"] = ";".join(groups)
        sheet.append([row.get(header, "") for header in headers])

    buffer = BytesIO()
    workbook.save(buffer)
    buffer.seek(0)
    return buffer


def _render_commvault_server_csv(jobs: list[dict[str, Any]]) -> StringIO:
    fieldnames = [
        "job_id",
        "job_type",
        "status",
        "localized_status",
        "localized_operation",
        "client_name",
        "client_id",
        "destination_client_name",
        "subclient_name",
        "backup_set_name",
        "application_name",
        "backup_level_name",
        "plan_name",
        "client_groups",
        "storage_policy_name",
        "start_time",
        "end_time",
        "elapsed_seconds",
        "size_of_application_bytes",
        "size_on_media_bytes",
        "total_num_files",
        "percent_complete",
        "percent_savings",
        "average_throughput_gb_per_hr",
        "retain_until",
    ]

    buffer = StringIO()
    writer = csv.DictWriter(buffer, fieldnames=fieldnames)
    writer.writeheader()
    for job in jobs:
        row = dict(job)
        groups = row.get("client_groups")
        if isinstance(groups, list):
            row["client_groups"] = ";".join(groups)
        writer.writerow({key: row.get(key, "") for key in fieldnames})
    buffer.seek(0)
    return buffer


def _mb_to_bytes(value: Any) -> int | None:
    number = _safe_int(value)
    if number is None:
        return None
    return number * 1024 * 1024


def _normalise_dedupe_ratio(value: Any) -> float | None:
    ratio = _safe_float(value)
    if ratio is None:
        return None
    if ratio <= 0:
        return None
    if ratio > 0 and ratio < 1:
        return ratio * 100.0
    return ratio


def _safe_int(value: Any) -> int | None:
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def _safe_float(value: Any) -> float | None:
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _load_commvault_storage() -> dict[str, Any]:
    path = _data_dir() / COMMVAULT_STORAGE_JSON
    if not path.exists():
        return {
            "pools": [],
            "generated_at": None,
            "total_cached": 0,
            "version": 1,
        }
    try:
        with path.open("r", encoding="utf-8") as handle:
            data = json.load(handle)
            if isinstance(data, Mapping):
                pools = list(data.get("pools") or [])
                generated = data.get("generated_at")
                total_cached = data.get("total_cached")
                try:
                    total_cached = int(total_cached)
                except (TypeError, ValueError):
                    total_cached = len(pools)
                return {
                    "pools": pools,
                    "generated_at": generated,
                    "total_cached": total_cached,
                    "version": data.get("version", 1),
                }
    except Exception:
        logger.exception("Failed to load cached Commvault storage pools", extra={"event": "commvault_storage_cache_error"})
    return {
        "pools": [],
        "generated_at": None,
        "total_cached": 0,
        "version": 1,
    }


def _summarise_storage_pool(
    pool: CommvaultStoragePool, details: Mapping[str, Any] | None
) -> dict[str, Any]:
    details = details or {}
    total_capacity_bytes = _mb_to_bytes(pool.total_capacity_mb)
    used_bytes = _mb_to_bytes(pool.size_on_disk_mb)
    free_bytes = _mb_to_bytes(pool.total_free_space_mb)
    usage_pct: float | None = None
    if total_capacity_bytes and used_bytes is not None:
        try:
            usage_pct = (used_bytes / total_capacity_bytes) * 100.0
        except ZeroDivisionError:
            usage_pct = None

    data_reduction = details.get("dataReduction") or details.get("dataReductionInfo")
    dedupe_ratio = None
    dedupe_savings_bytes = None
    if isinstance(data_reduction, Mapping):
        ratios = [
            data_reduction.get("dedupeRatio"),
            data_reduction.get("reductionRatio"),
            data_reduction.get("globalReductionPercent"),
        ]
        for candidate in ratios:
            dedupe_ratio = _normalise_dedupe_ratio(candidate)
            if dedupe_ratio is not None:
                break
        savings_keys = [
            data_reduction.get("dedupeSavingsInBytes"),
            data_reduction.get("totalSavingsInBytes"),
        ]
        for candidate in savings_keys:
            dedupe_savings_bytes = _safe_int(candidate)
            if dedupe_savings_bytes is not None:
                break

    logical_capacity_bytes = None
    if isinstance(details.get("usage"), Mapping):
        logical_capacity_bytes = _safe_int(details["usage"].get("logicalSpaceUsedInBytes"))

    return {
        "pool_id": pool.pool_id,
        "name": pool.name,
        "status": pool.status,
        "storage_type_code": pool.storage_type_code,
        "storage_pool_type_code": pool.storage_pool_type_code,
        "storage_sub_type_code": pool.storage_sub_type_code,
        "storage_policy_id": pool.storage_policy_id,
        "storage_policy_name": pool.storage_policy_name,
        "region_name": pool.region_display_name or pool.region_name,
        "total_capacity_bytes": total_capacity_bytes,
        "used_bytes": used_bytes,
        "free_bytes": free_bytes,
        "logical_capacity_bytes": logical_capacity_bytes,
        "usage_percent": usage_pct,
        "dedupe_ratio": dedupe_ratio,
        "dedupe_savings_bytes": dedupe_savings_bytes,
        "number_of_nodes": pool.number_of_nodes,
        "is_archive_storage": pool.is_archive_storage,
        "cloud_storage_class_name": pool.cloud_storage_class_name,
        "library_ids": list(pool.library_ids),
        "raw": pool.raw,
        "details": details,
    }


def _enrich_commvault_plan_row(row: dict[str, Any], detail: Mapping[str, Any]) -> None:
    plan_section = detail.get("plan") if isinstance(detail.get("plan"), Mapping) else detail
    summary = plan_section.get("summary") if isinstance(plan_section.get("summary"), Mapping) else {}
    plan_info = summary.get("plan") if isinstance(summary.get("plan"), Mapping) else {}

    plan_type_code = _safe_int(plan_info.get("planType") or summary.get("type"))
    if plan_type_code is not None:
        row["plan_type"] = _label_commvault_plan_type(plan_type_code, row.get("plan_type"))
    elif isinstance(row.get("plan_type"), str) and row["plan_type"].isdigit():
        row["plan_type"] = _label_commvault_plan_type(_safe_int(row["plan_type"]), row["plan_type"])

    status_flag = _safe_int(summary.get("planStatusFlag"))
    status_label = _plan_status_from_flag(status_flag)
    if status_label:
        row["status"] = status_label

    rpo_minutes = _safe_int(summary.get("rpoInMinutes"))
    if rpo_minutes is None:
        rpo_minutes = _safe_int(summary.get("slaInMinutes"))
    rpo_label = _format_minutes_label(rpo_minutes)
    if rpo_label:
        row["rpo"] = rpo_label

    if not row.get("associated_entities"):
        assoc = plan_section.get("associatedEntitiesCount")
        total = 0
        if isinstance(assoc, Mapping):
            all_clients = assoc.get("allClients")
            if isinstance(all_clients, Sequence):
                for item in all_clients:
                    if not isinstance(item, Mapping):
                        continue
                    count = _safe_int(item.get("count"))
                    if count:
                        total += count
        if total:
            row["associated_entities"] = total

    if not row.get("copy_count"):
        num_copies = _safe_int(plan_section.get("numCopies")) or _safe_int(summary.get("numCopies"))
        if num_copies:
            row["copy_count"] = num_copies


def _refresh_commvault_plans_sync(limit: int | None = None, plan_type: str | None = None) -> dict[str, Any]:
    fetch_limit = 500 if limit is None else max(1, limit)

    with _commvault_plans_lock:
        with task_logging("commvault.plans.refresh", limit=fetch_limit, plan_type=plan_type) as task_log:
            try:
                client = _commvault_client_from_env()
            except RuntimeError as exc:
                raise HTTPException(status_code=500, detail=str(exc)) from exc

            try:
                plans = client.list_plans(limit=fetch_limit, plan_type=plan_type)
            except CommvaultError as exc:
                raise HTTPException(status_code=502, detail=f"Commvault error: {exc}") from exc

            plan_details: dict[int, Mapping[str, Any]] = {}
            for plan in plans:
                if plan.plan_id is None:
                    continue
                try:
                    plan_details[plan.plan_id] = client.get_plan_details(plan.plan_id)
                except CommvaultError as exc:
                    logger.warning(
                        "Failed to load plan details",
                        extra={"event": "commvault_plan_detail_error", "plan_id": plan.plan_id, "error": str(exc)},
                    )

            rows: list[dict[str, Any]] = []
            for plan in plans:
                row = _serialise_commvault_plan(plan)
                detail = plan_details.get(plan.plan_id or -1)
                if detail:
                    _enrich_commvault_plan_row(row, detail)
                rows.append(row)

            plan_types = sorted(
                {
                    str(row.get("plan_type")).strip()
                    for row in rows
                    if isinstance(row.get("plan_type"), str) and str(row.get("plan_type")).strip()
                },
                key=str.casefold,
            )

            payload: dict[str, Any] = {
                "plans": rows,
                "generated_at": datetime.now(tz=UTC).isoformat(),
                "total_cached": len(rows),
                "plan_types": plan_types,
                "version": 1,
            }
            _write_commvault_plans_json(payload)
            task_log.add_success(plans=len(rows), plan_types=len(plan_types))
            return payload


def _refresh_commvault_storage_sync() -> dict[str, Any]:
    with _commvault_storage_lock:
        with task_logging("commvault.storage.refresh") as task_log:
            try:
                client = _commvault_client_from_env()
            except RuntimeError as exc:
                raise HTTPException(status_code=500, detail=str(exc)) from exc

            try:
                pools = list(client.list_storage_pools())
            except CommvaultError as exc:
                raise HTTPException(status_code=502, detail=f"Commvault error: {exc}") from exc

            summaries: list[dict[str, Any]] = []
            for pool in pools:
                details_payload: Mapping[str, Any] | None = None
                try:
                    details = client.get_storage_pool_details(pool.pool_id, summary=pool)
                    details_payload = details.details
                except CommvaultError as exc:
                    logger.warning(
                        "Failed to fetch storage pool details",
                        extra={"event": "commvault_storage_pool_details_error", "pool_id": pool.pool_id, "error": str(exc)},
                    )
                summaries.append(_summarise_storage_pool(pool, details_payload))

            payload: dict[str, Any] = {
                "pools": summaries,
                "generated_at": datetime.now(tz=UTC).isoformat(),
                "total_cached": len(summaries),
                "version": 1,
            }
            _write_commvault_storage_json(payload)
            task_log.add_success(pools=len(summaries))
            return payload


def _refresh_commvault_backups_sync(limit: int, since_hours: int) -> dict[str, Any]:
    lookback = max(since_hours, 0)
    since = datetime.now(tz=UTC) - timedelta(hours=lookback) if lookback > 0 else None

    with _commvault_backups_lock:
        with task_logging("commvault.backups.refresh", limit=limit, since_hours=since_hours) as task_log:
            try:
                client = _commvault_client_from_env()
            except RuntimeError as exc:
                raise HTTPException(status_code=500, detail=str(exc)) from exc

            cache_snapshot = _load_commvault_backups()
            jobs_by_id: dict[int, dict[str, Any]] = {}
            prev_digests: dict[int, str] = {}
            for job in cache_snapshot.get("jobs", []):
                job_id = job.get("job_id")
                if job_id is None:
                    continue
                try:
                    key = int(job_id)
                except (TypeError, ValueError):
                    continue
                payload = dict(job)
                jobs_by_id[key] = payload
                prev_digests[key] = _commvault_job_digest(payload)

            tracked_job_ids = {
                job_id for job_id, payload in jobs_by_id.items() if _commvault_job_is_active(payload)
            }
            latest_cached_start = _latest_commvault_start_time(jobs_by_id.values())

            fetch_cutoff = since
            if latest_cached_start and (fetch_cutoff is None or latest_cached_start < fetch_cutoff):
                fetch_cutoff = latest_cached_start

            tracked_min_time: datetime | None = None
            if tracked_job_ids:
                for job_id in tracked_job_ids:
                    payload = jobs_by_id.get(job_id)
                    if not isinstance(payload, Mapping):
                        continue
                    start_dt = _parse_job_datetime(payload.get("start_time"))
                    if start_dt and (tracked_min_time is None or start_dt < tracked_min_time):
                        tracked_min_time = start_dt
            if tracked_min_time and (fetch_cutoff is None or tracked_min_time < fetch_cutoff):
                fetch_cutoff = tracked_min_time

            cached_recent_count: int | None = None
            if fetch_cutoff is not None:
                recent_counter = 0
                for payload in jobs_by_id.values():
                    if not isinstance(payload, Mapping):
                        continue
                    start_dt = _parse_job_datetime(payload.get("start_time"))
                    if start_dt and start_dt >= fetch_cutoff:
                        recent_counter += 1
                cached_recent_count = recent_counter

            try:
                job_list = _collect_commvault_jobs_for_ui(
                    client,
                    limit=limit,
                    offset=0,
                    since=since,
                    tracked_job_ids=tracked_job_ids,
                    latest_start=latest_cached_start,
                    cached_recent_count=cached_recent_count,
                    cutoff_start=fetch_cutoff,
                )
            except CommvaultError as exc:
                raise HTTPException(status_code=502, detail=f"Commvault error: {exc}") from exc

            jobs = list(job_list.jobs)
            if since:
                jobs = [job for job in jobs if job.start_time and job.start_time >= since]

            rows = [_serialise_commvault_job(job) for job in jobs]
            new_jobs_count = 0
            updated_jobs_count = 0
            for row in rows:
                job_id = row.get("job_id")
                if job_id is None:
                    continue
                try:
                    key = int(job_id)
                except (TypeError, ValueError):
                    continue
                digest = _commvault_job_digest(row)
                previous_digest = prev_digests.get(key)
                if previous_digest is None:
                    new_jobs_count += 1
                elif previous_digest != digest:
                    updated_jobs_count += 1
                jobs_by_id[key] = row
                prev_digests[key] = digest

            merged_jobs = sorted(jobs_by_id.values(), key=lambda job: _parse_job_datetime(job.get("start_time")) or datetime.min.replace(tzinfo=UTC), reverse=True)

            cache_payload: dict[str, Any] = {
                "jobs": merged_jobs,
                "generated_at": datetime.now(tz=UTC).isoformat(),
                "total_cached": len(merged_jobs),
                "version": 2,
                "last_refresh_since_hours": since_hours,
                "last_refresh_limit": limit,
                "last_refresh_stats": {
                    "new_jobs": new_jobs_count,
                    "updated_jobs": updated_jobs_count,
                },
            }
            _write_commvault_backups_json(cache_payload)

            filtered_jobs = _filter_commvault_jobs(merged_jobs, since_hours)
            task_log.add_success(jobs=len(filtered_jobs), cached=len(merged_jobs))
            return {
                "jobs": filtered_jobs,
                "generated_at": cache_payload["generated_at"],
                "total_cached": cache_payload["total_cached"],
                "returned": len(filtered_jobs),
                "since_hours": since_hours,
                "limit": limit,
                "new_jobs": new_jobs_count,
                "updated_jobs": updated_jobs_count,
            }

SUGGESTIONS_FILE = "suggestions.csv"
SUGGESTION_FIELDS = [
    "id",
    "title",
    "summary",
    "classification",
    "status",
    "likes",
    "created_at",
    "updated_at",
    "comments",
]
SUGGESTION_CLASSIFICATIONS = {
    "Must have": {"color": "#7c3aed", "letter": "M"},
    "Should have": {"color": "#2563eb", "letter": "S"},
    "Could have": {"color": "#fbbf24", "letter": "C"},
    "Nice to have": {"color": "#22c55e", "letter": "N"},
    "Should not have": {"color": "#ef4444", "letter": "X"},
}
SUGGESTION_STATUSES = ["new", "accepted", "in progress", "done", "denied"]
_suggestions_lock = Lock()

ENV_SETTING_FIELDS: list[dict[str, Any]] = [
    # Zabbix
    {
        "key": "ZABBIX_API_URL",
        "label": "Zabbix API URL",
        "secret": False,
        "placeholder": "https://zabbix.example.com/api_jsonrpc.php",
        "category": "zabbix",
    },
    {
        "key": "ZABBIX_HOST",
        "label": "Zabbix Host",
        "secret": False,
        "placeholder": "https://zabbix.example.com",
        "category": "zabbix",
    },
    {
        "key": "ZABBIX_WEB_URL",
        "label": "Zabbix Web URL",
        "secret": False,
        "placeholder": "https://zabbix.example.com",
        "category": "zabbix",
    },
    {
        "key": "ZABBIX_API_TOKEN",
        "label": "Zabbix API Token",
        "secret": True,
        "placeholder": "paste-token-here",
        "category": "zabbix",
    },
    {
        "key": "ZABBIX_SEVERITIES",
        "label": "Zabbix Severities",
        "secret": False,
        "placeholder": "2,3,4",
        "category": "zabbix",
    },
    {
        "key": "ZABBIX_GROUP_ID",
        "label": "Zabbix Group ID",
        "secret": False,
        "placeholder": "optional",
        "category": "zabbix",
    },
    # NetBox & Atlassian
    {
        "key": "NETBOX_URL",
        "label": "NetBox URL",
        "secret": False,
        "placeholder": "https://netbox.example.com",
        "category": "net-atlassian",
    },
    {
        "key": "NETBOX_TOKEN",
        "label": "NetBox API Token",
        "secret": True,
        "placeholder": "paste-token-here",
        "category": "net-atlassian",
    },
    {
        "key": "NETBOX_DEBUG",
        "label": "NetBox Debug Logging",
        "secret": False,
        "placeholder": "0",
        "category": "net-atlassian",
    },
    {
        "key": "NETBOX_EXTRA_HEADERS",
        "label": "NetBox Extra Headers",
        "secret": False,
        "placeholder": "Key=Value;Other=Value",
        "category": "net-atlassian",
    },
    {
        "key": "NETBOX_DATA_DIR",
        "label": "NetBox Data Directory",
        "secret": False,
        "placeholder": "data",
        "category": "net-atlassian",
    },
    {
        "key": "ATLASSIAN_BASE_URL",
        "label": "Atlassian Base URL",
        "secret": False,
        "placeholder": "https://your-domain.atlassian.net",
        "category": "net-atlassian",
    },
    {
        "key": "ATLASSIAN_EMAIL",
        "label": "Atlassian Email",
        "secret": False,
        "placeholder": "user@example.com",
        "category": "net-atlassian",
    },
    {
        "key": "ATLASSIAN_API_TOKEN",
        "label": "Atlassian API Token",
        "secret": True,
        "placeholder": "paste-token-here",
        "category": "net-atlassian",
    },
    {
        "key": "CONFLUENCE_CMDB_PAGE_ID",
        "label": "Confluence CMDB Page ID",
        "secret": False,
        "placeholder": "981533033",
        "category": "net-atlassian",
    },
    {
        "key": "CONFLUENCE_DEVICES_PAGE_ID",
        "label": "Confluence Devices Page ID",
        "secret": False,
        "placeholder": "optional",
        "category": "net-atlassian",
    },
    {
        "key": "CONFLUENCE_VMS_PAGE_ID",
        "label": "Confluence VMs Page ID",
        "secret": False,
        "placeholder": "optional",
        "category": "net-atlassian",
    },
    {
        "key": "CONFLUENCE_ENABLE_TABLE_FILTER",
        "label": "Enable Table Filter Macro",
        "secret": False,
        "placeholder": "0 or 1",
        "category": "net-atlassian",
    },
    {
        "key": "CONFLUENCE_ENABLE_TABLE_SORT",
        "label": "Enable Table Sort Macro",
        "secret": False,
        "placeholder": "0 or 1",
        "category": "net-atlassian",
    },
    # Chat providers
    {"key": "OPENAI_API_KEY", "label": "OpenAI API Key", "secret": True, "placeholder": "sk-...", "category": "chat"},
    {
        "key": "CHAT_DEFAULT_MODEL_OPENAI",
        "label": "OpenAI Default Model",
        "secret": False,
        "placeholder": "gpt-4o-mini",
        "category": "chat",
    },
    {
        "key": "OPENROUTER_API_KEY",
        "label": "OpenRouter API Key",
        "secret": True,
        "placeholder": "or-...",
        "category": "chat",
    },
    {
        "key": "CHAT_DEFAULT_MODEL_OPENROUTER",
        "label": "OpenRouter Default Model",
        "secret": False,
        "placeholder": "openrouter/auto",
        "category": "chat",
    },
    {
        "key": "ANTHROPIC_API_KEY",
        "label": "Anthropic API Key",
        "secret": True,
        "placeholder": "api-key",
        "category": "chat",
    },
    {
        "key": "CHAT_DEFAULT_MODEL_CLAUDE",
        "label": "Anthropic Default Model",
        "secret": False,
        "placeholder": "claude-3-5-sonnet",
        "category": "chat",
    },
    {
        "key": "GOOGLE_API_KEY",
        "label": "Google (Gemini) API Key",
        "secret": True,
        "placeholder": "AIza...",
        "category": "chat",
    },
    {
        "key": "CHAT_DEFAULT_MODEL_GEMINI",
        "label": "Gemini Default Model",
        "secret": False,
        "placeholder": "gemini-1.5-flash",
        "category": "chat",
    },
    {
        "key": "CHAT_DEFAULT_PROVIDER",
        "label": "Default Chat Provider",
        "secret": False,
        "placeholder": "openai",
        "category": "chat",
    },
    {
        "key": "CHAT_DEFAULT_TEMPERATURE",
        "label": "Default Temperature",
        "secret": False,
        "placeholder": "0.2",
        "category": "chat",
    },
    {
        "key": "CHAT_SYSTEM_PROMPT",
        "label": "System Instructions",
        "secret": False,
        "placeholder": "Optional system prompt",
        "category": "chat",
    },
    # Export & reporting
    {
        "key": "NETBOX_XLSX_ORDER_FILE",
        "label": "Column Order Template",
        "secret": False,
        "placeholder": "path/to/column_order.xlsx",
        "category": "export",
    },
    # API & Web UI
    {"key": "LOG_LEVEL", "label": "Log Level", "secret": False, "placeholder": "INFO", "category": "api"},
    {
        "key": "ENREACH_API_TOKEN",
        "label": "Enreach API Token",
        "secret": True,
        "placeholder": "optional",
        "category": "api",
    },
    {
        "key": "ENREACH_UI_PASSWORD",
        "label": "UI Password",
        "secret": True,
        "placeholder": "optional",
        "category": "api",
    },
    {
        "key": "ENREACH_UI_SECRET",
        "label": "UI Session Secret",
        "secret": True,
        "placeholder": "auto-generated if empty",
        "category": "api",
    },
    {
        "key": "ENREACH_SSL_CERTFILE",
        "label": "SSL Certificate File",
        "secret": False,
        "placeholder": "certs/localhost.pem",
        "category": "api",
    },
    {
        "key": "ENREACH_SSL_KEYFILE",
        "label": "SSL Key File",
        "secret": False,
        "placeholder": "certs/localhost-key.pem",
        "category": "api",
    },
    {
        "key": "ENREACH_SSL_KEY_PASSWORD",
        "label": "SSL Key Password",
        "secret": True,
        "placeholder": "optional",
        "category": "api",
    },
    {"key": "UI_THEME_DEFAULT", "label": "Default Theme", "secret": False, "placeholder": "nebula", "category": "api"},
    # Backup
    {"key": "BACKUP_ENABLE", "label": "Backup Enabled", "secret": False, "placeholder": "1", "category": "backup"},
    {"key": "BACKUP_TYPE", "label": "Backup Type", "secret": False, "placeholder": "local", "category": "backup"},
    {
        "key": "BACKUP_HOST",
        "label": "Backup Host",
        "secret": False,
        "placeholder": "backup.example.com",
        "category": "backup",
    },
    {"key": "BACKUP_PORT", "label": "Backup Port", "secret": False, "placeholder": "22", "category": "backup"},
    {
        "key": "BACKUP_USERNAME",
        "label": "Backup Username",
        "secret": False,
        "placeholder": "backup_user",
        "category": "backup",
    },
    {
        "key": "BACKUP_PASSWORD",
        "label": "Backup Password",
        "secret": True,
        "placeholder": "password",
        "category": "backup",
    },
    {
        "key": "BACKUP_PRIVATE_KEY_PATH",
        "label": "Private Key Path",
        "secret": False,
        "placeholder": "~/.ssh/id_rsa",
        "category": "backup",
    },
    {
        "key": "BACKUP_REMOTE_PATH",
        "label": "Remote Path",
        "secret": False,
        "placeholder": "/backups/enreach-tools",
        "category": "backup",
    },
    {
        "key": "BACKUP_LOCAL_PATH",
        "label": "Local Backup Path",
        "secret": False,
        "placeholder": "backups",
        "category": "backup",
    },
    {
        "key": "BACKUP_CREATE_TIMESTAMPED_DIRS",
        "label": "Create Timestamped Directories",
        "secret": False,
        "placeholder": "false",
        "category": "backup",
    },
]


def _write_log(msg: str) -> None:
    try:
        ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        with LOG_PATH.open("a", encoding="utf-8") as fh:
            fh.write(f"[{ts}] {msg.rstrip()}\n")
    except Exception:
        # Logging failures should never crash the API
        pass


def _suggestions_path() -> Path:
    path = _csv_path(SUGGESTIONS_FILE)
    if not path.parent.exists():
        path.parent.mkdir(parents=True, exist_ok=True)
    if not path.exists():
        with path.open("w", newline="", encoding="utf-8") as fh:
            writer = csv.DictWriter(fh, fieldnames=SUGGESTION_FIELDS)
            writer.writeheader()
    return path


def _safe_classification(value: str | None) -> str:
    if not value:
        return "Could have"
    raw = str(value).strip().lower()
    for name in SUGGESTION_CLASSIFICATIONS:
        if raw == name.lower():
            return name
    return "Could have"


def _safe_status(value: str | None) -> str:
    if not value:
        return "new"
    raw = str(value).strip().lower()
    return raw if raw in SUGGESTION_STATUSES else "new"


def _require_classification(value: str | None) -> str:
    if value is None or str(value).strip() == "":
        return "Could have"
    raw = str(value).strip().lower()
    for name in SUGGESTION_CLASSIFICATIONS:
        if raw == name.lower():
            return name
    raise HTTPException(status_code=400, detail=f"Invalid classification: {value}")


def _require_status(value: str | None) -> str:
    if value is None or str(value).strip() == "":
        return "new"
    raw = str(value).strip().lower()
    if raw not in SUGGESTION_STATUSES:
        raise HTTPException(status_code=400, detail=f"Invalid status: {value}")
    return raw


def _now_iso() -> str:
    return datetime.now(UTC).replace(microsecond=0).isoformat()


def _decorate_suggestion(item: dict[str, Any]) -> dict[str, Any]:
    data = {**item}
    data["classification"] = _safe_classification(data.get("classification"))
    data["status"] = _safe_status(data.get("status"))
    data.setdefault("summary", "")
    try:
        data["likes"] = int(data.get("likes") or 0)
    except Exception:
        data["likes"] = 0
    comments_raw = data.get("comments") or []
    if isinstance(comments_raw, str):
        try:
            parsed = json.loads(comments_raw)
            comments = [c for c in parsed if isinstance(c, dict)] if isinstance(parsed, list) else []
        except json.JSONDecodeError:
            comments = []
    elif isinstance(comments_raw, list):
        comments = [c for c in comments_raw if isinstance(c, dict)]
    else:
        comments = []
    data["comments"] = comments
    meta = SUGGESTION_CLASSIFICATIONS.get(data["classification"], {})
    data["classification_color"] = meta.get("color")
    data["classification_letter"] = meta.get("letter")
    data["status_label"] = data["status"].title()
    return data


def _load_suggestions() -> list[dict[str, Any]]:
    path = _suggestions_path()
    try:
        sql = "SELECT * FROM read_csv_auto(?, header=True)"
        df = duckdb.query(sql, params=[path.as_posix()]).df()
    except Exception:
        df = pd.DataFrame(columns=SUGGESTION_FIELDS)
    if df.empty:
        return []
    for col in SUGGESTION_FIELDS:
        if col not in df.columns:
            df[col] = None
    df = df.replace([np.inf, -np.inf], np.nan)
    df = df.astype(object).where(pd.notnull(df), None)
    out: list[dict[str, Any]] = []
    for _, row in df.iterrows():
        item = {k: row.get(k) for k in SUGGESTION_FIELDS}
        out.append(_decorate_suggestion(item))

    # Newest first
    def _sort_key(it: dict[str, Any]):
        try:
            raw = str(it.get("created_at") or "")
            if raw.endswith("Z"):
                raw = raw[:-1]
            return datetime.fromisoformat(raw)
        except Exception:
            return datetime.min

    out.sort(key=_sort_key, reverse=True)
    return out


def _write_suggestions(rows: list[dict[str, Any]]) -> None:
    path = _suggestions_path()
    with path.open("w", newline="", encoding="utf-8") as fh:
        writer = csv.DictWriter(fh, fieldnames=SUGGESTION_FIELDS, extrasaction="ignore")
        writer.writeheader()
        for row in rows:
            payload = row.copy()
            payload["likes"] = int(payload.get("likes") or 0)
            payload["classification"] = _safe_classification(payload.get("classification"))
            payload["status"] = _safe_status(payload.get("status"))
            comments = payload.get("comments") or []
            if not isinstance(comments, list):
                comments = []
            payload["comments"] = json.dumps(comments, ensure_ascii=False)
            writer.writerow(payload)


def _env_file_path() -> Path:
    env_path = load_env()
    if not env_path.exists():
        env_path.touch()
    return env_path


def _read_env_values() -> dict[str, str]:
    env_path = _env_file_path()
    values = dotenv_values(env_path)
    # dotenv_values returns OrderedDict[str, Optional[str]]
    clean: dict[str, str] = {}
    for key, value in values.items():
        if value is None:
            continue
        clean[str(key)] = str(value)
    return clean


def _read_env_defaults() -> dict[str, str]:
    example_path = project_root() / ".env.example"
    if not example_path.exists():
        return {}
    values = dotenv_values(example_path)
    return {str(k): str(v) for k, v in values.items() if v is not None}


def _write_env_value(key: str, value: str | None) -> None:
    key = key.strip()
    env_path = _env_file_path()
    lines = env_path.read_text().splitlines() if env_path.exists() else []
    written = False
    new_lines: list[str] = []
    for line in lines:
        stripped = line.strip()
        if not stripped or stripped.startswith("#") or "=" not in line:
            new_lines.append(line)
            continue
        current_key, _, _ = line.partition("=")
        if current_key.strip() == key:
            if value is None:
                written = True
                continue  # remove line
            new_lines.append(f"{key}={value}")
            written = True
        else:
            new_lines.append(line)
    if not written and value is not None:
        if new_lines and new_lines[-1] != "":
            new_lines.append("")
        new_lines.append(f"{key}={value}")
    env_path.write_text("\n".join(new_lines).rstrip() + "\n")
    load_env(override=True)


def _chat_default_temperature() -> float | None:
    raw = os.getenv("CHAT_DEFAULT_TEMPERATURE", "").strip().lower()
    if raw in {"", "default", "auto"}:
        return None
    try:
        return float(raw)
    except Exception:
        return None


class SuggestionCreate(BaseModel):
    title: str
    summary: str | None = ""
    classification: str | None = None


class SuggestionUpdate(BaseModel):
    title: str | None = None
    summary: str | None = None
    classification: str | None = None
    status: str | None = None


class SuggestionLikeRequest(BaseModel):
    delta: int = 1


class SuggestionCommentCreate(BaseModel):
    text: str


class EnvUpdateRequest(BaseModel):
    key: str
    value: str | None = None


class EnvResetRequest(BaseModel):
    key: str


def _suggestion_meta() -> dict[str, Any]:
    classifications = [
        {
            "name": name,
            "color": meta.get("color"),
            "letter": meta.get("letter"),
        }
        for name, meta in SUGGESTION_CLASSIFICATIONS.items()
    ]
    statuses = [
        {
            "value": value,
            "label": value.title(),
        }
        for value in SUGGESTION_STATUSES
    ]
    return {"classifications": classifications, "statuses": statuses}


@app.get("/suggestions")
def suggestions_list() -> dict:
    with _suggestions_lock:
        items = _load_suggestions()
    meta = meta_to_dto(_suggestion_meta())
    dto = SuggestionListDTO(
        items=suggestions_to_dto(items),
        total=len(items),
        meta=meta,
    )
    return dto.dict_clean()


@app.get("/suggestions/{suggestion_id}")
def suggestions_detail(suggestion_id: str) -> dict:
    with _suggestions_lock:
        items = _load_suggestions()
    for item in items:
        if str(item.get("id")) == suggestion_id:
            meta = meta_to_dto(_suggestion_meta())
            dto = SuggestionItemDTO(
                item=suggestion_to_dto(item),
                meta=meta,
            )
            return dto.dict_clean()
    raise HTTPException(status_code=404, detail="Suggestion not found")


@app.post("/suggestions")
def suggestions_create(req: SuggestionCreate) -> dict:
    title = (req.title or "").strip()
    if not title:
        raise HTTPException(status_code=400, detail="Title is required")
    summary = (req.summary or "").strip()
    classification = _require_classification(req.classification)
    now = _now_iso()
    item = {
        "id": uuid.uuid4().hex,
        "title": title,
        "summary": summary,
        "classification": classification,
        "status": "new",
        "likes": 0,
        "created_at": now,
        "updated_at": now,
        "comments": [],
    }
    with _suggestions_lock:
        rows = _load_suggestions()
        rows.append(item)
        _write_suggestions(rows)
    decorated = _decorate_suggestion(item)
    dto = SuggestionItemDTO(item=suggestion_to_dto(decorated))
    return dto.dict_clean()


@app.put("/suggestions/{suggestion_id}")
def suggestions_update(suggestion_id: str, req: SuggestionUpdate) -> dict:
    with _suggestions_lock:
        rows = _load_suggestions()
        target = None
        target_index = -1
        for idx, existing in enumerate(rows):
            if str(existing.get("id")) == suggestion_id:
                target = existing
                target_index = idx
                break
        if target is None:
            raise HTTPException(status_code=404, detail="Suggestion not found")

        if req.title is not None:
            new_title = req.title.strip()
            if not new_title:
                raise HTTPException(status_code=400, detail="Title cannot be empty")
            target["title"] = new_title
        if req.summary is not None:
            target["summary"] = req.summary.strip() if req.summary else ""
        if req.classification is not None:
            target["classification"] = _require_classification(req.classification)
        if req.status is not None:
            target["status"] = _require_status(req.status)

        target["updated_at"] = _now_iso()
        rows[target_index] = target
        _write_suggestions(rows)
        saved = _decorate_suggestion(target)

    dto = SuggestionItemDTO(item=suggestion_to_dto(saved))
    return dto.dict_clean()


@app.post("/suggestions/{suggestion_id}/like")
def suggestions_like(suggestion_id: str, req: SuggestionLikeRequest) -> dict:
    delta = req.delta or 1
    with _suggestions_lock:
        rows = _load_suggestions()
        target = None
        target_index = -1
        for idx, existing in enumerate(rows):
            if str(existing.get("id")) == suggestion_id:
                target = existing
                target_index = idx
                break
        if target is None:
            raise HTTPException(status_code=404, detail="Suggestion not found")

        try:
            likes = int(target.get("likes") or 0)
        except Exception:
            likes = 0
        likes = max(0, likes + delta)
        target["likes"] = likes
        target["updated_at"] = _now_iso()
        rows[target_index] = target
        _write_suggestions(rows)
        saved = _decorate_suggestion(target)

    dto = SuggestionItemDTO(item=suggestion_to_dto(saved))
    return dto.dict_clean()


@app.post("/suggestions/{suggestion_id}/comments")
def suggestions_add_comment(suggestion_id: str, req: SuggestionCommentCreate) -> dict:
    text = (req.text or "").strip()
    if not text:
        raise HTTPException(status_code=400, detail="Comment text is required")
    with _suggestions_lock:
        rows = _load_suggestions()
        target = None
        target_index = -1
        for idx, existing in enumerate(rows):
            if str(existing.get("id")) == suggestion_id:
                target = existing
                target_index = idx
                break
        if target is None:
            raise HTTPException(status_code=404, detail="Suggestion not found")

        comments = target.get("comments") or []
        if not isinstance(comments, list):
            comments = []
        comment = {
            "id": uuid.uuid4().hex,
            "text": text,
            "created_at": _now_iso(),
        }
        comments.append(comment)
        target["comments"] = comments
        target["updated_at"] = _now_iso()
        rows[target_index] = target
        _write_suggestions(rows)
        saved = _decorate_suggestion(target)

    dto = SuggestionCommentResponseDTO(
        item=suggestion_to_dto(saved),
        comment=SuggestionCommentDTO.model_validate(comment),
    )
    return dto.dict_clean()


@app.delete("/suggestions/{suggestion_id}/comments/{comment_id}")
def suggestions_delete_comment(suggestion_id: str, comment_id: str) -> dict:
    with _suggestions_lock:
        rows = _load_suggestions()
        target = None
        target_index = -1
        for idx, existing in enumerate(rows):
            if str(existing.get("id")) == suggestion_id:
                target = existing
                target_index = idx
                break
        if target is None:
            raise HTTPException(status_code=404, detail="Suggestion not found")

        comments = target.get("comments") or []
        if not isinstance(comments, list):
            comments = []
        new_comments = [c for c in comments if str(c.get("id")) != comment_id]
        if len(new_comments) == len(comments):
            raise HTTPException(status_code=404, detail="Comment not found")
        target["comments"] = new_comments
        target["updated_at"] = _now_iso()
        rows[target_index] = target
        _write_suggestions(rows)
        saved = _decorate_suggestion(target)

    dto = SuggestionItemDTO(item=suggestion_to_dto(saved))
    return dto.dict_clean()


@app.delete("/suggestions/{suggestion_id}")
def suggestions_delete(suggestion_id: str) -> dict:
    with _suggestions_lock:
        rows = _load_suggestions()
        new_rows = [row for row in rows if str(row.get("id")) != suggestion_id]
        if len(new_rows) == len(rows):
            raise HTTPException(status_code=404, detail="Suggestion not found")
        _write_suggestions(new_rows)
    return {"ok": True}


def _build_admin_settings() -> list[dict[str, Any]]:
    env_values = _read_env_values()
    defaults = _read_env_defaults()
    settings: list[dict[str, Any]] = []
    for field in ENV_SETTING_FIELDS:
        key = field["key"]
        secret = bool(field.get("secret"))
        placeholder = field.get("placeholder") or ""
        category = field.get("category") or "general"
        current_value = os.getenv(key, "")
        has_value = bool(current_value)
        value = "" if secret else current_value
        placeholder_effective = placeholder
        if secret and has_value:
            placeholder_effective = "•••••• (hidden)"
        elif not placeholder_effective and not has_value:
            placeholder_effective = defaults.get(key, "")
        settings.append(
            {
                "key": key,
                "label": field.get("label", key),
                "secret": secret,
                "value": value,
                "has_value": has_value,
                "placeholder": placeholder,
                "placeholder_effective": placeholder_effective,
                "source": "file" if key in env_values else "env",
                "default": defaults.get(key, ""),
                "category": category,
            }
        )
    return settings


@app.get("/admin/env")
def admin_env_settings():
    settings = _build_admin_settings()
    backup_enabled = os.getenv("BACKUP_ENABLE", "1").strip().lower() not in {"0", "false", "no", "off"}
    backup_type = os.getenv("BACKUP_TYPE", "local").strip().lower()

    # Determine if backup is properly configured based on type
    backup_configured = False
    backup_target = ""

    if backup_type == "local":
        backup_path = os.getenv("BACKUP_LOCAL_PATH", "backups").strip()
        backup_configured = bool(backup_path)
        backup_target = backup_path
    elif backup_type in {"sftp", "scp"}:
        host = os.getenv("BACKUP_HOST", "").strip()
        username = os.getenv("BACKUP_USERNAME", "").strip()
        password = os.getenv("BACKUP_PASSWORD", "").strip()
        private_key = os.getenv("BACKUP_PRIVATE_KEY_PATH", "").strip()
        remote_path = os.getenv("BACKUP_REMOTE_PATH", "").strip()

        backup_configured = bool(host and username and (password or private_key))
        backup_target = (
            f"{username}@{host}:{remote_path}"
            if host and username and remote_path
            else f"{username}@{host}"
            if host and username
            else ""
        )

    settings_dto = [EnvSettingDTO.model_validate(setting) for setting in settings]
    backup_dto = BackupInfoDTO(
        enabled=backup_enabled,
        configured=backup_configured,
        type=backup_type,
        target=backup_target or None,
    )
    dto = AdminEnvResponseDTO(settings=settings_dto, backup=backup_dto)
    return dto.dict_clean()


@app.post("/admin/env")
def admin_env_update(req: EnvUpdateRequest):
    key = req.key.strip()
    if not key:
        raise HTTPException(status_code=400, detail="Key is required")
    field = next((f for f in ENV_SETTING_FIELDS if f["key"] == key), None)
    if field is None:
        raise HTTPException(status_code=404, detail="Unknown setting")
    if req.value is None:
        return admin_env_settings()
    _write_env_value(key, req.value)
    return admin_env_settings()


@app.post("/admin/env/reset")
def admin_env_reset(req: EnvResetRequest):
    key = req.key.strip()
    if not key:
        raise HTTPException(status_code=400, detail="Key is required")
    field = next((f for f in ENV_SETTING_FIELDS if f["key"] == key), None)
    if field is None:
        raise HTTPException(status_code=404, detail="Unknown setting")
    _write_env_value(key, None)
    return admin_env_settings()


@app.post("/admin/backup-sync")
def admin_backup_sync(user: OptionalUserDep = None):
    actor = getattr(user, "username", None)
    with task_logging("admin.backup_sync", actor=actor, trigger="ui") as task_log:
        try:
            result = backup_sync.sync_data_dir(note="manual-ui")
        except Exception as exc:  # pragma: no cover
            raise HTTPException(status_code=500, detail=str(exc)) from exc

        status = (result or {}).get("status")
        count = (result or {}).get("count")
        method = (result or {}).get("method")
        task_log.add_success(status=status, method=method, file_count=count)

        if status and status not in {"ok", "skipped"}:
            logger.warning(
                "Backup sync finished with non-OK status",
                extra={
                    "event": "task_warning",
                    "status": status,
                },
            )

        return result


@app.get("/config/ui")
def ui_config():
    theme = os.getenv("UI_THEME_DEFAULT", "nebula").strip() or "nebula"
    return {"theme_default": theme}


@app.get("/config/chat")
def chat_config():
    return {
        "system_prompt": os.getenv("CHAT_SYSTEM_PROMPT", ""),
        "temperature": _chat_default_temperature(),
    }


def _list_records(
    csv_name: str,
    limit: int | None,
    offset: int,
    order_by: str | None,
    order_dir: Literal["asc", "desc"],
) -> list[dict]:
    path = _csv_path(csv_name)
    if not path.exists():
        raise HTTPException(status_code=404, detail=f"{csv_name} not found")
    # Read headers to validate order_by
    with path.open(newline="", encoding="utf-8") as fh:
        reader = csv.reader(fh)
        headers = next(reader, [])
    if order_by and order_by not in headers:
        raise HTTPException(status_code=400, detail=f"Invalid order_by: {order_by}")

    ident = f'"{order_by}" {order_dir.upper()}' if order_by else None
    sql = "SELECT * FROM read_csv_auto(?, header=True)"
    params: list[Any] = [path.as_posix()]
    if ident:
        sql += f" ORDER BY {ident}"
    if limit is not None:
        sql += " LIMIT ? OFFSET ?"
        params.extend([int(limit), int(offset)])

    df = duckdb.query(sql, params=params).df()
    # Normalize to JSON‑safe values: NaN/NaT/±Inf -> None
    if not df.empty:
        df = df.replace([np.inf, -np.inf], np.nan)
        df = df.astype(object).where(pd.notnull(df), None)
    return df.to_dict(orient="records")


# ---------------------------
# Chat integration (simple)
# ---------------------------

# Chat sessions are now stored in the database
CHAT_QUERY_STOP_WORDS: set[str] = {
    "show",
    "list",
    "give",
    "get",
    "display",
    "please",
    "provide",
    "tell",
    "which",
    "what",
    "where",
    "find",
    "return",
    "me",
    "the",
    "latest",
    "recent",
    "top",
    "all",
    "any",
    "about",
    "for",
    "with",
    "those",
    "these",
    "new",
    "first",
    "last",
    "server",
    "servers",
    "device",
    "devices",
    "data",
    "information",
}


def _chat_env(*, db: Session | None = None, user: User | None = None) -> dict[str, Any]:
    close_db = False
    if db is None:
        db = SessionLocal()
        close_db = True

    try:
        env = {
            "openai": {
                "api_key": os.getenv("OPENAI_API_KEY", "").strip(),
                "default_model": os.getenv("CHAT_DEFAULT_MODEL_OPENAI", "gpt-5-mini"),
                "key_source": "env",
            },
            "openrouter": {
                "api_key": os.getenv("OPENROUTER_API_KEY", "").strip(),
                "default_model": os.getenv("CHAT_DEFAULT_MODEL_OPENROUTER", "openrouter/auto"),
                "key_source": "env",
            },
            "claude": {
                "api_key": os.getenv("ANTHROPIC_API_KEY", "").strip(),
                "default_model": os.getenv("CHAT_DEFAULT_MODEL_CLAUDE", "claude-3-5-sonnet-20240620"),
                "key_source": "env",
            },
            "gemini": {
                "api_key": os.getenv("GOOGLE_API_KEY", "").strip(),
                "default_model": os.getenv("CHAT_DEFAULT_MODEL_GEMINI", "gemini-1.5-flash"),
                "key_source": "env",
            },
            "default_provider": os.getenv("CHAT_DEFAULT_PROVIDER", "openai"),
        }

        for provider_id in ("openai", "openrouter", "claude", "gemini"):
            override_label: str | None = None
            override_secret: str | None = None
            if user:
                user_key = _get_user_api_key(db, user.id, provider_id)
                if user_key and user_key.secret:
                    override_secret = user_key.secret
                    override_label = user_key.label
                    env[provider_id]["key_source"] = "user"
            if not override_secret:
                global_key = _get_global_api_key(db, provider_id)
                if global_key and global_key.secret:
                    override_secret = global_key.secret
                    override_label = global_key.label or override_label
                    env[provider_id]["key_source"] = "global"
            if override_secret:
                env[provider_id]["api_key"] = override_secret
                if override_label:
                    env[provider_id]["label"] = override_label
            elif env[provider_id]["api_key"]:
                env[provider_id]["key_source"] = "env"
            else:
                env[provider_id]["key_source"] = None

        return env
    finally:
        if close_db:
            db.close()


def _format_responses_messages(messages: list[dict[str, str]]) -> list[dict[str, Any]]:
    formatted: list[dict[str, Any]] = []
    for msg in messages:
        role = msg.get("role", "user")
        text = str(msg.get("content", ""))
        # OpenAI Responses API requires input_text for prompts and output_text for responses
        content_type = "output_text" if role == "assistant" else "input_text"
        formatted.append(
            {
                "role": role,
                "content": [{"type": content_type, "text": text}],
            }
        )
    return formatted


def _get_chat_session(db: Session, session_id: str) -> ChatSession | None:
    """Get chat session by session_id."""
    from sqlalchemy import select

    stmt = select(ChatSession).where(ChatSession.session_id == session_id)
    return db.execute(stmt).scalar_one_or_none()


def _safe_to_str(value: Any) -> str | None:
    try:
        return str(value)
    except Exception as exc:  # pragma: no cover - defensive logging
        logger.debug("Skipping chat variable with non-stringable key", extra={"error": str(exc)})
        return None


def _safe_json_loads(data: str) -> Any | None:
    try:
        return json.loads(data)
    except Exception as exc:  # pragma: no cover - defensive logging
        logger.debug("Skipping streaming event with invalid JSON", extra={"error": str(exc)})
        return None


def _normalise_chat_variables(payload: Mapping[str, Any] | None) -> dict[str, Any]:
    if not payload:
        return {}
    normalised: dict[str, Any] = {}
    for raw_key, raw_value in payload.items():
        key = _safe_to_str(raw_key)
        if key is None:
            continue
        if raw_value is None:
            # Explicit None is treated as removal; skip during normalisation.
            normalised[key] = None
            continue
        if isinstance(raw_value, str | int | float | bool):
            normalised[key] = raw_value
            continue
        try:
            json.dumps(raw_value)
            normalised[key] = raw_value
        except TypeError:
            normalised[key] = str(raw_value)
    return normalised


def _apply_chat_variables(
    session: ChatSession,
    updates: Mapping[str, Any] | None,
    *,
    merge: bool = True,
) -> dict[str, Any]:
    if updates is None:
        return dict(session.context_variables or {})
    sanitised = _normalise_chat_variables(updates)
    base = dict(session.context_variables or {}) if merge else {}
    for key, value in sanitised.items():
        if value is None:
            base.pop(key, None)
        else:
            base[key] = value
    session.context_variables = base
    return dict(base)


def _create_chat_session(
    db: Session,
    session_id: str | None = None,
    title: str | None = None,
    user_id: str | None = None,
    variables: Mapping[str, Any] | None = None,
) -> ChatSession:
    """Create a new chat session."""
    if not session_id:
        session_id = "c_" + secrets.token_hex(8)

    context_vars = {key: value for key, value in _normalise_chat_variables(variables).items() if value is not None}
    session = ChatSession(
        session_id=session_id,
        title=title or "New chat",
        user_id=user_id,
        context_variables=context_vars,
    )
    db.add(session)
    db.commit()
    db.refresh(session)
    return session


def _update_session_title_from_message(db: Session, session: ChatSession, message: str) -> None:
    """Update session title based on first user message if title is still default."""
    if session.title in ("New chat", "", None) and message.strip():
        session.title = message.strip()[:60]
        db.add(session)
        db.commit()


def _add_chat_message(
    db: Session,
    session: ChatSession,
    role: str,
    content: str,
    *,
    usage: dict[str, int] | None = None,
) -> ChatMessage:
    """Add a message to a chat session."""
    if usage:
        try:
            import json as _json

            content_with_usage = content.rstrip() + "\n[[TOKENS " + _json.dumps(usage) + "]]"
        except Exception:
            content_with_usage = content
    else:
        content_with_usage = content
    message = ChatMessage(
        session_id=session.id,
        role=role,
        content=content_with_usage,
    )
    db.add(message)

    # Update session timestamp
    session.updated_at = datetime.now(UTC)
    db.add(session)

    db.commit()
    db.refresh(message)
    return message


def _serialize_chat_session(session: ChatSession) -> dict[str, Any]:
    """Serialize chat session for API response."""
    return {
        "session_id": session.session_id,
        "title": session.title,
        "created_at": session.created_at.isoformat() + "Z",
        "updated_at": session.updated_at.isoformat() + "Z",
        "variables": dict(session.context_variables or {}),
    }


def _serialize_chat_message(message: ChatMessage) -> dict[str, Any]:
    """Serialize chat message for API response."""
    content = message.content or ""
    usage = None
    if "[[TOKENS" in content:
        try:
            marker_start = content.rfind("[[TOKENS ")
            marker_end = content.rfind("]]", marker_start)
            if marker_start >= 0 and marker_end > marker_start:
                import json as _json

                payload = content[marker_start + len("[[TOKENS ") : marker_end]
                usage = _json.loads(payload)
                content = content[:marker_start].rstrip()
        except Exception:
            usage = None

    data = {
        "role": message.role,
        "content": content,
        "created_at": message.created_at.isoformat() + "Z",
    }
    if usage:
        data["usage"] = usage
    return data


class ChatRequest(BaseModel):
    provider: Literal["openai", "openrouter", "claude", "gemini"]
    model: str | None = None
    message: str
    session_id: str
    temperature: float | None = None
    system: str | None = None
    include_context: bool | None = False
    dataset: Literal["devices", "vms", "all", "merged"] | None = "merged"
    variables: dict[str, Any] | None = None
    tool: str | None = None


class ChatSessionCreate(BaseModel):
    name: str | None = None
    variables: dict[str, Any] | None = None


ChatRequestBody = Annotated[ChatRequest, Body(...)]
ChatSessionCreateBody = Annotated[ChatSessionCreate | None, Body()]
ToolSamplePayloadBody = Annotated[dict[str, Any] | None, Body()]


def _messages_for_provider(messages: list[dict[str, str]], max_turns: int = 16) -> list[dict[str, str]]:
    # Keep last N turns to keep payloads light
    if len(messages) <= max_turns:
        return messages
    # keep the last max_turns messages
    return messages[-max_turns:]


def _use_openai_responses(model: str) -> bool:
    m = (model or "").lower()
    return m.startswith("gpt-5")


def _responses_supports_temperature(model: str) -> bool:
    m = (model or "").lower()
    # Current gpt-5 family rejects 'temperature'; omit when using Responses API
    return not m.startswith("gpt-5")


def _is_openai_streaming_unsupported(ex: Exception) -> bool:
    """Heuristics to detect OpenAI responses error that disallows streaming for the model/org."""
    msg = str(ex).lower()
    if "must be verified to stream" in msg or ("stream" in msg and "unsupported" in msg):
        return True
    # Try to inspect HTTPError JSON payload
    try:
        resp = getattr(ex, "response", None)
        if resp is not None:
            try:
                data = resp.json()
            except Exception:
                data = None
            if isinstance(data, dict):
                err = data.get("error") or {}
                if isinstance(err, dict):
                    if err.get("param") == "stream" and err.get("code") == "unsupported_value":
                        return True
                    if (
                        isinstance(err.get("message"), str)
                        and "must be verified to stream" in err.get("message", "").lower()
                    ):
                        return True
    except Exception:
        pass
    return False


def _iter_chunks(text: str, size: int = 128):
    for i in range(0, len(text), size):
        yield text[i : i + size]


def _call_openai(
    model: str,
    messages: list[dict[str, str]],
    api_key: str,
    temperature: float = 0.2,
) -> ChatProviderResult:
    # Prefer the official SDK when available for reliability
    if OpenAI is not None:
        client = OpenAI(api_key=api_key)
        if _use_openai_responses(model):
            _kwargs: dict[str, Any] = {"model": model, "input": _format_responses_messages(messages)}
            if _responses_supports_temperature(model) and temperature is not None:
                _kwargs["temperature"] = temperature
            resp = client.responses.create(**_kwargs)
            usage = _normalise_usage(getattr(resp, "usage", None))
            try:
                text = getattr(resp, "output_text", None)
                if text:
                    return ChatProviderResult(str(text).strip(), usage)
                # Fallback: collect text parts
                chunks = []
                for item in getattr(resp, "output", []) or []:
                    for part in getattr(item, "content", []) or []:
                        if getattr(part, "type", "") == "output_text":
                            chunks.append(getattr(part, "text", ""))
                return ChatProviderResult("".join(chunks).strip(), usage)
            except Exception:
                return ChatProviderResult("", usage)
        else:
            resp = client.chat.completions.create(model=model, messages=messages, temperature=temperature)
            usage = _normalise_usage(getattr(resp, "usage", None))
            try:
                choice = (resp.choices or [None])[0]
                msg = getattr(choice, "message", None)
                return ChatProviderResult((getattr(msg, "content", "") or "").strip(), usage)
            except Exception:
                return ChatProviderResult("", usage)
    # SDK not available — fall back to HTTP
    if _use_openai_responses(model):
        url = "https://api.openai.com/v1/responses"
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
        payload = {"model": model, "input": _format_responses_messages(messages)}
        if _responses_supports_temperature(model) and temperature is not None:
            payload["temperature"] = temperature
        r = requests.post(url, headers=headers, json=payload, timeout=90)
        r.raise_for_status()
        data = r.json()
        usage = _normalise_usage(data.get("usage"))
        text = (data.get("output_text") or "").strip()
        if text:
            return ChatProviderResult(text, usage)
        try:
            outs = data.get("output", [])
            chunks = []
            for item in outs:
                parts = item.get("content", []) if isinstance(item, dict) else []
                for p in parts:
                    if isinstance(p, dict) and p.get("type") == "output_text":
                        chunks.append(p.get("text") or "")
            if chunks:
                return ChatProviderResult("".join(chunks).strip(), usage)
        except Exception:
            pass
        return ChatProviderResult("", usage)
    else:
        url = "https://api.openai.com/v1/chat/completions"
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
        payload = {
            "model": model,
            "messages": messages,
            "temperature": temperature,
        }
        r = requests.post(url, headers=headers, json=payload, timeout=60)
        r.raise_for_status()
        data = r.json()
        usage = _normalise_usage(data.get("usage"))
        text = (data.get("choices", [{}])[0].get("message", {}).get("content") or "").strip()
        return ChatProviderResult(text, usage)


def _call_openrouter(
    model: str,
    messages: list[dict[str, str]],
    api_key: str,
    temperature: float = 0.2,
) -> ChatProviderResult:
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        # Optional but recommended by OpenRouter
        "HTTP-Referer": os.getenv("OPENROUTER_REFERRER", ""),
        "X-Title": os.getenv("OPENROUTER_TITLE", "Enreach Tools"),
    }
    payload = {
        "model": model,
        "messages": messages,
        "temperature": temperature,
    }
    r = requests.post(url, headers=headers, json=payload, timeout=60)
    r.raise_for_status()
    data = r.json()
    usage = _normalise_usage(data.get("usage"))
    text = (data.get("choices", [{}])[0].get("message", {}).get("content") or "").strip()
    return ChatProviderResult(text, usage)


def _call_claude(
    model: str,
    messages: list[dict[str, str]],
    api_key: str,
    temperature: float = 0.2,
) -> ChatProviderResult:
    url = "https://api.anthropic.com/v1/messages"
    headers = {
        "x-api-key": api_key,
        "anthropic-version": "2023-06-01",
        "content-type": "application/json",
    }
    # Anthropic expects messages array without an initial system item; support optional separate system
    sys_prompt = None
    if messages and messages[0]["role"] == "system":
        sys_prompt = messages[0].get("content")
        messages = messages[1:]
    payload = {
        "model": model,
        "max_tokens": 800,
        "messages": messages,
        "temperature": temperature,
    }
    if sys_prompt:
        payload["system"] = sys_prompt
    r = requests.post(url, headers=headers, json=payload, timeout=60)
    r.raise_for_status()
    data = r.json()
    usage = _normalise_usage(data.get("usage"))
    # content is a list of blocks; take first text
    blocks = data.get("content", [])
    if isinstance(blocks, list) and blocks:
        part = blocks[0]
        if isinstance(part, dict) and part.get("type") == "text":
            return ChatProviderResult((part.get("text") or "").strip(), usage)
    # Fallback: try candidates
    return ChatProviderResult((data.get("output_text") or "").strip(), usage)


def _call_gemini(
    model: str,
    messages: list[dict[str, str]],
    api_key: str,
    temperature: float = 0.2,
) -> ChatProviderResult:
    # Convert to Gemini content format
    def to_parts(msgs: list[dict[str, str]]):
        parts = []
        for m in msgs:
            role = m.get("role", "user")
            text = m.get("content", "")
            if role == "system":
                # prepend system to first user message
                parts.append({"role": "user", "parts": [{"text": f"[SYSTEM]\n{text}"}]})
            elif role == "assistant":
                parts.append({"role": "model", "parts": [{"text": text}]})
            else:
                parts.append({"role": "user", "parts": [{"text": text}]})
        return parts

    base = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
    payload = {
        "contents": to_parts(messages),
        "generationConfig": {
            "temperature": temperature,
            "maxOutputTokens": 800,
        },
    }
    r = requests.post(base, json=payload, timeout=60)
    r.raise_for_status()
    data = r.json()
    usage_meta = data.get("usageMetadata")
    try:
        text = (data["candidates"][0]["content"]["parts"][0]["text"] or "").strip()
        return ChatProviderResult(text, _normalise_usage(usage_meta))
    except Exception:
        return ChatProviderResult((data.get("text") or "").strip(), _normalise_usage(usage_meta))


def _csv_for_dataset(dataset: str) -> Path | None:
    target = (dataset or "").strip().lower()
    if target == "devices":
        p = _csv_path("netbox_devices_export.csv")
    elif target == "vms":
        p = _csv_path("netbox_vms_export.csv")
    else:
        # "all" (legacy) and "merged" both map to the merged export
        p = _csv_path("netbox_merged_export.csv")
    return p if p.exists() else None


def _build_data_context(dataset: str, query: str, max_rows: int = 6, max_chars: int = 1800) -> str:
    """Return a compact textual context from CSV based on a keyword query.
    Includes columns and up to N matching rows across all columns (case-insensitive LIKE).
    """
    p = _csv_for_dataset(dataset)
    if not p:
        return ""
    try:
        # Read headers
        import csv as _csv

        with p.open("r", encoding="utf-8", errors="ignore") as fh:
            rdr = _csv.reader(fh)
            headers = next(rdr, [])
        if not headers:
            return ""

        tokens = re.findall(r"[A-Za-z0-9]+", query.lower()) if query else []
        text_keywords: list[str] = []
        numeric_keywords: list[str] = []
        numeric_tokens: list[int] = []
        seen_text: set[str] = set()
        seen_numeric: set[str] = set()
        for token in tokens:
            if not token:
                continue
            if token.isdigit():
                try:
                    numeric_tokens.append(int(token))
                except ValueError:
                    continue
                if len(token) >= 2 and token not in seen_numeric:
                    numeric_keywords.append(token)
                    seen_numeric.add(token)
                continue
            if token in CHAT_QUERY_STOP_WORDS:
                continue
            if len(token) < 3:
                continue
            if token not in seen_text:
                text_keywords.append(token)
                seen_text.add(token)

        keywords = text_keywords if text_keywords else numeric_keywords

        limit = max_rows
        if numeric_tokens:
            limit = max(3, min(max(numeric_tokens), 20))

        where_clauses: list[str] = []
        if keywords:
            for kw in keywords[:5]:
                safe_kw = kw.replace("'", "''")
                ors = [f"lower(CAST(\"{h}\" AS VARCHAR)) LIKE '%{safe_kw}%'" for h in headers]
                where_clauses.append("(" + " OR ".join(ors) + ")")

        def _run(where_clause: str | None = None) -> pd.DataFrame:
            sql = "SELECT * FROM read_csv_auto(?, header=True)"
            params: list[Any] = [p.as_posix()]
            if where_clause:
                sql += f" WHERE {where_clause}"
            sql += " LIMIT ?"
            params.append(int(limit))
            return duckdb.query(sql, params=params).df()

        if where_clauses:
            df = _run(" OR ".join(where_clauses))
        elif query.strip():
            safe = query.replace("'", "''")
            ors = [f"lower(CAST(\"{h}\" AS VARCHAR)) LIKE lower('%{safe}%')" for h in headers]
            df = _run(" OR ".join(ors))
            if df.empty:
                df = _run()
        else:
            df = _run()
        # Render context text
        parts: list[str] = []
        parts.append(f"Source file: {p.name}")
        parts.append(f"Columns: {', '.join(map(str, headers))}")
        if not df.empty:
            parts.append("Relevant rows:")
            preferred = [
                "Name",
                "Status",
                "Tenant",
                "Site",
                "Location",
                "Rack",
                "Rack Position",
                "Role",
                "Manufacturer",
                "Type",
                "Platform",
                "IP Address",
                "IPv4 Address",
                "IPv6 Address",
                "ID",
                "Serial number",
                "Asset tag",
                "Region",
                "Server Group",
                "Cluster",
                "DTAP state",
                "CPU",
                "VCPUs",
                "Memory",
                "Memory (MB)",
                "Disk",
                "Harddisk",
                "Backup",
            ]
            preferred_lower = [p.lower() for p in preferred]

            def normalise_value(value: Any) -> Any:
                if value is None:
                    return None
                if isinstance(value, pd.Timestamp):
                    return value.isoformat()
                if hasattr(value, "isoformat") and not isinstance(value, str | int | float | bool):
                    try:
                        return value.isoformat()
                    except Exception:
                        return str(value)
                if isinstance(value, float) and np.isnan(value):
                    return None
                if isinstance(value, str | int | float | bool):
                    return value
                return str(value)

            for idx, (_, row) in enumerate(df.iterrows(), start=1):
                try:
                    obj = {str(k): normalise_value(v) for k, v in row.items()}
                except Exception:
                    obj = {str(k): normalise_value(str(v)) for k, v in row.items()}

                # Filter out empty/null values
                non_empty = {k: v for k, v in obj.items() if v not in (None, "", "null", "None")}
                if not non_empty:
                    continue

                title = non_empty.get("Name") or non_empty.get("Device") or non_empty.get("ID") or f"Row {idx}"
                parts.append(f"- {title}")

                seen: set[str] = set()
                for field in preferred:
                    if field in non_empty:
                        value = non_empty[field]
                        parts.append(f"    - **{field}:** {value}")
                        seen.add(field)
                for field, value in non_empty.items():
                    if field in seen:
                        continue
                    if field.lower() in preferred_lower:
                        continue
                    parts.append(f"    - **{field}:** {value}")
        context = "\n".join(parts)
        if len(context) > max_chars:
            context = context[: max_chars - 20] + "\n…"
        return context
    except Exception:
        return ""


@app.get("/chat/providers")
def chat_providers(request: Request, user: CurrentUserDep, db: DbSessionDep):
    require_permission(request, "chat.use")
    env = _chat_env(db=db, user=user)
    out = []
    for pid in ["openai", "openrouter", "claude", "gemini"]:
        cfg = env.get(pid, {})
        out.append(
            {
                "id": pid,
                "configured": bool(cfg.get("api_key")),
                "default_model": cfg.get("default_model"),
                "key_source": cfg.get("key_source"),
                "label": cfg.get("label"),
            }
        )
    return {"providers": out, "default_provider": env.get("default_provider", "openai")}


@app.get("/chat/history")
def chat_history(
    request: Request,
    db: DbSessionDep,
    session_id: str = Query(...),
):
    require_permission(request, "chat.use")
    session = _get_chat_session(db, session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Chat session not found")

    messages = [_serialize_chat_message(msg) for msg in session.messages]
    return {
        "session_id": session_id,
        "messages": messages,
        "variables": dict(session.context_variables or {}),
    }


@app.get("/chat/sessions")
def chat_sessions(
    request: Request,
    db: DbSessionDep,
    user: OptionalUserDep,
    limit: int | None = Query(None, ge=1, le=200),
):
    require_permission(request, "chat.use")
    from sqlalchemy import select

    stmt = select(ChatSession).order_by(ChatSession.updated_at.desc())
    if user:
        stmt = stmt.where(ChatSession.user_id == user.id)
    if limit:
        stmt = stmt.limit(limit)

    sessions = db.execute(stmt).scalars().all()
    return {"sessions": [_serialize_chat_session(session) for session in sessions]}


@app.post("/chat/session")
def chat_session_create(
    request: Request,
    db: DbSessionDep,
    user: OptionalUserDep,
    req: ChatSessionCreateBody = None,
):
    require_permission(request, "chat.use")
    title = (req.name if req else "") or "New chat"
    variables = req.variables if req else None
    session = _create_chat_session(
        db,
        title=title,
        user_id=user.id if user else None,
        variables=variables or None,
    )
    return _serialize_chat_session(session)


@app.delete("/chat/session/{session_id}")
def chat_session_delete(
    request: Request,
    session_id: str,
    db: DbSessionDep,
    user: OptionalUserDep,
):
    """Delete a chat session and all its messages."""
    require_permission(request, "chat.use")
    session = _get_chat_session(db, session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Chat session not found")

    # Optional: Only allow users to delete their own sessions
    if user and session.user_id and session.user_id != user.id:
        raise HTTPException(status_code=403, detail="You can only delete your own chat sessions")

    # Delete the session (messages will be deleted automatically due to cascade)
    db.delete(session)
    db.commit()

    return {"status": "deleted", "session_id": session_id}


@app.post("/chat/complete")
def chat_complete(
    request: Request,
    req: ChatRequestBody,
    user: OptionalUserDep,
    db: DbSessionDep,
):
    require_permission(request, "chat.use")
    env = _chat_env(db=db, user=user)
    pid = req.provider
    if pid not in env:
        raise HTTPException(status_code=400, detail=f"Unknown provider: {pid}")
    api_key = env[pid].get("api_key")
    if not api_key:
        raise HTTPException(status_code=400, detail=f"Provider '{pid}' not configured (missing API key)")
    model = (req.model or env[pid].get("default_model") or "").strip()
    if not model:
        raise HTTPException(status_code=400, detail=f"No model specified for provider '{pid}'")

    # Get or create chat session
    session = _get_chat_session(db, req.session_id)
    if not session:
        session = _create_chat_session(db, req.session_id, user_id=user.id if user else None)

    history = [_serialize_chat_message(msg) for msg in session.messages]

    system_prompt = (req.system or os.getenv("CHAT_SYSTEM_PROMPT", "")).strip()
    if system_prompt:
        if history and history[0].get("role") == "system":
            history[0]["content"] = system_prompt[:4000]
        else:
            history.insert(0, {"role": "system", "content": system_prompt[:4000]})

    if req.include_context:
        ctx_text = _build_data_context(req.dataset or "all", req.message)
        if ctx_text:
            history.append(
                {
                    "role": "system",
                    "content": f"Data context from {req.dataset or 'all'}:\n" + ctx_text,
                }
            )

    merged_variables = _apply_chat_variables(session, req.variables or None)
    db.add(session)

    user_message = str(req.message)[:8000]
    _add_chat_message(db, session, "user", user_message)
    _update_session_title_from_message(db, session, user_message)

    history.append({"role": "user", "content": user_message})
    agent_history = history[:-1]
    agent_input = history[-1]["content"] if history else user_message

    temperature = req.temperature if req.temperature is not None else _chat_default_temperature()

    actor = getattr(user, "username", None)
    try:
        runtime = AgentRuntime(provider=pid, model=model, api_key=api_key, temperature=temperature)
    except AgentRuntimeError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    with task_logging(
        "chat.complete",
        actor=actor,
        provider=pid,
        model=model,
        session_id=req.session_id,
    ) as task_log:
        try:
            run_result = runtime.run(
                session_id=session.session_id,
                user_label=actor,
                variables=merged_variables,
                history=agent_history,
                message=agent_input,
                tool_hint=req.tool,
            )
        except AgentRuntimeError as exc:
            err = f"Agent error: {exc}"
            _add_chat_message(db, session, "assistant", err)
            task_log.add_success(status="agent_error")
            return {"session_id": req.session_id, "provider": pid, "model": model, "reply": err}
        except Exception as exc:
            err = f"Error: {exc}"
            _add_chat_message(db, session, "assistant", err)
            task_log.add_success(status="error")
            return {"session_id": req.session_id, "provider": pid, "model": model, "reply": err}

        assistant_text = (run_result.text or "").strip() or "(no response)"
        usage = run_result.usage or None
        _add_chat_message(db, session, "assistant", assistant_text, usage=usage)

        task_log.add_success(
            reply_chars=len(assistant_text),
            agent_domain=run_result.decision.domain,
            agent_tool=run_result.decision.tool_key,
            agent_steps=len(run_result.intermediate_steps),
        )
        if usage:
            task_log.add_success(
                prompt_tokens=usage.get("prompt_tokens"),
                completion_tokens=usage.get("completion_tokens"),
                total_tokens=usage.get("total_tokens"),
            )

        logger.info(
            "Agent completed turn",
            extra={
                "event": "chat_agent_completed",
                "agent_domain": run_result.decision.domain,
                "agent_tool": run_result.decision.tool_key,
                "agent_reason": run_result.decision.reason,
            },
        )

        response: dict[str, Any] = {
            "session_id": req.session_id,
            "provider": pid,
            "model": model,
            "reply": assistant_text,
            "agent": {
                "domain": run_result.decision.domain,
                "tool": run_result.decision.tool_key,
                "reason": run_result.decision.reason,
            },
        }
        if run_result.tool_outputs:
            response["tool_outputs"] = run_result.tool_outputs
        if usage:
            response["usage"] = usage
        return response


def _stream_openai_text(
    model: str,
    messages: list[dict[str, str]],
    api_key: str,
    temperature: float = 0.2,
    usage_target: list[dict[str, int]] | None = None,
):
    # Prefer SDK streaming
    if OpenAI is not None:
        client = OpenAI(api_key=api_key)
        if _use_openai_responses(model):
            try:
                kwargs: dict[str, Any] = {
                    "model": model,
                    "input": _format_responses_messages(messages),
                }
                if _responses_supports_temperature(model) and temperature is not None:
                    kwargs["temperature"] = temperature
                with client.responses.stream(**kwargs) as stream:
                    for event in stream:
                        event_type = getattr(event, "type", "")
                        if event_type == "response.output_text.delta":
                            delta = getattr(event, "delta", "")
                            if delta:
                                yield delta
                        if usage_target is not None:
                            usage = _normalise_usage(getattr(event, "usage", None))
                            if usage:
                                usage_target.clear()
                                usage_target.append(usage)
                    final_response = stream.get_final_response()
                    if usage_target is not None:
                        usage = _normalise_usage(getattr(final_response, "usage", None))
                        if usage:
                            usage_target.clear()
                            usage_target.append(usage)
                    return
            except Exception:
                # bubble up to caller; do not fall back to raw HTTP when SDK is present
                raise
        else:
            try:
                gen = client.chat.completions.create(
                    model=model,
                    messages=messages,
                    temperature=temperature,
                    stream=True,
                )
                for chunk in gen:
                    choice = (chunk.choices or [None])[0]
                    delta = getattr(choice, "delta", None)
                    if delta is not None:
                        text = getattr(delta, "content", None)
                        if text:
                            yield text
                    if usage_target is not None:
                        usage = _normalise_usage(getattr(chunk, "usage", None))
                        if usage:
                            usage_target.clear()
                            usage_target.append(usage)
                return
            except Exception:
                raise
    # HTTP fallback (SSE parsing)
    if _use_openai_responses(model):
        url = "https://api.openai.com/v1/responses"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
            "Accept": "text/event-stream",
        }
        payload = {
            "model": model,
            "input": _format_responses_messages(messages),
            "stream": True,
        }
        if _responses_supports_temperature(model) and temperature is not None:
            payload["temperature"] = temperature
        with requests.post(url, headers=headers, json=payload, stream=True, timeout=300) as r:
            r.raise_for_status()
            for raw_line in r.iter_lines(decode_unicode=True):
                if not raw_line:
                    continue
                line = raw_line
                if isinstance(line, bytes):
                    try:
                        line = line.decode("utf-8", errors="ignore")
                    except Exception:
                        line = str(line)
                if not isinstance(line, str):
                    line = str(line)
                if not line.startswith("data:"):
                    continue
                data = line[5:].strip()
                if data == "[DONE]":
                    break
                obj = _safe_json_loads(data)
                if not isinstance(obj, dict):
                    continue
                if obj.get("type") == "response.output_text.delta":
                    delta = obj.get("delta") or ""
                    if delta:
                        yield delta
                if usage_target is not None:
                    usage = _normalise_usage(obj.get("usage"))
                    if usage:
                        usage_target.clear()
                        usage_target.append(usage)
    else:
        url = "https://api.openai.com/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
            "Accept": "text/event-stream",
        }
        payload = {
            "model": model,
            "messages": messages,
            "temperature": temperature,
            "stream": True,
        }
        with requests.post(url, headers=headers, json=payload, stream=True, timeout=300) as r:
            r.raise_for_status()
            for raw_line in r.iter_lines(decode_unicode=True):
                if not raw_line:
                    continue
                line = raw_line
                if isinstance(line, bytes):
                    try:
                        line = line.decode("utf-8", errors="ignore")
                    except Exception:
                        line = str(line)
                if not isinstance(line, str):
                    line = str(line)
                if not line.startswith("data:"):
                    continue
                data = line[5:].strip()
                if data == "[DONE]":
                    break
                obj = _safe_json_loads(data)
                if not isinstance(obj, dict):
                    continue
                delta = obj.get("choices", [{}])[0].get("delta", {}).get("content")
                if delta:
                    yield delta
                if usage_target is not None:
                    usage = _normalise_usage(obj.get("usage"))
                    if usage:
                        usage_target.clear()
                        usage_target.append(usage)


def _stream_openrouter_text(
    model: str,
    messages: list[dict[str, str]],
    api_key: str,
    temperature: float = 0.2,
    usage_target: list[dict[str, int]] | None = None,
):
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "Accept": "text/event-stream",
        "HTTP-Referer": os.getenv("OPENROUTER_REFERRER", ""),
        "X-Title": os.getenv("OPENROUTER_TITLE", "Enreach Tools"),
    }
    payload = {
        "model": model,
        "messages": messages,
        "temperature": temperature,
        "stream": True,
    }
    with requests.post(url, headers=headers, json=payload, stream=True, timeout=300) as r:
        r.raise_for_status()
        for raw_line in r.iter_lines(decode_unicode=True):
            if not raw_line:
                continue
            line = raw_line
            if isinstance(line, bytes):
                try:
                    line = line.decode("utf-8", errors="ignore")
                except Exception:
                    line = str(line)
            if not isinstance(line, str):
                line = str(line)
            if not line.startswith("data:"):
                continue
            data = line[5:].strip()
            if data == "[DONE]":
                break
            obj = _safe_json_loads(data)
            if not isinstance(obj, dict):
                continue
            delta = obj.get("choices", [{}])[0].get("delta", {}).get("content")
            if delta:
                yield delta
            if usage_target is not None:
                usage = _normalise_usage(obj.get("usage"))
                if usage:
                    usage_target.clear()
                    usage_target.append(usage)


@app.post("/chat/stream")
def chat_stream(
    request: Request,
    req: ChatRequestBody,
    user: OptionalUserDep,
    db: DbSessionDep,
):
    require_permission(request, "chat.use")
    try:
        result = chat_complete(req, user, db, request)
    except HTTPException:
        raise
    except Exception as exc:  # pragma: no cover - defensive fallback
        raise HTTPException(status_code=500, detail=str(exc)) from exc

    reply_text = str(result.get("reply") or "")

    async def generator():
        yield reply_text

    return StreamingResponse(generator(), media_type="text/plain; charset=utf-8")


def _build_sample_prompt(definition: ToolDefinition, args: Mapping[str, Any]) -> str:
    rendered_args = json.dumps(dict(args), indent=2, ensure_ascii=False, default=str) if args else "{}"
    return (
        f"Use the {definition.agent} agent's `{definition.key}` tool with these arguments:\n"
        f"{rendered_args}\n"
        "Provide a concise operational summary highlighting key metrics."
    )


@app.post("/tools/{tool_key}/sample")
def tools_run_sample(
    tool_key: str,
    user: OptionalUserDep,
    db: DbSessionDep,
    request: Request,
    payload: ToolSamplePayloadBody = None,
    provider: str | None = Query(None, description="Override the provider used for the sample run."),
    model: str | None = Query(None, description="Override the model used for the sample run."),
):
    require_permission(request, "tools.use")
    catalog = tools_router._catalog_index()
    definition = catalog.get(tool_key)
    if definition is None:
        raise HTTPException(status_code=404, detail="Tool not found")

    if payload is not None and not isinstance(payload, Mapping):
        raise HTTPException(status_code=400, detail="Sample payload must be a JSON object")
    args: dict[str, Any] = {}
    if payload:
        args = dict(payload)
    elif definition.sample:
        args = dict(definition.sample)

    env = _chat_env(db=db, user=user)
    provider_id = (provider or env.get("default_provider") or "openai").strip()
    if provider_id not in env:
        raise HTTPException(status_code=400, detail=f"Unknown provider: {provider_id}")
    provider_cfg = env[provider_id]
    api_key = provider_cfg.get("api_key")
    if not api_key:
        raise HTTPException(status_code=400, detail=f"Provider '{provider_id}' not configured (missing API key)")
    model_name = (model or provider_cfg.get("default_model") or "").strip()
    if not model_name:
        raise HTTPException(status_code=400, detail=f"No model configured for provider '{provider_id}'")

    try:
        runtime = AgentRuntime(
            provider=provider_id,
            model=model_name,
            api_key=api_key,
            temperature=_chat_default_temperature(),
        )
    except AgentRuntimeError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    message = _build_sample_prompt(definition, args)
    actor = getattr(user, "username", None)

    with task_logging(
        "tools.sample",
        actor=actor,
        provider=provider_id,
        model=model_name,
        tool_key=tool_key,
    ) as task_log:
        try:
            run_result = runtime.run(
                session_id=None,
                user_label=actor,
                variables={},
                history=[],
                message=message,
                tool_hint=tool_key,
            )
        except AgentRuntimeError as exc:
            task_log.add_success(status="agent_error")
            raise HTTPException(status_code=502, detail=str(exc)) from exc
        except Exception as exc:
            task_log.add_success(status="error")
            raise HTTPException(status_code=500, detail=str(exc)) from exc

        summary_text = (run_result.text or "").strip()
        task_log.add_success(
            agent_domain=run_result.decision.domain,
            agent_tool=run_result.decision.tool_key,
            summary_chars=len(summary_text),
        )

    response: dict[str, Any] = {
        "summary": summary_text,
        "agent": {
            "domain": run_result.decision.domain,
            "tool": run_result.decision.tool_key,
            "reason": run_result.decision.reason,
        },
        "inputs": args,
    }
    if run_result.tool_outputs:
        response["tool_outputs"] = run_result.tool_outputs
        primary = run_result.tool_outputs[0].get("output")
        if primary is not None:
            response["data"] = primary
    return response


# ---------------------------
# Zabbix integration (read-only)
# ---------------------------


def _zbx_base_url() -> str | None:
    raw = os.getenv("ZABBIX_API_URL", "").strip()
    if raw:
        return raw
    host = os.getenv("ZABBIX_HOST", "").strip()
    if host:
        if host.endswith("/api_jsonrpc.php"):
            return host
        return host.rstrip("/") + "/api_jsonrpc.php"
    return None


def _zbx_web_base() -> str | None:
    web = os.getenv("ZABBIX_WEB_URL", "").strip()
    if web:
        return web.rstrip("/")
    api = _zbx_base_url()
    if api and api.endswith("/api_jsonrpc.php"):
        return api[: -len("/api_jsonrpc.php")]
    return None


def _zabbix_client() -> ZabbixClient:
    api_url = _zbx_base_url()
    if not api_url:
        raise HTTPException(status_code=400, detail="ZABBIX_API_URL or ZABBIX_HOST not configured")
    token = os.getenv("ZABBIX_API_TOKEN", "").strip() or None
    config = ZabbixClientConfig(
        api_url=api_url,
        api_token=token,
        web_url=_zbx_web_base(),
        timeout=30.0,
    )
    try:
        return ZabbixClient(config)
    except ZabbixConfigError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc


def _zbx_rpc(method: str, params: dict, *, client: ZabbixClient | None = None) -> dict:
    client = client or _zabbix_client()
    try:
        result = client.rpc(method, params)
    except ZabbixAuthError as err:
        raise HTTPException(status_code=401, detail=f"Zabbix error: {err}") from err
    except ZabbixError as err:
        raise HTTPException(status_code=502, detail=f"Zabbix error: {err}") from err
    if isinstance(result, dict | list):
        return result  # type: ignore[return-value]
    return {}


def _zbx_expand_groupids(base_group_ids: list[int]) -> list[int]:
    if not base_group_ids:
        return base_group_ids
    client = _zabbix_client()
    try:
        expanded = client.expand_groupids(base_group_ids)
    except ZabbixError:
        return base_group_ids
    return list(expanded)


@app.get("/zabbix/problems")
def zabbix_problems(
    severities: str | None = Query(None, description="Comma-separated severities 0..5 (e.g. '2,3,4')"),
    groupids: str | None = Query(None, description="Comma-separated group IDs"),
    hostids: str | None = Query(None, description="Comma-separated host IDs"),
    unacknowledged: int = Query(0, ge=0, le=1),
    suppressed: int = Query(0, ge=0, le=1),
    limit: int = Query(300, ge=1, le=2000),
    include_subgroups: int = Query(0, ge=0, le=1, description="When filtering by groupids, include all subgroup IDs"),
):
    """Return problems from Zabbix using problem.get with basic filters."""
    try:
        sev_list = [int(s) for s in (severities.split(",") if severities else []) if str(s).strip() != ""]
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid severities")
    try:
        grp_list = [int(s) for s in (groupids.split(",") if groupids else []) if str(s).strip() != ""]
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid groupids")
    try:
        host_list = [int(s) for s in (hostids.split(",") if hostids else []) if str(s).strip() != ""]
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid hostids")

    if not sev_list:
        env_sev = os.getenv("ZABBIX_SEVERITIES", "2,3,4").strip()
        if env_sev:
            try:
                sev_list = [int(x) for x in env_sev.split(",") if x.strip()]
            except Exception:
                sev_list = [2, 3, 4]
        else:
            sev_list = [2, 3, 4]
    if not grp_list:
        gid = os.getenv("ZABBIX_GROUP_ID", "").strip()
        if gid.isdigit():
            grp_list = [int(gid)]

    client = _zabbix_client()
    if grp_list and include_subgroups == 1:
        grp_list = list(client.expand_groupids(grp_list))

    try:
        problem_list = client.get_problems(
            severities=sev_list,
            groupids=grp_list,
            hostids=host_list,
            unacknowledged=bool(unacknowledged),
            suppressed=bool(suppressed) if suppressed in (0, 1) else None,
            limit=limit,
        )
    except ZabbixAuthError as err:
        raise HTTPException(status_code=401, detail=f"Zabbix error: {err}") from err
    except ZabbixError as err:
        raise HTTPException(status_code=502, detail=f"Zabbix error: {err}") from err

    rows = []
    for problem in problem_list.items:
        rows.append(
            {
                "eventid": problem.event_id,
                "name": problem.name,
                "opdata": problem.opdata,
                "severity": problem.severity,
                "acknowledged": int(problem.acknowledged),
                "clock": problem.clock,
                "clock_iso": problem.clock_iso,
                "tags": list(problem.tags),
                "suppressed": int(problem.suppressed),
                "status": problem.status,
                "host": problem.host_name,
                "hostid": problem.host_id,
                "host_url": problem.host_url,
                "problem_url": problem.problem_url,
            }
        )
    return {"items": rows, "count": len(rows)}


@app.get("/zabbix/host")
def zabbix_host(hostid: int = Query(..., description="Host ID")):
    """Return extended information about a single host for debugging/analysis."""
    try:
        host = _zabbix_client().get_host(hostid)
    except ZabbixAuthError as err:
        raise HTTPException(status_code=401, detail=f"Zabbix error: {err}") from err
    except ZabbixError as err:
        raise HTTPException(status_code=502, detail=f"Zabbix error: {err}") from err
    if not host.raw:
        raise HTTPException(status_code=404, detail="Host not found")
    return {"host": host.raw}


class ZabbixAckRequest(BaseModel):
    eventids: list[str] | list[int]
    message: str | None = None


@app.post("/zabbix/ack")
def zabbix_ack(req: ZabbixAckRequest, request: Request):
    """Acknowledge one or more events in Zabbix.

    Uses event.acknowledge with action=6 (acknowledge + message). Requires API token.
    """
    require_permission(request, "zabbix.ack")
    ids = [str(x) for x in (req.eventids or []) if str(x).strip()]
    if not ids:
        raise HTTPException(status_code=400, detail="No event IDs provided")
    try:
        result = _zabbix_client().acknowledge(ids, message=req.message)
    except ZabbixAuthError as err:
        raise HTTPException(status_code=401, detail=f"Zabbix error: {err}") from err
    except ZabbixError as err:
        raise HTTPException(status_code=502, detail=f"Zabbix error: {err}") from err
    return {"ok": True, "eventids": list(result.succeeded), "result": result.response}


@app.get("/zabbix/history")
def zabbix_history(
    q: str | None = Query(None, description="Optional keyword matched against the problem or host name"),
    severities: str | None = Query(None, description="Comma-separated severities 0..5"),
    groupids: str | None = Query(None, description="Comma-separated group IDs"),
    hostids: str | None = Query(None, description="Comma-separated host IDs"),
    include_subgroups: int = Query(0, ge=0, le=1, description="Include subgroup IDs when filtering by group IDs"),
    hours: int = Query(168, ge=1, le=24 * 90, description="Number of hours to look back"),
    limit: int = Query(100, ge=1, le=500, description="Maximum number of results"),
):
    """Search recent Zabbix problems (including resolved ones) for analysis or AI tooling."""

    try:
        sev_list = [int(s) for s in (severities.split(",") if severities else []) if str(s).strip() != ""]
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid severities")
    try:
        grp_list = [int(s) for s in (groupids.split(",") if groupids else []) if str(s).strip() != ""]
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid groupids")
    try:
        host_list = [int(s) for s in (hostids.split(",") if hostids else []) if str(s).strip() != ""]
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid hostids")

    if not sev_list:
        env_sev = os.getenv("ZABBIX_SEVERITIES", "2,3,4").strip()
        if env_sev:
            try:
                sev_list = [int(x) for x in env_sev.split(",") if x.strip()]
            except Exception:
                sev_list = [2, 3, 4]
        else:
            sev_list = [2, 3, 4]
    if not grp_list:
        gid = os.getenv("ZABBIX_GROUP_ID", "").strip()
        if gid.isdigit():
            grp_list = [int(gid)]

    lookback = datetime.now(UTC) - timedelta(hours=hours)
    time_from = int(lookback.timestamp()) if hours else None
    client = _zabbix_client()
    if grp_list and include_subgroups == 1:
        grp_list = list(client.expand_groupids(grp_list))

    # Gebruik API-zoekopdracht op naam; val terug op host-match als niets wordt gevonden.
    search_term = (q or "").strip()
    primary_search = search_term or None
    try:
        problem_list = client.get_problems(
            severities=sev_list,
            groupids=grp_list,
            hostids=host_list,
            limit=limit,
            recent=True,
            search=primary_search,
            time_from=time_from,
        )
    except ZabbixAuthError as err:
        raise HTTPException(status_code=401, detail=f"Zabbix error: {err}") from err
    except ZabbixError as err:
        raise HTTPException(status_code=502, detail=f"Zabbix error: {err}") from err

    items = list(problem_list.items)
    if search_term:
        term_lower = search_term.lower()
        filtered = [
            it for it in items if term_lower in (it.name or "").lower() or term_lower in (it.host_name or "").lower()
        ]
        if filtered:
            items = filtered
        else:
            try:
                alt = client.get_problems(
                    severities=sev_list,
                    groupids=grp_list,
                    hostids=host_list,
                    limit=limit,
                    recent=True,
                    time_from=time_from,
                )
                items = [
                    it
                    for it in alt.items
                    if term_lower in (it.name or "").lower() or term_lower in (it.host_name or "").lower()
                ]
            except ZabbixError:
                items = []

    rows = []
    for problem in items:
        rows.append(
            {
                "eventid": problem.event_id,
                "name": problem.name,
                "opdata": problem.opdata,
                "severity": problem.severity,
                "acknowledged": int(problem.acknowledged),
                "clock": problem.clock,
                "clock_iso": problem.clock_iso,
                "tags": list(problem.tags),
                "suppressed": int(problem.suppressed),
                "status": problem.status,
                "host": problem.host_name,
                "hostid": problem.host_id,
                "host_url": problem.host_url,
                "problem_url": problem.problem_url,
            }
        )

    return {
        "items": rows,
        "count": len(rows),
        "hours": hours,
        "query": search_term,
        "limit": limit,
    }


# Serve favicon from project package location (png) as /favicon.ico
@app.get("/favicon.ico")
def favicon_ico():
    # Prefer png present at src/enreach_tools/api/favicon.png
    png_path = Path(__file__).parent / "favicon.png"
    if png_path.exists():
        # Serve PNG under .ico path; browsers accept image/png
        return FileResponse(png_path, media_type="image/png")
    # Else, try static path under /app
    static_png = Path(__file__).parent / "static" / "favicon.png"
    if static_png.exists():
        return FileResponse(static_png, media_type="image/png")
    raise HTTPException(status_code=404, detail="favicon not found")


@app.get("/health")
def health():
    d = _data_dir()
    dev = _csv_path("netbox_devices_export.csv")
    vms = _csv_path("netbox_vms_export.csv")
    merged = _csv_path("netbox_merged_export.csv")
    return {
        "status": "ok",
        "data_dir": str(d),
        "files": {
            "devices_csv": dev.exists(),
            "vms_csv": vms.exists(),
            "merged_csv": merged.exists(),
        },
    }


@app.get("/column-order")
def column_order() -> list[str]:
    """Return preferred column order based on Systems CMDB.xlsx if available.

    Falls back to merged CSV headers if Excel not found; otherwise empty list.
    """
    try:
        # Prefer the Excel produced by merge step
        xlsx_path = _csv_path("Systems CMDB.xlsx")
        if xlsx_path.exists():
            try:
                from openpyxl import load_workbook

                wb = load_workbook(xlsx_path, read_only=True, data_only=True)
                ws = wb.worksheets[0]
                headers: list[str] = []
                for cell in ws[1]:
                    v = cell.value
                    if v is not None:
                        headers.append(str(v))
                wb.close()
                if headers:
                    return headers
            except Exception:
                pass
        # Fallback to merged CSV header
        csv_path = _csv_path("netbox_merged_export.csv")
        if csv_path.exists():
            with csv_path.open("r", encoding="utf-8") as fh:
                import csv as _csv

                reader = _csv.reader(fh)
                headers = next(reader, [])
                return [str(h) for h in headers if h]
    except Exception:
        pass
    return []


@app.get("/")
def root_redirect():
    # Serve the frontend at /app/
    return RedirectResponse(url="/app/")


# Mount static frontend
_static_dir = Path(__file__).parent / "static"
app.mount("/app", StaticFiles(directory=_static_dir, html=True), name="app")


@app.get("/logs/tail")
def logs_tail(n: int = Query(200, ge=1, le=5000)) -> dict:
    """
    Return the last N lines of the export log.
    Response: { "lines": ["..", ".."] }
    """
    if not LOG_PATH.exists():
        return {"lines": []}
    try:
        with LOG_PATH.open("r", encoding="utf-8", errors="ignore") as fh:
            lines = fh.readlines()[-n:]
        return {"lines": [ln.rstrip("\n") for ln in lines]}
    except Exception:
        return {"lines": []}


@app.get("/export/stream")
async def export_stream(
    request: Request,
    dataset: Literal["devices", "vms", "all"] = "devices",
    user: OptionalUserDep = None,
):
    """
    Stream the output of an export run for the given dataset.
    - devices -> uv run enreach export devices
    - vms     -> uv run enreach export vms
    - all     -> uv run enreach export update
    """
    require_permission(request, "export.run")

    args_map = {
        "devices": ["enreach", "export", "devices"],
        "vms": ["enreach", "export", "vms"],
        "all": ["enreach", "export", "update"],
    }
    sub = args_map.get(dataset, args_map["devices"])
    if shutil.which("uv"):
        cmd = ["uv", "run", *sub]
    else:
        # Fallback to Python module invocation if uv isn't available
        cmd = [sys.executable, "-m", "enreach_tools.cli", *sub]

    command_str = " ".join(cmd)
    actor = getattr(user, "username", None)

    async def runner():
        with task_logging(
            "export.stream",
            dataset=dataset,
            command=command_str,
            actor=actor,
        ) as task_log:
            start_cmd = f"$ {command_str}"
            yield start_cmd + "\n"
            _write_log(start_cmd)
            proc = await asyncio.create_subprocess_exec(
                *cmd,
                cwd=str(project_root()),
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.STDOUT,
            )
            try:
                if proc.stdout is None:
                    raise RuntimeError("export runner missing stdout pipe")
                while True:
                    line = await proc.stdout.readline()
                    if not line:
                        break
                    try:
                        txt = line.decode(errors="ignore")
                    except Exception:
                        txt = str(line)
                    yield txt
                    _write_log(txt)
            finally:
                rc = await proc.wait()
                exit_line = f"[exit {rc}]"
                yield f"\n{exit_line}\n"
                _write_log(exit_line)
                task_log.add_success(return_code=rc)
                if rc != 0:
                    logger.warning(
                        "Export stream finished with non-zero exit code",
                        extra={
                            "event": "task_warning",
                            "return_code": rc,
                            "dataset": dataset,
                        },
                    )

    return StreamingResponse(runner(), media_type="text/plain")


@app.get("/commvault/backups")
def commvault_backups_data(
    since_hours: int = Query(24, ge=0, le=24 * 90, description="Only include jobs newer than this window (hours)."),
) -> dict[str, Any]:
    """Return cached Commvault backup jobs for the dashboard."""

    cache = _load_commvault_backups()
    jobs = cache.get("jobs") or []
    filtered = _filter_commvault_jobs(jobs, since_hours)
    return {
        "jobs": filtered,
        "generated_at": cache.get("generated_at"),
        "total_cached": cache.get("total_cached", len(jobs)),
        "returned": len(filtered),
        "since_hours": since_hours,
    }


@app.get("/commvault/plans")
def commvault_plans_data(
    refresh: int = Query(0, ge=0, le=1, description="Force refresh of the Commvault plan cache."),
    plan_type: str | None = Query(None, description="Optional case-insensitive plan type filter."),
    limit: int | None = Query(None, ge=1, le=1000, description="Maximum number of plans to return."),
) -> dict[str, Any]:
    """Return cached Commvault plan definitions for the dashboard."""

    cache = _load_commvault_plans()
    if refresh == 1 or not cache.get("plans"):
        try:
            cache = _refresh_commvault_plans_sync(limit=None, plan_type=None)
        except HTTPException as exc:
            if cache.get("plans"):
                logger.warning(
                    "Commvault plan refresh failed; serving cached data",
                    extra={"event": "commvault_plan_refresh_failed", "error": exc.detail if hasattr(exc, "detail") else str(exc)},
                )
            else:
                raise

    plans_payload = list(cache.get("plans") or [])
    plan_type_norm = plan_type.strip().casefold() if plan_type else None
    filtered: list[Mapping[str, Any]] = []
    if plan_type_norm:
        for plan in plans_payload:
            plan_type_value = str(plan.get("plan_type") or "").strip().casefold()
            if plan_type_value == plan_type_norm:
                filtered.append(plan)
    else:
        filtered = plans_payload

    if limit is not None:
        filtered = filtered[:limit]

    return {
        "plans": filtered,
        "total_cached": cache.get("total_cached", len(plans_payload)),
        "generated_at": cache.get("generated_at"),
        "plan_types": cache.get("plan_types") or [],
        "requested_plan_type": plan_type,
        "returned": len(filtered),
        "version": cache.get("version", 1),
    }


@app.post("/commvault/plans/refresh")
async def commvault_plans_refresh(
    limit: int = Query(0, ge=0, le=1000, description="Maximum number of plans to fetch (0 = default)."),
    plan_type: str | None = Query(None, description="Restrict refresh to a specific plan type."),
) -> dict[str, Any]:
    fetch_limit = None if limit == 0 else limit
    return await asyncio.to_thread(_refresh_commvault_plans_sync, fetch_limit, plan_type)


@app.post("/commvault/backups/refresh")
async def commvault_backups_refresh(
    limit: int = Query(0, ge=0, le=500, description="Maximum number of jobs to retain (0 = all)"),
    since_hours: int = Query(24, ge=0, le=24 * 90, description="Only include jobs newer than this window (hours)."),
) -> dict[str, Any]:
    """Fetch fresh Commvault backup jobs and cache them to disk."""

    return await asyncio.to_thread(_refresh_commvault_backups_sync, limit, since_hours)


@app.get("/commvault/backups/file")
def commvault_backups_file(
    file_format: Literal["json", "csv"] = Query("json", description="File format to download"),
) -> FileResponse:
    if file_format == "json":
        path = _data_dir() / COMMVAULT_BACKUPS_JSON
        if not path.exists():
            raise HTTPException(status_code=404, detail="Commvault export not available")
        return FileResponse(path, media_type="application/json", filename=path.name)

    # Generate CSV on demand from cached dataset
    cache = _load_commvault_backups()
    jobs = cache.get("jobs") or []
    if not jobs:
        raise HTTPException(status_code=404, detail="Commvault export not available")

    fieldnames = [
        "job_id",
        "job_type",
        "status",
        "localized_status",
        "localized_operation",
        "client_name",
        "client_id",
        "destination_client_name",
        "subclient_name",
        "backup_set_name",
        "application_name",
        "backup_level_name",
        "plan_name",
        "client_groups",
        "storage_policy_name",
        "start_time",
        "end_time",
        "elapsed_seconds",
        "size_of_application_bytes",
        "size_on_media_bytes",
        "total_num_files",
        "percent_complete",
        "percent_savings",
        "average_throughput_gb_per_hr",
        "retain_until",
    ]

    def _iter_rows():
        import csv as _csv
        from io import StringIO

        buffer = StringIO()
        writer = _csv.DictWriter(buffer, fieldnames=fieldnames)
        writer.writeheader()
        yield buffer.getvalue()
        buffer.seek(0)
        buffer.truncate(0)

        writer = _csv.DictWriter(buffer, fieldnames=fieldnames)
        for job in jobs:
            row = dict(job)
            groups = row.get("client_groups")
            if isinstance(groups, list):
                row["client_groups"] = ";".join(groups)
            writer.writerow({k: row.get(k, "") for k in fieldnames})
            yield buffer.getvalue()
            buffer.seek(0)
            buffer.truncate(0)

    return StreamingResponse(_iter_rows(), media_type="text/csv", headers={"Content-Disposition": "attachment; filename=commvault_backups.csv"})


@app.get("/commvault/storage")
def commvault_storage_data(
    refresh: int = Query(0, ge=0, le=1, description="Force refresh of storage pool cache before returning data."),
) -> dict[str, Any]:
    """Return cached Commvault storage pool data."""

    if refresh == 1:
        return _refresh_commvault_storage_sync()

    cache = _load_commvault_storage()
    if cache.get("pools"):
        return cache

    return _refresh_commvault_storage_sync()


@app.get("/commvault/servers/search")
def commvault_servers_search(
    q: str = Query(..., min_length=1, description="Client name fragment or ID."),
    limit: int = Query(10, ge=1, le=200, description="Maximum number of matches to return."),
) -> dict[str, Any]:
    needle = q.strip().casefold()
    if not needle:
        raise HTTPException(status_code=400, detail="Query must not be empty")

    clients, _, _ = _cached_commvault_clients()
    if not clients:
        return {"results": []}

    matches: list[dict[str, Any]] = []
    for record in clients:
        variants = record.get("name_variants") or set()
        client_id = record.get("client_id")
        if needle.isdigit() and client_id == int(needle):
            matches = [
                {
                    "client_id": client_id,
                    "name": record.get("name"),
                    "display_name": record.get("display_name"),
                    "job_count": record.get("job_count", 0),
                }
            ]
            break
        if any(needle in variant for variant in variants if variant):
            matches.append(
                {
                    "client_id": client_id,
                    "name": record.get("name"),
                    "display_name": record.get("display_name"),
                    "job_count": record.get("job_count", 0),
                }
            )
        if len(matches) >= limit:
            break

    return {"results": matches[:limit]}


@app.get("/commvault/servers/summary")
def commvault_servers_summary(
    client: str = Query(..., description="Commvault client name or ID."),
    job_limit: int = Query(500, ge=0, le=2000, description="Maximum jobs to include (0 = all)."),
    since_hours: int = Query(0, ge=0, le=24 * 365, description="Lookback window in hours (0 = all)."),
    retained_only: bool = Query(True, description="Only include jobs with retention metadata."),
    refresh_cache: bool = Query(False, description="Force bypass of cached job metrics."),
) -> dict[str, Any]:
    summary, metrics, jobs, stats = _load_commvault_server_data(
        client,
        job_limit=job_limit,
        since_hours=since_hours,
        retained_only=retained_only,
        refresh_cache=refresh_cache,
    )
    return {
        "client": summary,
        "job_metrics": metrics,
        "stats": stats,
        "jobs": jobs,
    }


@app.get("/commvault/servers/export")
def commvault_servers_export(
    client: str = Query(..., description="Commvault client name or ID."),
    file_format: Literal["text", "markdown", "html", "docx", "xlsx", "csv"] = Query(
        "xlsx", description="Export format"
    ),
    job_limit: int = Query(500, ge=0, le=2000, description="Maximum jobs to include (0 = all)."),
    since_hours: int = Query(0, ge=0, le=24 * 365, description="Lookback window in hours (0 = all)."),
    retained_only: bool = Query(True, description="Only include jobs with retention metadata."),
    refresh_cache: bool = Query(False, description="Force bypass of cached job metrics."),
) -> Response:
    summary, metrics, jobs, stats = _load_commvault_server_data(
        client,
        job_limit=job_limit,
        since_hours=since_hours,
        retained_only=retained_only,
        refresh_cache=refresh_cache,
    )
    basename_source = summary.get("display_name") or summary.get("name") or f"client-{summary.get('client_id')}"
    basename = _slugify_commvault_name(str(basename_source))

    extensions = {
        "text": "txt",
        "markdown": "md",
        "html": "html",
        "docx": "docx",
        "xlsx": "xlsx",
        "csv": "csv",
    }
    ext = extensions[file_format]
    filename = f"{basename}.{ext}"
    disposition = {"Content-Disposition": f"attachment; filename={filename}"}

    if file_format == "text":
        content = _render_commvault_server_text_report(summary, stats, jobs, metrics)
        return PlainTextResponse(content, headers=disposition)

    if file_format == "markdown":
        content = _render_commvault_server_markdown_report(summary, stats, jobs, metrics)
        return PlainTextResponse(content, headers=disposition, media_type="text/markdown")

    if file_format == "html":
        content = _render_commvault_server_html_report(summary, stats, jobs, metrics)
        return HTMLResponse(content, headers=disposition)

    if file_format == "docx":
        buffer = _render_commvault_server_docx(summary, stats, jobs, metrics)
        return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers=disposition)

    if file_format == "xlsx":
        buffer = _render_commvault_server_xlsx(summary, jobs)
        return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers=disposition)

    buffer = _render_commvault_server_csv(jobs)
    csv_bytes = BytesIO(buffer.getvalue().encode("utf-8"))
    csv_bytes.seek(0)
    return StreamingResponse(csv_bytes, media_type="text/csv", headers=disposition)


# ---------------------------
# Search aggregator (Zabbix, Jira, Confluence, NetBox, vCenter)
# ---------------------------


def _collect_vm_search_values(vm: Any) -> list[str]:
    values: list[str] = []

    def push(raw: Any) -> None:
        if raw is None:
            return
        text = str(raw).strip().lower()
        if text:
            values.append(text)

    push(getattr(vm, "vm_id", None))
    push(getattr(vm, "name", None))
    attrs = [
        "power_state",
        "guest_os",
        "tools_status",
        "host",
        "cluster",
        "datacenter",
        "resource_pool",
        "folder",
        "instance_uuid",
        "bios_uuid",
        "guest_family",
        "guest_name",
        "guest_full_name",
        "guest_host_name",
        "guest_ip_address",
        "tools_run_state",
        "tools_version",
        "tools_version_status",
        "tools_install_type",
        "vcenter_url",
    ]
    for attr in attrs:
        push(getattr(vm, attr, None))
    for seq in (
        getattr(vm, "ip_addresses", ()) or (),
        getattr(vm, "mac_addresses", ()) or (),
        getattr(vm, "tags", ()) or (),
        getattr(vm, "network_names", ()) or (),
    ):
        for item in seq:
            push(item)
    custom_attrs = getattr(vm, "custom_attributes", None)
    if isinstance(custom_attrs, Mapping):
        for key, value in custom_attrs.items():
            push(key)
            push(value)
    return values


def _vcenter_vm_matches(vm: Any, tokens: list[str]) -> bool:
    if not tokens:
        return True
    values = _collect_vm_search_values(vm)
    if not values:
        return False
    return all(any(token in value for value in values) for token in tokens)


def _iso_datetime(value: Any) -> str | None:
    if isinstance(value, datetime):
        return value.astimezone(UTC).isoformat().replace("+00:00", "Z")
    if isinstance(value, str):
        return value
    return None


def _ts_iso(ts: int | str | None) -> str:
    try:
        t = int(ts or 0)
        if t <= 0:
            return ""
        return datetime.utcfromtimestamp(t).strftime("%Y-%m-%d %H:%M:%S")
    except Exception:
        return ""


@app.get("/search/aggregate")
def search_aggregate(
    request: Request,
    q: str = Query(..., description="Object name to search across systems"),
    zlimit: int = Query(10, ge=0, le=500, description="Max Zabbix items per list (0 = no limit)"),
    jlimit: int = Query(10, ge=0, le=200, description="Max Jira issues (0 = no limit, capped upstream)"),
    climit: int = Query(10, ge=0, le=200, description="Max Confluence results (0 = no limit, capped upstream)"),
    vlimit: int = Query(10, ge=0, le=500, description="Max vCenter matches (0 = no limit)"),
):
    out: dict[str, Any] = {"q": q}

    permissions = getattr(request.state, "permissions", frozenset())
    user = getattr(request.state, "user", None)
    user_role = getattr(user, "role", "") if user else ""
    can_view_vcenter = bool(user) and (user_role == "admin" or "vcenter.view" in permissions)
    vcenter_payload: dict[str, Any] = {
        "items": [],
        "errors": [],
        "permitted": can_view_vcenter,
        "has_more": False,
        "total": 0,
    }
    out["vcenter"] = vcenter_payload

    if can_view_vcenter and vlimit != 0:
        tokens = [token for token in q.lower().split() if token.strip()]
        try:
            with SessionLocal() as db:
                service = create_vcenter_service(db)
                configs_with_meta = service.list_configs_with_status()
                match_count = 0
                for config, _meta in configs_with_meta:
                    friendly_name = config.name or config.base_url or config.id
                    try:
                        _, vms, meta_payload = service.get_inventory(config.id, refresh=False)
                    except Exception as exc:  # pragma: no cover - integration path
                        vcenter_payload["errors"].append(f"{friendly_name}: {exc}")
                        continue

                    generated_at = None
                    if isinstance(meta_payload, Mapping):
                        generated_at = _iso_datetime(meta_payload.get("generated_at"))

                    for vm in vms:
                        if tokens and not _vcenter_vm_matches(vm, tokens):
                            continue
                        match_count += 1
                        if vlimit > 0 and len(vcenter_payload["items"]) >= vlimit:
                            vcenter_payload["has_more"] = True
                            break

                        vcenter_payload["items"].append(
                            {
                                "id": vm.vm_id,
                                "name": vm.name,
                                "config_id": config.id,
                                "config_name": friendly_name,
                                "power_state": vm.power_state,
                                "guest_os": vm.guest_os,
                                "tools_status": vm.tools_status,
                                "guest_host_name": vm.guest_host_name,
                                "guest_ip_address": vm.guest_ip_address,
                                "ip_addresses": list(vm.ip_addresses),
                                "mac_addresses": list(vm.mac_addresses),
                                "tags": list(vm.tags),
                                "network_names": list(vm.network_names),
                                "instance_uuid": vm.instance_uuid,
                                "bios_uuid": vm.bios_uuid,
                                "vcenter_url": vm.vcenter_url,
                                "detail_url": f"/app/vcenter/view.html?config={config.id}&vm={vm.vm_id}",
                                "generated_at": generated_at,
                            }
                        )

                    if vcenter_payload["has_more"]:
                        break

                vcenter_payload["total"] = match_count
        except Exception as exc:  # pragma: no cover - defensive fallback
            vcenter_payload["errors"].append(str(exc))

        if vcenter_payload["items"]:
            vcenter_payload["items"].sort(
                key=lambda item: (
                    (item.get("name") or item.get("id") or "").lower(),
                    item.get("config_name") or "",
                )
            )

    # Zabbix: active (problems) and historical (events)
    try:
        client = _zabbix_client()
        hostids: list[int] = []
        try:
            # Fuzzy host search on both 'name' and 'host', allow partial matches and wildcards
            patt = f"*{q}*"
            res = _zbx_rpc(
                "host.get",
                {
                    "output": ["hostid", "host", "name"],
                    "search": {"name": patt, "host": patt},
                    "searchByAny": 1,
                    "searchWildcardsEnabled": 1,
                    "limit": 200,
                },
                client=client,
            )
            for h in res or []:
                try:
                    hostids.append(int(h.get("hostid")))
                except Exception:
                    pass
            # If q looks like an IP, match host interfaces by IP as well
            import re as _re

            if _re.match(r"^\d{1,3}(?:\.\d{1,3}){3}$", q.strip()):
                try:
                    intfs = _zbx_rpc(
                        "hostinterface.get",
                        {"output": ["interfaceid", "hostid", "ip"], "search": {"ip": q.strip()}, "limit": 200},
                        client=client,
                    )
                    for itf in intfs or []:
                        try:
                            hostids.append(int(itf.get("hostid")))
                        except Exception:
                            pass
                except Exception:
                    pass
            # Deduplicate
            hostids = sorted({i for i in hostids if isinstance(i, int)})
        except Exception:
            hostids = []
        zbx = {"active": [], "historical": []}
        base_web = client.web_base or _zbx_web_base() or ""
        # Active problems (prefer hostids; fallback to name search)
        p_params = {
            "output": ["eventid", "name", "severity", "clock", "acknowledged", "r_eventid"],
            "selectTags": "extend",
            "limit": 200,
        }
        if hostids:
            p_params["hostids"] = hostids
        else:
            p_params["search"] = {"name": f"*{q}*"}
            p_params["searchWildcardsEnabled"] = 1
        # Also request hosts to allow client-side fallback filtering
        p_params["selectHosts"] = ["host", "name", "hostid"]
        p = _zbx_rpc("problem.get", p_params, client=client)
        items = []
        try:
            p = sorted(p or [], key=lambda x: int(x.get("clock") or 0), reverse=True)
        except Exception:
            p = p or []
        # Apply limit
        lim = int(zlimit) if int(zlimit) > 0 else len(p)
        for it in p[:lim]:
            items.append(
                {
                    "eventid": it.get("eventid"),
                    "name": it.get("name"),
                    "severity": it.get("severity"),
                    "clock": _ts_iso(it.get("clock")),
                    "acknowledged": it.get("acknowledged"),
                    "resolved": 1 if (str(it.get("r_eventid") or "") not in ("", "0")) else 0,
                    "status": ("ACTIVE" if str(it.get("r_eventid") or "").strip() in ("", "0") else "RESOLVED"),
                    "problem_url": (
                        f"{base_web}/zabbix.php?action=problem.view&eventid={it.get('eventid')}"
                        if base_web and it.get("eventid")
                        else None
                    ),
                    "host_url": (
                        f"{base_web}/zabbix.php?action=host.view&hostid={(it.get('hosts') or [{}])[0].get('hostid')}"
                        if base_web and (it.get("hosts") or [{}])[0].get("hostid")
                        else None
                    ),
                }
            )
        # Extra fallback: if still empty and we didn't have hostids, try a broader recent scan and filter locally
        if not items and not hostids:
            try:
                alt = _zbx_rpc(
                    "problem.get",
                    {
                        "output": ["eventid", "name", "severity", "clock", "acknowledged", "r_eventid"],
                        "selectHosts": ["host", "name", "hostid"],
                        "limit": 200,
                        "sortfield": ["clock"],
                        "sortorder": "DESC",
                    },
                    client=client,
                )
                ql = q.lower().strip()
                for it in alt or []:
                    host_list = it.get("hosts", []) or []
                    host_match = any(
                        (str(h.get("host") or "") + " " + str(h.get("name") or "")).lower().find(ql) >= 0
                        for h in host_list
                    )
                    if host_match or (str(it.get("name") or "").lower().find(ql) >= 0):
                        items.append(
                            {
                                "eventid": it.get("eventid"),
                                "name": it.get("name"),
                                "severity": it.get("severity"),
                                "clock": _ts_iso(it.get("clock")),
                                "acknowledged": it.get("acknowledged"),
                                "resolved": 1 if (str(it.get("r_eventid") or "") not in ("", "0")) else 0,
                            }
                        )
            except Exception:
                pass
        zbx["active"] = items
        # Historical events (prefer hostids; fallback to name search)
        ev_params = {
            "output": ["eventid", "name", "clock", "value"],
            "selectTags": "extend",
            "source": 0,  # triggers
            "limit": 200,
        }
        if hostids:
            ev_params["hostids"] = hostids
        else:
            ev_params["search"] = {"name": f"*{q}*"}
            ev_params["searchWildcardsEnabled"] = 1
        ev = _zbx_rpc("event.get", ev_params, client=client)
        ev_items = []
        try:
            ev = sorted(ev or [], key=lambda x: int(x.get("clock") or 0), reverse=True)
        except Exception:
            ev = ev or []
        limh = int(zlimit) if int(zlimit) > 0 else len(ev)
        for it in ev[:limh]:
            ev_items.append(
                {
                    "eventid": it.get("eventid"),
                    "name": it.get("name"),
                    "clock": _ts_iso(it.get("clock")),
                    "value": it.get("value"),
                    "status": ("PROBLEM" if str(it.get("value") or "").strip() == "1" else "OK"),
                    "event_url": (
                        f"{base_web}/zabbix.php?action=event.view&eventid={it.get('eventid')}"
                        if base_web and it.get("eventid")
                        else None
                    ),
                    "host_url": (
                        f"{base_web}/zabbix.php?action=host.view&hostid={(it.get('hosts') or [{}])[0].get('hostid')}"
                        if base_web and (it.get("hosts") or [{}])[0].get("hostid")
                        else None
                    ),
                }
            )
        zbx["historical"] = ev_items
        out["zabbix"] = zbx
    except HTTPException as ex:
        out["zabbix"] = {"error": ex.detail}
    except Exception as ex:
        out["zabbix"] = {"error": str(ex)}

    # Jira: tickets containing text (last 365d to be practical)
    try:
        mr = int(jlimit) if int(jlimit) > 0 else 50
        res = jira_search(
            q=q,
            jql=None,
            project=None,
            status=None,
            assignee=None,
            priority=None,
            issuetype=None,
            updated="-365d",
            team=None,
            only_open=0,
            max_results=mr,
        )
        out["jira"] = {"total": res.get("total", 0), "issues": res.get("issues", [])}
    except HTTPException as ex:
        out["jira"] = {"error": ex.detail}
    except Exception as ex:
        out["jira"] = {"error": str(ex)}

    # Confluence: pages mentioning the object (last 365d)
    try:
        mc = int(climit) if int(climit) > 0 else 50
        res = confluence_search(q=q, space=None, ctype="page", labels=None, updated="-365d", max_results=mc)
        out["confluence"] = {"total": res.get("total", 0), "results": res.get("results", [])}
    except HTTPException as ex:
        out["confluence"] = {"error": ex.detail}
    except Exception as ex:
        out["confluence"] = {"error": str(ex)}

    # NetBox: objects matching the name; also include IPs when dataset=all
    try:
        # NetBox: no limit by default
        res = netbox_search(dataset="all", q=q, limit=0)
        out["netbox"] = {"total": res.get("total", 0), "items": res.get("rows", [])}
    except HTTPException as ex:
        out["netbox"] = {"error": ex.detail}
    except Exception as ex:
        out["netbox"] = {"error": str(ex)}

    return out


# ---------------------------
# Jira integration (search)
# ---------------------------


def _jira_cfg() -> dict[str, str]:
    """Return Atlassian (Jira) credentials.

    Preferred envs: ATLASSIAN_BASE_URL, ATLASSIAN_EMAIL, ATLASSIAN_API_TOKEN
    Backwards-compatible fallbacks: JIRA_BASE_URL, JIRA_EMAIL, JIRA_API_TOKEN
    """
    base = os.getenv("ATLASSIAN_BASE_URL", "").strip() or os.getenv("JIRA_BASE_URL", "").strip()
    email = os.getenv("ATLASSIAN_EMAIL", "").strip() or os.getenv("JIRA_EMAIL", "").strip()
    token = os.getenv("ATLASSIAN_API_TOKEN", "").strip() or os.getenv("JIRA_API_TOKEN", "").strip()
    return {"base": base, "email": email, "token": token}


def _jira_configured() -> bool:
    cfg = _jira_cfg()
    return bool(cfg["base"] and cfg["email"] and cfg["token"])


def _jira_session() -> tuple[requests.Session, str]:
    cfg = _jira_cfg()
    if not (cfg["base"] and cfg["email"] and cfg["token"]):
        raise HTTPException(
            status_code=400,
            detail="Jira not configured: set ATLASSIAN_BASE_URL, ATLASSIAN_EMAIL, ATLASSIAN_API_TOKEN in .env",
        )
    sess = requests.Session()
    sess.auth = (cfg["email"], cfg["token"])  # Basic auth for Jira Cloud
    sess.headers.update({"Accept": "application/json"})
    base = cfg["base"].rstrip("/")
    return sess, base


def _jira_build_jql(
    q: str | None = None,
    project: str | None = None,
    status: str | None = None,
    assignee: str | None = None,
    priority: str | None = None,
    issuetype: str | None = None,
    updated: str | None = None,
    team: str | None = None,
    only_open: bool = True,
) -> str:
    parts: list[str] = []
    # Only open maps better via statusCategory != Done (workflow agnostic)
    if only_open:
        parts.append("statusCategory != Done")
    if project:
        # Accept both key and name
        p = project.strip()
        if p:
            if any(ch.isspace() for ch in p) or not p.isalnum():
                parts.append(f'project = "{p}"')
            else:
                parts.append(f"project = {p}")
    if status:
        s = status.strip()
        if s:
            # Allow comma separated
            if "," in s:
                vals = ",".join([f'"{v.strip()}"' for v in s.split(",") if v.strip()])
                if vals:
                    parts.append(f"status in ({vals})")
            else:
                parts.append(f'status = "{s}"')
    if assignee:
        a = assignee.strip()
        if a:
            parts.append(f'assignee = "{a}"')
    if priority:
        pr = priority.strip()
        if pr:
            if "," in pr:
                vals = ",".join([f'"{v.strip()}"' for v in pr.split(",") if v.strip()])
                if vals:
                    parts.append(f"priority in ({vals})")
            else:
                parts.append(f'priority = "{pr}"')
    if issuetype:
        it = issuetype.strip()
        if it:
            parts.append(f'issuetype = "{it}"')
    # Custom field: Team (Service Desk) -> cf[10575]
    if team:
        tv = team.strip()
        if tv:
            if "," in tv:
                vals = ",".join([f'"{v.strip()}"' for v in tv.split(",") if v.strip()])
                if vals:
                    parts.append(f"cf[10575] in ({vals})")
            else:
                parts.append(f'cf[10575] = "{tv}"')
    if updated:
        up = updated.strip()
        if up:
            # Accept absolute date (YYYY-MM-DD) or relative (-7d / -4w)
            parts.append(f"updated >= {up}")
    # Jira /search/jql requires bounded queries; if user provided no limiting filters,
    # apply a safe default of last 30 days to avoid 400 errors.
    if not any(
        [project, status, assignee, priority, issuetype, team, (updated and updated.strip()), (q and q.strip())]
    ):
        parts.append("updated >= -30d")
    if q and q.strip():
        # text ~ search across summary, description, comments (Cloud behavior)
        # Escape quotes in q
        qq = q.replace('"', '\\"')
        parts.append(f'text ~ "{qq}"')
    jql = " AND ".join(parts) if parts else "order by updated desc"
    if "order by" not in jql.lower():
        jql += " ORDER BY updated DESC"
    return jql


@app.get("/jira/config")
def jira_config():
    cfg = _jira_cfg()
    return {"configured": _jira_configured(), "base_url": cfg.get("base")}


@app.get("/jira/search")
def jira_search(
    q: str | None = Query(None, description="Free-text search (text ~ '...')"),
    jql: str | None = Query(None, description="Explicit JQL overrides other filters"),
    project: str | None = Query(None),
    status: str | None = Query(None),
    assignee: str | None = Query(None),
    priority: str | None = Query(None),
    issuetype: str | None = Query(None),
    updated: str | None = Query(None, description=">= constraint, e.g. -14d or 2025-01-01"),
    team: str | None = Query(None, description='Team (Servicedesk), e.g. "Systems Infrastructure"'),
    only_open: int = Query(1, ge=0, le=1),
    max_results: int = Query(50, ge=1, le=200),
):
    sess, base = _jira_session()
    # Build JQL
    jql_str = (
        jql.strip()
        if jql and jql.strip()
        else _jira_build_jql(
            q=q,
            project=project,
            status=status,
            assignee=assignee,
            priority=priority,
            issuetype=issuetype,
            updated=updated,
            team=team,
            only_open=bool(only_open),
        )
    )
    fields = [
        "key",
        "summary",
        "status",
        "assignee",
        "priority",
        "updated",
        "created",
        "issuetype",
        "project",
    ]

    # Use the new /search/jql endpoint with GET + query params (legacy /search is removed)
    data: dict[str, Any] | None = None
    used_endpoint = ""
    try:
        url_jql = f"{base}/rest/api/3/search/jql"
        params = {
            "jql": jql_str,
            "startAt": 0,
            "maxResults": int(max_results),
            "fields": ",".join(fields),
        }
        r = sess.get(url_jql, params=params, timeout=60)
        if r.status_code == 401:
            raise HTTPException(status_code=401, detail="Unauthorized: check ATLASSIAN_EMAIL/ATLASSIAN_API_TOKEN")
        if r.status_code == 403:
            raise HTTPException(status_code=403, detail="Forbidden: missing permissions for this JQL/fields")
        if r.status_code == 400:
            # Jira may return a generic 400 for unbounded queries; surface detail
            raise HTTPException(status_code=400, detail=r.text or "Bad request to Jira /search/jql")
        r.raise_for_status()
        data = r.json()
        used_endpoint = "/rest/api/3/search/jql (GET)"
    except HTTPException:
        raise
    except requests.HTTPError as ex:
        msg = getattr(ex.response, "text", "")
        raise HTTPException(status_code=502, detail=f"Jira /search/jql error: {ex} {msg[:300]}")
    except Exception as ex:
        raise HTTPException(status_code=502, detail=f"Jira /search/jql error: {ex}")

    # Normalize issues list from either shape
    issues = []
    if isinstance(data, dict):
        if isinstance(data.get("issues"), list):
            issues = data.get("issues")
        elif isinstance(data.get("results"), list) and data["results"] and isinstance(data["results"][0], dict):
            issues = data["results"][0].get("issues", [])
    out: list[dict[str, Any]] = []
    for it in issues:
        if not isinstance(it, Mapping):
            continue
        k = str(it.get("key") or "")
        fields = it.get("fields")
        f = fields if isinstance(fields, Mapping) else {}
        out.append(
            {
                "key": k,
                "summary": str(f.get("summary") or ""),
                "status": str((f.get("status") or {}).get("name") or ""),
                "assignee": str((f.get("assignee") or {}).get("displayName") or ""),
                "priority": str((f.get("priority") or {}).get("name") or ""),
                "issuetype": str((f.get("issuetype") or {}).get("name") or ""),
                "project": (
                    str((f.get("project") or {}).get("key") or "")
                    or str((f.get("project") or {}).get("name") or "")
                ),
                "updated": str(f.get("updated") or ""),
                "created": str(f.get("created") or ""),
                "url": f"{base}/browse/{k}" if k else "",
            }
        )
    total = 0
    if isinstance(data, dict):
        # New endpoint may not return 'total'; compute from page or use provided
        total = int(data.get("total", 0) or 0)
        if not total and isinstance(data.get("isLast"), bool):
            total = len(out)
        if (
            not total
            and isinstance(data.get("results"), list)
            and data["results"]
            and isinstance(data["results"][0], dict)
        ):
            total = int(data["results"][0].get("total", 0) or 0)
        if not total:
            total = len(out)
    else:
        total = len(out)
    return {"total": total, "issues": out, "jql": jql_str, "endpoint": used_endpoint}


"""
Confluence integration (read-only search)
- Uses same ATLASSIAN_* credentials
- Queries CQL via /wiki/rest/api/search (GET) with bounded defaults
"""


def _conf_session() -> tuple[requests.Session, str]:
    sess, base = _jira_session()
    # Confluence Cloud REST base is under /wiki
    wiki = base.rstrip("/") + "/wiki"
    return sess, wiki


def _cql_build(
    q: str | None = None,
    space: str | None = None,
    ctype: str | None = None,
    labels: str | None = None,
    updated: str | None = None,
) -> str:
    parts: list[str] = []
    if space and space.strip():
        s = space.strip()
        # If looks like a key (no spaces), use space = "KEY"; otherwise match by title
        esc = s.replace('"', '\\"')
        if any(ch.isspace() for ch in s):
            parts.append(f'space.title = "{esc}"')
        else:
            parts.append(f'space = "{esc}"')
    if ctype and ctype.strip():
        # Confluence types: page, blogpost, attachment, comment, etc.
        parts.append(f"type = {ctype.strip()}")
    if labels and labels.strip():
        arr = [v.strip() for v in labels.split(",") if v.strip()]
        if len(arr) == 1:
            parts.append(f"label = '{arr[0]}'")
        elif arr:
            parts.append("(" + " OR ".join([f"label = '{v}'" for v in arr]) + ")")
    if updated and updated.strip():
        up = updated.strip()
        if up.startswith("-"):
            parts.append(f"lastmodified >= now('{up}')")
        else:
            parts.append(f"lastmodified >= '{up}'")
    # Add text query last to help relevance
    if q and q.strip():
        qq = q.replace('"', '\\"')
        parts.append(f'text ~ "{qq}"')
    # Bound the query if still empty (avoid unbounded errors/pagination surprises)
    if not parts:
        parts.append("lastmodified >= now(-90d)")
    # Order by last modified desc
    cql = " AND ".join(parts)
    cql += " order by lastmodified desc"
    return cql


@app.get("/confluence/config")
def confluence_config():
    cfg = _jira_cfg()
    ok = bool(cfg.get("base") and cfg.get("email") and cfg.get("token"))
    base = (cfg.get("base") or "").rstrip("/")
    return {"configured": ok, "base_url": base + "/wiki" if ok else base}


@app.get("/confluence/search")
def confluence_search(
    q: str | None = Query(None, description="Full-text query"),
    space: str | None = Query(None, description="Space key (e.g., DOCS)"),
    ctype: str | None = Query("page", description="Type: page, blogpost, attachment"),
    labels: str | None = Query(None, description="Comma-separated labels"),
    updated: str | None = Query(None, description="-30d or 2025-01-01"),
    max_results: int = Query(50, ge=1, le=100),
):
    sess, wiki = _conf_session()

    # Resolve space names to keys when needed (names often contain spaces; CQL expects keys)
    def _resolve_space_keys(raw: str) -> list[str]:
        toks = [t.strip() for t in (raw or "").split(",") if t.strip()]
        keys: list[str] = []
        for t in toks:
            # Likely a key if no spaces and matches typical key charset
            if t and all((ch.isalnum() or ch in ("_", "-")) for ch in t) and (not any(ch.isspace() for ch in t)):
                keys.append(t)
                continue
            # 1) Lookup by name using CQL; then keep only exact title/name matches
            exact_keys: list[str] = []
            try:
                esc = t.replace('"', '\\"')
                url_s = wiki + "/rest/api/search"
                r_s = sess.get(url_s, params={"cql": f'type = space AND title ~ "{esc}"', "limit": 50}, timeout=30)
                if r_s.ok:
                    data_s = r_s.json()
                    for it in data_s.get("results", []) or []:
                        sp = it.get("space", {}) if isinstance(it, dict) else {}
                        name = (
                            (sp.get("name") or it.get("title") or "")
                            if isinstance(sp, dict)
                            else (it.get("title") or "")
                        )
                        if isinstance(name, str) and name.strip().lower() == t.strip().lower():
                            k = sp.get("key") if isinstance(sp, dict) else None
                            if k and (k not in exact_keys):
                                exact_keys.append(k)
            except Exception:
                pass
            # 2) Fallback: spaces REST listing filtered by q, then exact name match
            if not exact_keys:
                try:
                    rs = sess.get(wiki + "/rest/api/space", params={"q": t, "limit": 50}, timeout=30)
                    if rs.ok:
                        ds = rs.json()
                        for sp in ds.get("results", []) or []:
                            nm = sp.get("name") or ""
                            if isinstance(nm, str) and nm.strip().lower() == t.strip().lower():
                                k = sp.get("key") or ""
                                if k and (k not in exact_keys):
                                    exact_keys.append(k)
                except Exception:
                    pass
            # Only add exact match keys to avoid partial-space spills
            keys.extend([k for k in exact_keys if k and k not in keys])
        return keys

    space_keys: list[str] = []
    if space and space.strip():
        space_keys = _resolve_space_keys(space)

    cql = _cql_build(q=q, space=None, ctype=ctype, labels=labels, updated=updated)
    if space and space.strip() and not space_keys:
        # Space provided but not resolved exactly -> return empty set
        return {"total": 0, "cql": f"space unresolved: {space}", "results": []}
    if space_keys:
        quoted_keys = ", ".join(f'"{k}"' for k in space_keys)
        cql = f"space in ({quoted_keys}) AND " + cql
    url = wiki + "/rest/api/search"
    # Ask Confluence to include space + history info so we can display Space and Updated reliably
    params = {"cql": cql, "limit": int(max_results), "expand": "content.space,content.history"}
    try:
        r = sess.get(url, params=params, timeout=60)
        if r.status_code == 401:
            raise HTTPException(status_code=401, detail="Unauthorized: check ATLASSIAN_EMAIL/ATLASSIAN_API_TOKEN")
        if r.status_code == 403:
            raise HTTPException(status_code=403, detail="Forbidden: missing permissions for this CQL/fields")
        r.raise_for_status()
        data = r.json()
    except requests.HTTPError as ex:
        msg = getattr(ex.response, "text", "")
        raise HTTPException(status_code=502, detail=f"Confluence error: {ex} {msg[:300]}")
    except Exception as ex:
        raise HTTPException(status_code=502, detail=f"Confluence error: {ex}")

    items = data.get("results", []) if isinstance(data, dict) else []
    out: list[dict[str, Any]] = []
    for it in items or []:
        if not isinstance(it, Mapping):
            continue
        content_raw = it.get("content")
        content = content_raw if isinstance(content_raw, Mapping) else {}
        title = str(content.get("title") or it.get("title") or "")
        ctype_val = str(content.get("type") or it.get("type") or "")
        space_obj_raw = content.get("space")
        space_obj = space_obj_raw if isinstance(space_obj_raw, Mapping) else {}
        space_key = space_obj.get("key") if isinstance(space_obj, Mapping) else None
        space_name = space_obj.get("name") if isinstance(space_obj, Mapping) else None
        if not space_name:
            rgc_raw = it.get("resultGlobalContainer")
            rgc = rgc_raw if isinstance(rgc_raw, Mapping) else {}
            if rgc:
                space_name = space_name or rgc.get("title")
                disp = rgc.get("displayUrl") or ""
                if (not space_key) and isinstance(disp, str) and "/spaces/" in disp:
                    parts = disp.split("/spaces/")
                    if len(parts) > 1:
                        tail = parts[1]
                        space_key = tail.split("/")[0]
        links_raw = content.get("_links") or it.get("_links")
        links = links_raw if isinstance(links_raw, Mapping) else {}
        webui = links.get("webui") or links.get("base")
        link = (
            wiki + webui
            if isinstance(webui, str) and webui.startswith("/")
            else (wiki + "/" + webui if isinstance(webui, str) and webui else "")
        )
        lastmod = None
        hist_raw = content.get("history")
        hist = hist_raw if isinstance(hist_raw, Mapping) else {}
        last = hist.get("lastUpdated") if isinstance(hist, Mapping) else None
        if isinstance(last, Mapping):
            lastmod = last.get("when")
        if not lastmod:
            lastmod = it.get("lastModified") or it.get("friendlyLastModified") or ""
        out.append(
            {
                "title": title,
                "type": ctype_val,
                "space": (space_name or space_key or ""),
                "space_key": (space_key or ""),
                "space_name": (space_name or ""),
                "updated": lastmod or "",
                "url": link,
            }
        )
    total = int(data.get("size", 0) or 0) if isinstance(data, dict) else len(out)
    if not total:
        total = len(out)
    return {"total": total, "cql": cql, "results": out}


@app.get("/devices")
def devices(
    limit: int | None = Query(None, ge=1, description="Max rows to return (omit for all)"),
    offset: int = Query(0, ge=0),
    order_by: str | None = Query(None, description="Column name to order by"),
    order_dir: Literal["asc", "desc"] = Query("asc"),
):
    return _list_records("netbox_devices_export.csv", limit, offset, order_by, order_dir)


@app.get("/vms")
def vms(
    limit: int | None = Query(None, ge=1, description="Max rows to return (omit for all)"),
    offset: int = Query(0, ge=0),
    order_by: str | None = Query(None, description="Column name to order by"),
    order_dir: Literal["asc", "desc"] = Query("asc"),
):
    return _list_records("netbox_vms_export.csv", limit, offset, order_by, order_dir)


@app.get("/all")
def all_merged(
    limit: int | None = Query(None, ge=1, description="Max rows to return (omit for all)"),
    offset: int = Query(0, ge=0),
    order_by: str | None = Query(None, description="Column name to order by"),
    order_dir: Literal["asc", "desc"] = Query("asc"),
):
    return _list_records("netbox_merged_export.csv", limit, offset, order_by, order_dir)


@app.get("/netbox/config")
def netbox_config():
    """Return minimal NetBox config for the UI (base URL only)."""
    base = os.getenv("NETBOX_URL", "").strip()
    return {"configured": bool(base), "base_url": base}


def _nb_session() -> tuple[requests.Session, str]:
    base = os.getenv("NETBOX_URL", "").strip()
    token = os.getenv("NETBOX_TOKEN", "").strip()
    if not base or not token:
        raise HTTPException(status_code=400, detail="NETBOX_URL/NETBOX_TOKEN not configured in .env")
    sess = requests.Session()
    sess.headers.update({"Authorization": f"Token {token}", "Accept": "application/json"})
    try:
        from .env import apply_extra_headers as _apply

        _apply(sess)
    except Exception:
        pass
    return sess, base.rstrip("/")


@app.get("/netbox/search")
def netbox_search(
    dataset: Literal["devices", "vms", "all"] = Query("all"),
    q: str = Query("", description="Full-text query passed to NetBox ?q="),
    limit: int = Query(50, ge=0, le=5000, description="0 = no limit (fetch all pages)"),
):
    """Search NetBox live (no CSV) using the built-in ?q= filter.

    Returns rows with common fields across devices/VMs and a suggested column list.
    """
    if not (q and q.strip()):
        return {"columns": [], "rows": [], "total": 0}
    sess, base = _nb_session()

    def _status_label(x):
        if isinstance(x, dict):
            return x.get("label") or x.get("value") or x.get("name") or ""
        return str(x or "")

    def _get(addr):
        r = sess.get(addr, timeout=30)
        if r.status_code in {401, 403}:
            raise HTTPException(status_code=r.status_code, detail=f"NetBox auth failed: {r.text[:200]}")
        r.raise_for_status()
        return r.json()

    def _collect(endpoint: str, q: str, max_items: int | None) -> list[dict]:
        items: list[dict] = []
        # NetBox uses DRF pagination: limit/offset/next
        page_limit = 200  # reasonable page size
        url = f"{base}{endpoint}?q={requests.utils.quote(q)}&limit={page_limit}&offset=0"
        while url:
            data = _get(url)
            results = data.get("results", []) if isinstance(data, dict) else []
            if not isinstance(results, list):
                break
            items.extend(results)
            if max_items is not None and len(items) >= max_items:
                return items[:max_items]
            url = data.get("next") if isinstance(data, dict) else None
        return items

    def _map_device(it):
        name = it.get("name") or ""
        site = (it.get("site") or {}).get("name") or ""
        tenant = (it.get("tenant") or {}).get("name") or ""
        role = (it.get("device_role") or it.get("role") or {}).get("name") or ""
        status = _status_label(it.get("status"))
        pip4 = (it.get("primary_ip4") or {}).get("address") or ""
        pip6 = (it.get("primary_ip6") or {}).get("address") or ""
        pip = pip4 or pip6
        platform = (it.get("platform") or {}).get("name") or ""
        dtype = (it.get("device_type") or {}).get("model") or ""
        # Try to find an explicit out-of-band management IP from custom fields if present
        cf = it.get("custom_fields") or {}
        oob = ""
        try:
            if isinstance(cf, dict):
                # Common variants people use for OOB/IPMI management
                for key in [
                    "oob_ip",
                    "oob_ip4",
                    "oob_ip6",
                    "out_of_band_ip",
                    "out_of_band",
                    "management_ip",
                    "mgmt_ip",
                    "mgmt_ip4",
                    "mgmt_ip6",
                ]:
                    val = cf.get(key)
                    if isinstance(val, str | int | float) and str(val).strip():
                        oob = str(val).strip()
                        break
        except Exception:
            pass
        if not oob:
            oob = pip  # fallback to primary IP when no explicit OOB is found
        ui_path = f"/dcim/devices/{it.get('id')}/" if it.get("id") is not None else ""
        updated = it.get("last_updated") or it.get("last_updated") or ""
        return {
            "Name": name,
            "Status": status,
            "Site": site,
            "Role": role,
            "Tenant": tenant,
            "Primary IP": pip,
            "Out-of-band IP": oob,
            "Platform": platform,
            "Device Type": dtype,
            "Updated": updated,
            "ui_path": ui_path,
        }

    def _map_vm(it):
        name = it.get("name") or ""
        status = _status_label(it.get("status"))
        tenant = (it.get("tenant") or {}).get("name") or ""
        role = (it.get("role") or {}).get("name") or ""
        cluster = (it.get("cluster") or {}).get("name") or ""
        pip4 = (it.get("primary_ip4") or {}).get("address") or ""
        pip6 = (it.get("primary_ip6") or {}).get("address") or ""
        pip = pip4 or pip6
        ui_path = f"/virtualization/virtual-machines/{it.get('id')}/" if it.get("id") is not None else ""
        updated = it.get("last_updated") or ""
        return {
            "Name": name,
            "Status": status,
            "Cluster": cluster,
            "Role": role,
            "Tenant": tenant,
            "Primary IP": pip,
            "Updated": updated,
            "Out-of-band IP": "",
            "ui_path": ui_path,
        }

    rows: list[dict[str, Any]] = []
    try:
        max_items = None if int(limit) == 0 else int(limit)
        if dataset in ("devices", "all"):
            results = _collect("/api/dcim/devices/", q, max_items)
            for it in results:
                d = _map_device(it)
                if dataset == "all":
                    d["Type"] = "device"
                rows.append(d)
        if dataset in ("vms", "all"):
            results = _collect("/api/virtualization/virtual-machines/", q, max_items)
            for it in results:
                v = _map_vm(it)
                if dataset == "all":
                    v["Type"] = "vm"
                rows.append(v)
        # Always include IP addresses when searching 'all'
        if dataset == "all":

            def _map_ip(it: dict) -> dict[str, Any]:
                addr = it.get("address") or ""
                status = _status_label(it.get("status"))
                vrf = (it.get("vrf") or {}).get("name") or ""
                assigned = ""
                ao = it.get("assigned_object") or {}
                if isinstance(ao, dict):
                    assigned = ao.get("display") or ao.get("name") or ""
                ui_path = f"/ipam/ip-addresses/{it.get('id')}/" if it.get("id") is not None else ""
                updated = it.get("last_updated") or ""
                return {
                    "Name": addr,
                    "Status": status,
                    "VRF": vrf,
                    "Assigned Object": assigned,
                    "Primary IP": "",
                    "Out-of-band IP": "",
                    "Type": "ip address",
                    "Updated": updated,
                    "ui_path": ui_path,
                }

            ip_results = _collect("/api/ipam/ip-addresses/", q, max_items)
            for it in ip_results:
                rows.append(_map_ip(it))
    except HTTPException:
        raise
    except Exception as ex:
        raise HTTPException(status_code=502, detail=f"NetBox search error: {ex}")

    # Determine columns from first row
    columns: list[str] = []
    if rows:
        keys = list(rows[0].keys())
        # Hide internal helper field from table
        if "ui_path" in keys:
            keys.remove("ui_path")
        columns = keys
    return {"columns": columns, "rows": rows, "total": len(rows)}
