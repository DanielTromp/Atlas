"""Export handlers for generating files (xlsx, csv, txt, docx).

These handlers generate files that can be uploaded to chat platforms.
Supports automatic cleanup of old exports.
"""

from __future__ import annotations

import csv
import io
import tempfile
import time
from datetime import datetime
from pathlib import Path
from typing import Any

from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

from infrastructure_atlas.infrastructure.logging import get_logger

logger = get_logger(__name__)

# Export directory
EXPORT_DIR = Path(tempfile.gettempdir()) / "atlas_exports"

# Cleanup threshold (24 hours in seconds)
CLEANUP_THRESHOLD_SECONDS = 24 * 60 * 60


def get_export_dir() -> Path:
    """Get the export directory, creating it if needed."""
    EXPORT_DIR.mkdir(exist_ok=True)
    return EXPORT_DIR


def cleanup_old_exports(max_age_seconds: int = CLEANUP_THRESHOLD_SECONDS) -> int:
    """Remove export files older than max_age_seconds.

    Args:
        max_age_seconds: Maximum age in seconds (default: 24 hours)

    Returns:
        Number of files deleted
    """
    if not EXPORT_DIR.exists():
        return 0

    deleted = 0
    now = time.time()

    for file_path in EXPORT_DIR.iterdir():
        if file_path.is_file():
            file_age = now - file_path.stat().st_mtime
            if file_age > max_age_seconds:
                try:
                    file_path.unlink()
                    deleted += 1
                    logger.debug(f"Deleted old export: {file_path.name}")
                except OSError as e:
                    logger.warning(f"Failed to delete {file_path}: {e}")

    if deleted > 0:
        logger.info(f"Cleaned up {deleted} old export files")

    return deleted


def generate_export(
    data: list[dict[str, Any]] | list[list[Any]] | list[str],
    filename: str,
    file_format: str = "xlsx",
    title: str | None = None,
    sheet_name: str = "Data",
) -> dict[str, Any]:
    """Generate an export file in the specified format.

    Args:
        data: Data to export
        filename: Base filename (without extension)
        file_format: Output format (xlsx, csv, txt, docx)
        title: Optional title for the document
        sheet_name: Sheet name for xlsx exports

    Returns:
        Dictionary with file_path, filename, row_count, etc.
    """
    # Cleanup old exports first (non-blocking)
    try:
        cleanup_old_exports()
    except Exception as e:
        logger.warning(f"Cleanup failed: {e}")

    # Route to appropriate handler
    format_lower = file_format.lower().strip(".")
    
    if format_lower in ("xlsx", "excel", "xls"):
        return generate_xlsx(data, filename, sheet_name=sheet_name, title=title)
    elif format_lower == "csv":
        return generate_csv(data, filename, title=title)
    elif format_lower in ("txt", "text"):
        return generate_txt(data, filename, title=title)
    elif format_lower in ("docx", "doc", "word"):
        return generate_docx(data, filename, title=title)
    else:
        return {
            "error": f"Unsupported format: {file_format}. Supported: xlsx, csv, txt, docx",
            "success": False,
        }


def generate_xlsx(
    data: list[dict[str, Any]] | list[list[Any]] | list[str],
    filename: str,
    sheet_name: str = "Data",
    title: str | None = None,
) -> dict[str, Any]:
    """Generate a formatted Excel file from data.

    Args:
        data: Data to export. Can be:
            - List of dictionaries (each dict is a row)
            - List of lists (first list is headers)
            - List of strings (will be split by common delimiters)
        filename: Base filename (without extension)
        sheet_name: Name for the worksheet
        title: Optional title row at the top

    Returns:
        Dictionary with file_path, filename, and row_count
    """
    if not data:
        return {
            "error": "No data to export",
            "success": False,
        }

    # Normalize data to list of dicts
    normalized_data = _normalize_data(data)
    if not normalized_data:
        return {
            "error": "Could not parse data format",
            "success": False,
        }

    logger.debug(f"Normalized {len(data)} items to {len(normalized_data)} rows")

    # Create workbook
    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name[:31]  # Excel sheet name limit

    # Styles
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    title_font = Font(bold=True, size=14)
    title_fill = PatternFill(start_color="D9E2F3", end_color="D9E2F3", fill_type="solid")

    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )

    current_row = 1

    # Add title row if provided
    if title:
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(normalized_data[0]) if normalized_data else 1)
        title_cell = ws.cell(row=1, column=1, value=title)
        title_cell.font = title_font
        title_cell.fill = title_fill
        title_cell.alignment = Alignment(horizontal="center", vertical="center")
        current_row = 2
        # Add empty row after title
        current_row = 3

    # Get all unique headers from data
    headers: list[str] = []
    for row in normalized_data:
        for key in row.keys():
            if key not in headers:
                headers.append(key)

    # Write headers
    header_row = current_row
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=header_row, column=col_idx, value=_format_header(header))
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        cell.border = thin_border

    # Write data rows
    for row_idx, row_data in enumerate(normalized_data, header_row + 1):
        for col_idx, header in enumerate(headers, 1):
            value = row_data.get(header, "")
            cell = ws.cell(row=row_idx, column=col_idx, value=_format_value(value))
            cell.border = thin_border
            cell.alignment = Alignment(vertical="center", wrap_text=True)

            # Apply number formatting for dates
            if isinstance(value, datetime):
                cell.number_format = "YYYY-MM-DD HH:MM"

    # Auto-size columns
    for col_idx, header in enumerate(headers, 1):
        column_letter = get_column_letter(col_idx)

        # Calculate max width
        max_length = len(str(header))
        for row in normalized_data:
            value = row.get(header, "")
            cell_length = len(str(value)) if value else 0
            max_length = max(max_length, min(cell_length, 50))  # Cap at 50 chars

        # Set column width (add padding)
        ws.column_dimensions[column_letter].width = min(max_length + 2, 50)

    # Freeze header row
    ws.freeze_panes = ws.cell(row=header_row + 1, column=1)

    # Add auto-filter
    ws.auto_filter.ref = f"A{header_row}:{get_column_letter(len(headers))}{header_row + len(normalized_data)}"

    # Generate filename with timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_filename = "".join(c for c in filename if c.isalnum() or c in "-_").strip()
    full_filename = f"{safe_filename}_{timestamp}.xlsx"

    # Save to temp file
    temp_dir = get_export_dir()
    file_path = temp_dir / full_filename

    wb.save(file_path)

    logger.info(f"Generated xlsx export: {file_path} ({len(normalized_data)} rows)")

    return {
        "success": True,
        "file_path": str(file_path),
        "filename": full_filename,
        "row_count": len(normalized_data),
        "column_count": len(headers),
        "file_type": "xlsx",
        "message": f"Excel file generated with {len(normalized_data)} rows and {len(headers)} columns",
    }


def generate_csv(
    data: list[dict[str, Any]] | list[list[Any]] | list[str],
    filename: str,
    title: str | None = None,
    delimiter: str = ",",
) -> dict[str, Any]:
    """Generate a CSV file from data.

    Args:
        data: Data to export
        filename: Base filename (without extension)
        title: Optional title (will be added as comment)
        delimiter: CSV delimiter (default: comma)

    Returns:
        Dictionary with file_path, filename, and row_count
    """
    if not data:
        return {
            "error": "No data to export",
            "success": False,
        }

    # Normalize data to list of dicts
    normalized_data = _normalize_data(data)
    if not normalized_data:
        return {
            "error": "Could not parse data format",
            "success": False,
        }

    # Get all unique headers
    headers: list[str] = []
    for row in normalized_data:
        for key in row.keys():
            if key not in headers:
                headers.append(key)

    # Generate filename with timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_filename = "".join(c for c in filename if c.isalnum() or c in "-_").strip()
    full_filename = f"{safe_filename}_{timestamp}.csv"

    # Write CSV
    temp_dir = get_export_dir()
    file_path = temp_dir / full_filename

    with open(file_path, "w", newline="", encoding="utf-8") as f:
        # Add title as comment if provided
        if title:
            f.write(f"# {title}\n")

        writer = csv.DictWriter(f, fieldnames=headers, delimiter=delimiter, extrasaction="ignore")
        writer.writeheader()

        for row in normalized_data:
            # Convert complex values to strings
            clean_row = {k: _format_value(v) for k, v in row.items()}
            writer.writerow(clean_row)

    logger.info(f"Generated csv export: {file_path} ({len(normalized_data)} rows)")

    return {
        "success": True,
        "file_path": str(file_path),
        "filename": full_filename,
        "row_count": len(normalized_data),
        "column_count": len(headers),
        "file_type": "csv",
        "message": f"CSV file generated with {len(normalized_data)} rows and {len(headers)} columns",
    }


def generate_txt(
    data: list[dict[str, Any]] | list[list[Any]] | list[str],
    filename: str,
    title: str | None = None,
) -> dict[str, Any]:
    """Generate a plain text file with tabular data.

    Args:
        data: Data to export
        filename: Base filename (without extension)
        title: Optional title

    Returns:
        Dictionary with file_path, filename, and row_count
    """
    if not data:
        return {
            "error": "No data to export",
            "success": False,
        }

    # Normalize data to list of dicts
    normalized_data = _normalize_data(data)
    if not normalized_data:
        return {
            "error": "Could not parse data format",
            "success": False,
        }

    # Get all unique headers
    headers: list[str] = []
    for row in normalized_data:
        for key in row.keys():
            if key not in headers:
                headers.append(key)

    # Calculate column widths
    col_widths = {h: len(_format_header(h)) for h in headers}
    for row in normalized_data:
        for header in headers:
            value = str(_format_value(row.get(header, "")))
            col_widths[header] = max(col_widths[header], min(len(value), 40))

    # Build text content
    lines = []

    # Title
    if title:
        lines.append(title)
        lines.append("=" * len(title))
        lines.append("")

    # Header row
    header_line = " | ".join(_format_header(h).ljust(col_widths[h]) for h in headers)
    lines.append(header_line)
    lines.append("-" * len(header_line))

    # Data rows
    for row in normalized_data:
        row_values = []
        for header in headers:
            value = str(_format_value(row.get(header, "")))
            # Truncate long values
            if len(value) > 40:
                value = value[:37] + "..."
            row_values.append(value.ljust(col_widths[header]))
        lines.append(" | ".join(row_values))

    # Generate filename with timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_filename = "".join(c for c in filename if c.isalnum() or c in "-_").strip()
    full_filename = f"{safe_filename}_{timestamp}.txt"

    # Write file
    temp_dir = get_export_dir()
    file_path = temp_dir / full_filename

    with open(file_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    logger.info(f"Generated txt export: {file_path} ({len(normalized_data)} rows)")

    return {
        "success": True,
        "file_path": str(file_path),
        "filename": full_filename,
        "row_count": len(normalized_data),
        "column_count": len(headers),
        "file_type": "txt",
        "message": f"Text file generated with {len(normalized_data)} rows and {len(headers)} columns",
    }


def generate_docx(
    data: list[dict[str, Any]] | list[list[Any]] | list[str],
    filename: str,
    title: str | None = None,
) -> dict[str, Any]:
    """Generate a Word document with a formatted table.

    Args:
        data: Data to export
        filename: Base filename (without extension)
        title: Optional document title

    Returns:
        Dictionary with file_path, filename, and row_count
    """
    try:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
    except ImportError:
        return {
            "error": "python-docx not installed. Run: pip install python-docx",
            "success": False,
        }

    if not data:
        return {
            "error": "No data to export",
            "success": False,
        }

    # Normalize data to list of dicts
    normalized_data = _normalize_data(data)
    if not normalized_data:
        return {
            "error": "Could not parse data format",
            "success": False,
        }

    # Get all unique headers
    headers: list[str] = []
    for row in normalized_data:
        for key in row.keys():
            if key not in headers:
                headers.append(key)

    # Create document
    doc = Document()

    # Add title if provided
    if title:
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add generation info
    doc.add_paragraph(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | "
        f"Rows: {len(normalized_data)} | Columns: {len(headers)}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Spacer

    # Create table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        cell = header_cells[idx]
        cell.text = _format_header(header)
        # Style header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Data rows
    for row_data in normalized_data:
        row_cells = table.add_row().cells
        for idx, header in enumerate(headers):
            value = _format_value(row_data.get(header, ""))
            row_cells[idx].text = str(value) if value else ""
            # Style data cells
            for paragraph in row_cells[idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    # Generate filename with timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_filename = "".join(c for c in filename if c.isalnum() or c in "-_").strip()
    full_filename = f"{safe_filename}_{timestamp}.docx"

    # Save file
    temp_dir = get_export_dir()
    file_path = temp_dir / full_filename

    doc.save(file_path)

    logger.info(f"Generated docx export: {file_path} ({len(normalized_data)} rows)")

    return {
        "success": True,
        "file_path": str(file_path),
        "filename": full_filename,
        "row_count": len(normalized_data),
        "column_count": len(headers),
        "file_type": "docx",
        "message": f"Word document generated with {len(normalized_data)} rows and {len(headers)} columns",
    }


def _normalize_data(data: list[Any]) -> list[dict[str, Any]]:
    """Normalize various data formats to list of dictionaries.

    Handles:
    - List of dicts: returns as-is
    - List of lists: first list becomes headers
    - List of strings: attempts to parse as CSV/delimited or JSON
    - Single string: attempts to parse as JSON
    """
    import json

    if not data:
        return []

    # If data is a single string, try to parse as JSON first
    if isinstance(data, str):
        try:
            parsed = json.loads(data)
            if isinstance(parsed, list):
                data = parsed
            else:
                logger.warning(f"JSON string parsed to non-list: {type(parsed)}")
                return [{"value": str(parsed)}]
        except json.JSONDecodeError:
            logger.warning("Could not parse string as JSON")
            return [{"value": data}]

    first_item = data[0]

    # Already list of dicts
    if isinstance(first_item, dict):
        return data

    # List of lists - first row is headers
    if isinstance(first_item, (list, tuple)):
        if len(data) < 2:
            # Only headers, no data
            return []
        headers = [str(h) for h in first_item]
        return [
            {headers[i]: row[i] if i < len(row) else "" for i in range(len(headers))}
            for row in data[1:]
        ]

    # List of strings - try to parse as JSON first, then CSV
    if isinstance(first_item, str):
        # Check if the first item looks like JSON (list or object)
        first_stripped = first_item.strip()
        if first_stripped.startswith("[") or first_stripped.startswith("{"):
            # Try to parse the entire list as a single JSON string
            full_text = "".join(str(item) for item in data)
            try:
                parsed = json.loads(full_text)
                if isinstance(parsed, list) and parsed and isinstance(parsed[0], dict):
                    return parsed
            except json.JSONDecodeError:
                pass

        # Try common delimiters
        for delimiter in [",", "\t", ";", "|"]:
            if delimiter in first_item:
                rows = [line.split(delimiter) for line in data if line.strip()]
                if len(rows) >= 2:
                    headers = [h.strip() for h in rows[0]]
                    return [
                        {headers[i]: row[i].strip() if i < len(row) else "" for i in range(len(headers))}
                        for row in rows[1:]
                    ]

        # No delimiter found - treat each string as a single-column row
        return [{"value": str(item)} for item in data]

    # Unknown format - try to convert to string
    logger.warning(f"Unknown data format: {type(first_item)}, converting to strings")
    return [{"value": str(item)} for item in data]


def _format_header(header: str) -> str:
    """Format a header string for display.

    Converts snake_case and camelCase to Title Case.
    """
    # Replace underscores with spaces
    formatted = header.replace("_", " ")
    # Add space before capital letters (camelCase)
    import re
    formatted = re.sub(r"([a-z])([A-Z])", r"\1 \2", formatted)
    # Title case
    return formatted.title()


def _format_value(value: Any) -> Any:
    """Format a value for export.

    Handles special types like lists, dicts, None, etc.
    """
    if value is None:
        return ""
    if isinstance(value, (list, tuple)):
        return ", ".join(str(v) for v in value)
    if isinstance(value, dict):
        return str(value)
    return value
