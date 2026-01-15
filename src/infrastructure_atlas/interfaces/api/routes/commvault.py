"""Commvault API routes."""

from __future__ import annotations

import asyncio
import csv
import html
import json
import math
import os
import re
import warnings
from collections import Counter
from collections.abc import Collection, Iterable, Mapping, Sequence
from datetime import UTC, datetime, timedelta
from io import BytesIO, StringIO
from pathlib import Path
from threading import Lock
from typing import Any, Literal
from zoneinfo import ZoneInfo

from fastapi import APIRouter, HTTPException, Query
from fastapi.responses import FileResponse, HTMLResponse, PlainTextResponse, StreamingResponse

from infrastructure_atlas.domain.integrations.commvault import (
    CommvaultJob,
    CommvaultJobList,
    CommvaultPlan,
    CommvaultStoragePool,
)
from infrastructure_atlas.env import project_root
from infrastructure_atlas.infrastructure.external.commvault_client import (
    CommvaultClient,
    CommvaultClientConfig,
    CommvaultError,
)
from infrastructure_atlas.infrastructure.logging import get_logger, logging_context

router = APIRouter(prefix="/commvault", tags=["commvault"])
logger = get_logger(__name__)

# Timezone for date formatting
AMS_TZ = ZoneInfo("Europe/Amsterdam")

# Constants
COMMVAULT_BACKUPS_JSON = "commvault_backups.json"
COMMVAULT_STORAGE_JSON = "commvault_storage.json"
COMMVAULT_PLANS_JSON = "commvault_plans.json"
COMMVAULT_DEFAULT_PLAN_NAME = "Unassigned"
COMMVAULT_PLAN_TYPE_LABELS: dict[int, str] = {
    2: "Server",
}
_commvault_backups_lock = Lock()
_commvault_storage_lock = Lock()
_commvault_plans_lock = Lock()

_ACTIVE_STATUS_KEYWORDS = ("running", "pending", "waiting", "queued", "active", "suspended", "in progress")
_FAILURE_STATUS_KEYWORDS = ("fail", "error", "denied", "invalid", "timeout", "timed out", "kill")
_SUCCESS_STATUS_KEYWORDS = ("complete", "success", "ok", "done")
_WARNING_STATUS_KEYWORDS = ("warning", "partial", "skipped")

COMMVAULT_SERVER_EXPORT_COLUMNS: list[tuple[str, str]] = [
    ("job_id", "Job ID"),
    ("start", "Start"),
    ("status", "Status"),
    ("plan", "Plan"),
    ("app_size", "App Size"),
    ("media_size", "Media Size"),
    ("savings", "Savings"),
]

# Optional urllib3 imports for TLS warning suppression
try:
    from urllib3 import disable_warnings as _disable_urllib3_warnings
    from urllib3.exceptions import InsecureRequestWarning as _InsecureRequestWarning
except Exception:
    _disable_urllib3_warnings = None  # type: ignore[assignment]
    _InsecureRequestWarning = None  # type: ignore[assignment]


# Shared utility functions
def _data_dir() -> Path:
    """Get data directory path."""
    root = project_root()
    raw = os.getenv("NETBOX_DATA_DIR", "data")
    path = Path(raw) if os.path.isabs(raw) else (root / raw)
    path.mkdir(parents=True, exist_ok=True)
    return path


def _env_flag(name: str, default: bool = False) -> bool:
    """Parse boolean environment variable."""
    value = os.getenv(name)
    if value is None:
        return default
    return value.strip().lower() in {"1", "true", "yes", "on"}


# Task logging utilities (duplicated from app.py for module independence)
class _TaskLogger:
    """Mutable container passed into task_logging for success metadata."""

    __slots__ = ("_success_extra",)

    def __init__(self) -> None:
        self._success_extra: dict[str, Any] = {}

    def add_success(self, **kwargs: Any) -> None:
        for key, value in kwargs.items():
            if value is not None:
                self._success_extra[key] = value

    @property
    def success_extra(self) -> dict[str, Any]:
        return self._success_extra


def task_logging(task: str, **context: Any):
    """Log lifecycle events around long-running UI-triggered tasks."""
    import time
    from contextlib import contextmanager

    @contextmanager
    def _task_logging_ctx():
        start = time.perf_counter()
        tracker = _TaskLogger()
        safe_context = {k: v for k, v in context.items() if v not in (None, "")}

        with logging_context(task=task, **safe_context):
            logger.info("Task started", extra={"event": "task_started"})
            try:
                yield tracker
            except Exception:
                duration_ms = int((time.perf_counter() - start) * 1000)
                with logging_context(duration_ms=duration_ms, **tracker.success_extra):
                    logger.exception("Task failed", extra={"event": "task_failed"})
                raise
            else:
                duration_ms = int((time.perf_counter() - start) * 1000)
                with logging_context(duration_ms=duration_ms, **tracker.success_extra):
                    logger.info("Task completed", extra={"event": "task_completed"})

    return _task_logging_ctx()


# Helper functions
def _commvault_client_from_env() -> CommvaultClient:
    """Create Commvault client from environment configuration."""
    base_url = os.getenv("COMMVAULT_BASE_URL", "").strip()
    if not base_url:
        raise RuntimeError("Commvault integration is not configured (COMMVAULT_BASE_URL missing)")

    authtoken = os.getenv("COMMVAULT_API_TOKEN", "").strip() or None
    username = os.getenv("COMMVAULT_EMAIL", "").strip() or None
    password = os.getenv("COMMVAULT_PASSWORD", "")
    if password:
        password = password.strip()
    if not authtoken and not (username and password):
        raise RuntimeError("Commvault credentials are not configured (token or email/password required)")

    verify_tls = _env_flag("COMMVAULT_VERIFY_TLS", True)
    timeout_raw = os.getenv("COMMVAULT_TIMEOUT", "").strip()
    try:
        timeout = float(timeout_raw) if timeout_raw else 30.0
    except ValueError:
        timeout = 30.0

    if not verify_tls:
        _maybe_disable_commvault_tls_warnings()

    config = CommvaultClientConfig(
        base_url=base_url,
        authtoken=authtoken,
        username=username,
        password=password or None,
        verify_tls=verify_tls,
        timeout=timeout,
    )
    return CommvaultClient(config)


def _maybe_disable_commvault_tls_warnings() -> None:
    """Mute urllib3 TLS warnings when verification is intentionally disabled."""
    if _disable_urllib3_warnings and _InsecureRequestWarning:
        _disable_urllib3_warnings(_InsecureRequestWarning)
        return
    warnings.filterwarnings("ignore", category=Warning, module="urllib3")


def _collect_commvault_jobs_for_ui(
    client: CommvaultClient,
    *,
    limit: int,
    offset: int,
    since: datetime | None,
    tracked_job_ids: Collection[int] | None = None,
    latest_start: datetime | None = None,
    cached_recent_count: int | None = None,
    cutoff_start: datetime | None = None,
) -> CommvaultJobList:
    """Collect Commvault jobs for UI display with pagination and filtering."""
    if limit < 0:
        raise ValueError("limit must be zero or positive")

    job_type = None
    page_target = 200 if limit == 0 else max(1, min(500, limit))
    tracked_remaining: set[int] = {int(job_id) for job_id in (tracked_job_ids or [])}
    threshold: datetime | None = since
    if latest_start and (threshold is None or latest_start < threshold):
        threshold = latest_start
    if cutoff_start and (threshold is None or cutoff_start < threshold):
        threshold = cutoff_start

    def _list_jobs_with_retry(*, fetch_limit: int, fetch_offset: int) -> CommvaultJobList:
        attempt_limit = max(1, fetch_limit)
        last_error: CommvaultError | None = None
        while attempt_limit >= 1:
            try:
                return client.list_jobs(
                    limit=attempt_limit,
                    offset=fetch_offset,
                    job_type=job_type,
                )
            except CommvaultError as exc:
                last_error = exc
                if attempt_limit <= 25:
                    break
                attempt_limit = max(25, attempt_limit // 2)
        if last_error is not None:
            raise last_error
        return CommvaultJobList(total_available=0, jobs=())

    try:
        initial = _list_jobs_with_retry(fetch_limit=1, fetch_offset=0)
    except CommvaultError:
        initial = _list_jobs_with_retry(fetch_limit=25, fetch_offset=0)

    total_available = initial.total_available or len(initial.jobs)
    if not total_available:
        return CommvaultJobList(total_available=0, jobs=())

    target_end = max(0, total_available - offset)
    if target_end <= 0:
        return CommvaultJobList(total_available=total_available, jobs=())

    approximate_needed = page_target
    if cached_recent_count is not None:
        approximate_needed = max(approximate_needed, min(total_available, cached_recent_count + page_target))
    if tracked_remaining:
        approximate_needed = max(approximate_needed, min(total_available, len(tracked_remaining) + page_target))
    if limit > 0:
        approximate_needed = max(approximate_needed, min(total_available, limit + page_target))

    target_start = max(0, target_end - approximate_needed)
    jobs: list[CommvaultJob] = []
    seen: set[int] = set()
    current_end = target_end

    while current_end > target_start and (limit == 0 or len(jobs) < limit):
        request_limit = min(page_target, current_end - target_start)
        try:
            response = _list_jobs_with_retry(
                fetch_limit=request_limit,
                fetch_offset=current_end - request_limit,
            )
        except CommvaultError:
            break

        batch = list(response.jobs)
        if not batch:
            break

        for job in reversed(batch):
            try:
                job_id = int(job.job_id)
            except (TypeError, ValueError):
                job_id = None
            if job_id is not None and job_id in seen:
                continue
            jobs.append(job)
            if job_id is not None:
                seen.add(job_id)
                tracked_remaining.discard(job_id)

        current_end -= len(batch)

        if limit > 0 and len(jobs) >= limit:
            break
        if tracked_remaining:
            continue
        if threshold:
            oldest = batch[0].start_time if batch else None
            if oldest and oldest < threshold:
                break
        if len(batch) < request_limit:
            break

    if limit > 0 and len(jobs) > limit:
        jobs = jobs[:limit]
    return CommvaultJobList(total_available=total_available, jobs=tuple(jobs))


def _serialise_commvault_job(job: CommvaultJob) -> dict[str, Any]:
    """Serialize Commvault job to dict."""
    def iso(dt: datetime | None) -> str | None:
        if not dt:
            return None
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=UTC)
        return dt.astimezone(UTC).isoformat()

    return {
        "job_id": job.job_id,
        "job_type": job.job_type,
        "status": job.status,
        "localized_status": job.localized_status,
        "localized_operation": job.localized_operation,
        "client_name": job.client_name,
        "client_id": job.client_id,
        "destination_client_name": job.destination_client_name,
        "subclient_name": job.subclient_name,
        "backup_set_name": job.backup_set_name,
        "application_name": job.application_name,
        "backup_level_name": job.backup_level_name,
        "plan_name": job.plan_name,
        "client_groups": list(job.client_groups),
        "storage_policy_name": job.storage_policy_name,
        "start_time": iso(job.start_time),
        "end_time": iso(job.end_time),
        "elapsed_seconds": job.elapsed_seconds,
        "size_of_application_bytes": job.size_of_application_bytes,
        "size_on_media_bytes": job.size_on_media_bytes,
        "total_num_files": job.total_num_files,
        "percent_complete": job.percent_complete,
        "percent_savings": job.percent_savings,
        "average_throughput_gb_per_hr": job.average_throughput,
        "retain_until": iso(job.retain_until),
    }


def _serialise_commvault_plan(plan: CommvaultPlan) -> dict[str, Any]:
    """Serialize Commvault plan to dict."""
    return {
        "plan_id": plan.plan_id,
        "name": plan.name,
        "plan_type": plan.plan_type,
        "associated_entities": plan.associated_entities,
        "rpo": plan.rpo,
        "copy_count": plan.copy_count,
        "status": plan.status,
        "tags": list(plan.tags),
        "raw": plan.raw,
    }


def _write_commvault_backups_json(payload: Mapping[str, Any]) -> None:
    """Write Commvault backups JSON to disk."""
    path = _data_dir() / COMMVAULT_BACKUPS_JSON
    tmp = path.with_suffix(path.suffix + ".tmp")
    with tmp.open("w", encoding="utf-8") as handle:
        json.dump(payload, handle, indent=2)
    tmp.replace(path)


def _write_commvault_storage_json(payload: Mapping[str, Any]) -> None:
    """Write Commvault storage JSON to disk."""
    path = _data_dir() / COMMVAULT_STORAGE_JSON
    tmp = path.with_suffix(path.suffix + ".tmp")
    with tmp.open("w", encoding="utf-8") as handle:
        json.dump(payload, handle, indent=2)
    tmp.replace(path)


def _write_commvault_plans_json(payload: Mapping[str, Any]) -> None:
    """Write Commvault plans JSON to disk."""
    path = _data_dir() / COMMVAULT_PLANS_JSON
    tmp = path.with_suffix(path.suffix + ".tmp")
    with tmp.open("w", encoding="utf-8") as handle:
        json.dump(payload, handle, indent=2)
    tmp.replace(path)


def _load_commvault_backups() -> dict[str, Any]:
    """Load cached Commvault backups from disk."""
    path = _data_dir() / COMMVAULT_BACKUPS_JSON
    if not path.exists():
        return {
            "jobs": [],
            "generated_at": None,
            "total_cached": 0,
            "version": 2,
        }
    try:
        with path.open("r", encoding="utf-8") as handle:
            data = json.load(handle)
            if isinstance(data, Mapping):
                jobs = list(data.get("jobs") or [])
                generated = data.get("generated_at")
                total_cached = data.get("total_cached")
                try:
                    total_cached = int(total_cached)
                except (TypeError, ValueError):
                    total_cached = len(jobs)
                return {
                    "jobs": jobs,
                    "generated_at": generated,
                    "total_cached": total_cached,
                    "version": data.get("version", 2),
                }
    except Exception:
        logger.exception("Failed to load cached Commvault backups", extra={"event": "commvault_cache_error"})
    return {
        "jobs": [],
        "generated_at": None,
        "total_cached": 0,
        "version": 2,
    }


def _load_commvault_plans() -> dict[str, Any]:
    """Load cached Commvault plans from disk."""
    path = _data_dir() / COMMVAULT_PLANS_JSON
    if not path.exists():
        return {
            "plans": [],
            "generated_at": None,
            "total_cached": 0,
            "plan_types": [],
            "version": 1,
        }
    try:
        with path.open("r", encoding="utf-8") as handle:
            data = json.load(handle)
            if isinstance(data, Mapping):
                plans = list(data.get("plans") or [])
                plan_types = list(data.get("plan_types") or [])
                return {
                    "plans": plans,
                    "generated_at": data.get("generated_at"),
                    "total_cached": data.get("total_cached", len(plans)),
                    "plan_types": plan_types,
                    "version": data.get("version", 1),
                }
    except Exception:
        logger.exception("Failed to load cached Commvault plans", extra={"event": "commvault_plan_cache_error"})
    return {
        "plans": [],
        "generated_at": None,
        "total_cached": 0,
        "plan_types": [],
        "version": 1,
    }


def _parse_job_datetime(value: Any) -> datetime | None:
    """Parse datetime from job data."""
    if isinstance(value, datetime):
        return value if value.tzinfo else value.replace(tzinfo=UTC)
    if not value:
        return None
    try:
        parsed = datetime.fromisoformat(str(value))
    except ValueError:
        return None
    if parsed.tzinfo is None:
        parsed = parsed.replace(tzinfo=UTC)
    return parsed


def _filter_commvault_jobs(jobs: Iterable[Mapping[str, Any]], since_hours: int) -> list[dict[str, Any]]:
    """Filter jobs by time window."""
    if since_hours <= 0:
        return [dict(job) for job in jobs]
    cutoff = datetime.now(tz=UTC) - timedelta(hours=since_hours)
    filtered: list[dict[str, Any]] = []
    for job in jobs:
        start = _parse_job_datetime(job.get("start_time"))
        if start and start >= cutoff:
            filtered.append(dict(job))
    return filtered


def _commvault_job_digest(job: Mapping[str, Any]) -> str:
    """Create digest/hash of job for change detection."""
    try:
        return json.dumps(job, sort_keys=True, default=str)
    except Exception:
        return repr(sorted(job.items()))


def _commvault_job_is_active(job: Mapping[str, Any]) -> bool:
    """Check if job is still active/running."""
    status_raw = job.get("localized_status") or job.get("status") or ""
    status = str(status_raw).strip().lower()
    if not status:
        return True
    for keyword in _ACTIVE_STATUS_KEYWORDS:
        if keyword in status:
            return True
    end_time = _parse_job_datetime(job.get("end_time"))
    return end_time is None


def _normalise_commvault_plan_name(value: Any) -> str:
    """Normalize plan name."""
    if isinstance(value, str):
        text = value.strip()
        if text:
            return text
    if value is None:
        return COMMVAULT_DEFAULT_PLAN_NAME
    text = str(value).strip()
    return text or COMMVAULT_DEFAULT_PLAN_NAME


def _normalise_commvault_status(value: Any) -> str:
    """Normalize status value."""
    if isinstance(value, str):
        text = value.strip()
        if text:
            return text
    if value is None:
        return "Unknown"
    text = str(value).strip()
    return text or "Unknown"


def _commvault_status_failed(status: str) -> bool:
    """Check if status indicates failure."""
    lowered = status.casefold()
    return any(keyword in lowered for keyword in _FAILURE_STATUS_KEYWORDS)


def _commvault_status_successful(status: str) -> bool:
    """Check if status indicates success."""
    if _commvault_status_failed(status):
        return False
    lowered = status.casefold()
    return any(keyword in lowered for keyword in _SUCCESS_STATUS_KEYWORDS)


def _commvault_status_warning(status: str) -> bool:
    """Check if status indicates warning."""
    if _commvault_status_failed(status) or _commvault_status_successful(status):
        return False
    lowered = status.casefold()
    return any(keyword in lowered for keyword in _WARNING_STATUS_KEYWORDS)


def _summarise_commvault_plans(jobs: Iterable[Mapping[str, Any]]) -> dict[str, Any]:
    """Summarize jobs by plan."""
    plan_records: dict[str, Any] = {}
    total_jobs = 0
    active_jobs = 0
    global_status = Counter()
    global_success = 0
    global_failure = 0
    unique_clients: set[str] = set()
    now = datetime.now(tz=UTC)

    for payload in jobs:
        if not isinstance(payload, Mapping):
            continue

        total_jobs += 1
        plan_display = _normalise_commvault_plan_name(payload.get("plan_name"))
        plan_key = plan_display.casefold()
        bucket = plan_records.setdefault(
            plan_key,
            {
                "plan_name": plan_display,
                "job_count": 0,
                "clients": Counter(),
                "applications": Counter(),
                "subclients": Counter(),
                "status_counts": Counter(),
                "success_count": 0,
                "failed_count": 0,
                "warning_count": 0,
                "active_jobs": 0,
                "retained_jobs": 0,
                "total_application_bytes": 0,
                "total_media_bytes": 0,
                "savings_total": 0.0,
                "savings_samples": 0,
                "latest_start": None,
                "latest_job": None,
            },
        )
        if bucket["plan_name"] == COMMVAULT_DEFAULT_PLAN_NAME and plan_display != COMMVAULT_DEFAULT_PLAN_NAME:
            bucket["plan_name"] = plan_display

        bucket["job_count"] += 1

        status_label = _normalise_commvault_status(payload.get("localized_status") or payload.get("status"))
        bucket["status_counts"][status_label] += 1
        global_status[status_label] += 1
        if _commvault_status_failed(status_label):
            bucket["failed_count"] += 1
            global_failure += 1
        elif _commvault_status_successful(status_label):
            bucket["success_count"] += 1
            global_success += 1
        elif _commvault_status_warning(status_label):
            bucket["warning_count"] += 1

        if _commvault_job_is_active(payload):
            bucket["active_jobs"] += 1
            active_jobs += 1

        client_label = payload.get("client_name") or payload.get("destination_client_name")
        if client_label:
            client_name = str(client_label).strip()
            if client_name:
                bucket["clients"][client_name] += 1
                unique_clients.add(client_name)

        application = payload.get("application_name")
        if application:
            app_name = str(application).strip()
            if app_name:
                bucket["applications"][app_name] += 1

        subclient_label = payload.get("subclient_name") or payload.get("backup_set_name")
        if subclient_label:
            subclient_name = str(subclient_label).strip()
            if subclient_name:
                bucket["subclients"][subclient_name] += 1

        retain_until = _parse_job_datetime(payload.get("retain_until"))
        if retain_until and retain_until > now:
            bucket["retained_jobs"] += 1

        try:
            bucket["total_application_bytes"] += int(payload.get("size_of_application_bytes") or 0)
        except (TypeError, ValueError):
            pass
        try:
            bucket["total_media_bytes"] += int(payload.get("size_on_media_bytes") or 0)
        except (TypeError, ValueError):
            pass

        try:
            savings_value = float(payload.get("percent_savings"))
        except (TypeError, ValueError):
            savings_value = None
        if savings_value is not None:
            bucket["savings_total"] += savings_value
            bucket["savings_samples"] += 1

        start_dt = _parse_job_datetime(payload.get("start_time"))
        if start_dt and (bucket["latest_start"] is None or start_dt > bucket["latest_start"]):
            bucket["latest_start"] = start_dt
            bucket["latest_job"] = {
                "job_id": payload.get("job_id"),
                "job_type": payload.get("job_type"),
                "status": payload.get("status"),
                "localized_status": payload.get("localized_status"),
                "client_name": payload.get("client_name"),
                "destination_client_name": payload.get("destination_client_name"),
                "subclient_name": payload.get("subclient_name"),
                "application_name": payload.get("application_name"),
                "backup_level_name": payload.get("backup_level_name"),
                "plan_name": plan_display,
                "start_time": payload.get("start_time"),
                "end_time": payload.get("end_time"),
                "size_of_application_bytes": payload.get("size_of_application_bytes"),
                "size_on_media_bytes": payload.get("size_on_media_bytes"),
                "percent_savings": payload.get("percent_savings"),
            }

    plans: list[dict[str, Any]] = []
    for bucket in plan_records.values():
        clients_counter: Counter[str] = bucket["clients"]
        applications_counter: Counter[str] = bucket["applications"]
        subclients_counter: Counter[str] = bucket["subclients"]
        status_counter: Counter[str] = bucket["status_counts"]
        savings_count = bucket["savings_samples"] or 0
        average_savings = bucket["savings_total"] / savings_count if savings_count else None
        job_count = bucket["job_count"] or 0

        plans.append(
            {
                "plan_name": bucket["plan_name"],
                "job_count": job_count,
                "client_count": len(clients_counter),
                "clients": [
                    {"name": name, "count": count}
                    for name, count in sorted(
                        clients_counter.items(), key=lambda item: (-item[1], item[0].casefold())
                    )
                ],
                "applications": [
                    {"name": name, "count": count}
                    for name, count in sorted(
                        applications_counter.items(), key=lambda item: (-item[1], item[0].casefold())
                    )
                ],
                "subclients": [
                    {"name": name, "count": count}
                    for name, count in sorted(
                        subclients_counter.items(), key=lambda item: (-item[1], item[0].casefold())
                    )
                ],
                "status_counts": [
                    {"status": name, "count": count}
                    for name, count in sorted(
                        status_counter.items(), key=lambda item: (-item[1], item[0].casefold())
                    )
                ],
                "success_count": bucket["success_count"],
                "failed_count": bucket["failed_count"],
                "warning_count": bucket["warning_count"],
                "active_jobs": bucket["active_jobs"],
                "retained_jobs": bucket["retained_jobs"],
                "total_application_bytes": bucket["total_application_bytes"],
                "total_media_bytes": bucket["total_media_bytes"],
                "average_savings_percent": average_savings,
                "latest_job": bucket["latest_job"],
                "latest_job_start": bucket["latest_start"].isoformat() if bucket["latest_start"] else None,
                "success_rate": (bucket["success_count"] / job_count * 100.0) if job_count else None,
                "failure_rate": (bucket["failed_count"] / job_count * 100.0) if job_count else None,
            }
        )

    plans.sort(key=lambda plan: (-plan["job_count"], plan["plan_name"].casefold()))

    total_success = sum(plan["success_count"] for plan in plans)
    total_failure = sum(plan["failed_count"] for plan in plans)

    return {
        "plans": plans,
        "total_plans": len(plans),
        "total_jobs": total_jobs,
        "active_jobs": active_jobs,
        "unique_clients": len(unique_clients),
        "status_counts": [
            {"status": name, "count": count}
            for name, count in sorted(global_status.items(), key=lambda item: (-item[1], item[0].casefold()))
        ],
        "success_rate": (total_success / total_jobs * 100.0) if total_jobs else None,
        "failure_rate": (total_failure / total_jobs * 100.0) if total_jobs else None,
    }


def _latest_commvault_start_time(jobs: Iterable[Mapping[str, Any]]) -> datetime | None:
    """Get latest job start time."""
    latest: datetime | None = None
    for job in jobs:
        start = _parse_job_datetime(job.get("start_time"))
        if start and (latest is None or start > latest):
            latest = start
    return latest


def _safe_commvault_client_id(value: Any) -> int | None:
    """Safely parse client ID."""
    try:
        client_id = int(value)
    except (TypeError, ValueError):
        return None
    return client_id if client_id >= 0 else None


def _normalise_commvault_client_name(value: Any) -> str:
    """Normalize client name."""
    if value is None:
        return ""
    return str(value).strip()


def _cached_commvault_clients() -> tuple[list[dict[str, Any]], list[Mapping[str, Any]], str | None]:
    """Get cached Commvault clients from backups data."""
    cache = _load_commvault_backups()
    jobs_payload = cache.get("jobs")
    if not isinstance(jobs_payload, list):
        jobs_payload = []

    clients: dict[tuple[int | None, str], dict[str, Any]] = {}
    for raw in jobs_payload:
        if not isinstance(raw, Mapping):
            continue
        client_id = _safe_commvault_client_id(raw.get("client_id"))
        client_name = _normalise_commvault_client_name(raw.get("client_name"))
        dest_name = _normalise_commvault_client_name(raw.get("destination_client_name"))
        primary = client_name or dest_name
        display = dest_name or client_name or primary
        key_name = (primary or display or "").casefold()
        key = (client_id, key_name)
        entry = clients.get(key)
        if not entry:
            label = display or primary or (f"Client {client_id}" if client_id is not None else "Unnamed client")
            entry = {
                "client_id": client_id,
                "name": primary or label,
                "display_name": label,
                "name_variants": set(),
                "job_count": 0,
            }
            clients[key] = entry
        entry["job_count"] += 1
        for candidate in (client_name, dest_name, primary, display):
            if candidate:
                entry["name_variants"].add(candidate.casefold())

    client_list = list(clients.values())
    for entry in client_list:
        entry["name_variants"].add((entry["name"] or "").casefold())
        entry["name_variants"].add((entry["display_name"] or "").casefold())
    client_list.sort(key=lambda item: (-item["job_count"], (item["display_name"] or item["name"] or "").lower()))

    jobs = [job for job in jobs_payload if isinstance(job, Mapping)]
    generated_at = cache.get("generated_at")
    return client_list, jobs, generated_at


def _match_cached_commvault_client(identifier: str, clients: list[dict[str, Any]]) -> dict[str, Any]:
    """Match client by identifier (ID or name)."""
    ident = (identifier or "").strip()
    if not ident:
        raise HTTPException(status_code=400, detail="Client identifier is required")

    if ident.isdigit():
        client_id = int(ident)
        for entry in clients:
            if entry["client_id"] == client_id:
                return entry
        raise HTTPException(status_code=404, detail=f"No cached client with ID {client_id} found")

    needle = ident.casefold()

    def _merge_clients(candidates: list[dict[str, Any]]) -> dict[str, Any]:
        """Merge multiple client records into one virtual client."""
        if not candidates:
            raise ValueError("No candidates to merge")
        first = candidates[0]
        merged = dict(first)
        # Use the union of all name variants to ensure we match jobs for any of the clients
        all_variants = set()
        for c in candidates:
            all_variants.update(c.get("name_variants", set()) or set())
        merged["name_variants"] = all_variants
        return merged

    exact = [entry for entry in clients if needle in entry["name_variants"]]
    if len(exact) >= 1:
        return _merge_clients(exact)

    matches = [
        entry
        for entry in clients
        if any(needle in variant for variant in entry["name_variants"] if variant)
    ]
    if not matches:
        raise HTTPException(status_code=404, detail=f"No cached client matching '{identifier}' found")
    
    # Merge fuzzy matches as well
    return _merge_clients(matches)


def _job_matches_cached_client(job: Mapping[str, Any], client_record: dict[str, Any]) -> bool:
    """Check if job belongs to client."""
    if client_record.get("client_id") is not None:
        job_id = _safe_commvault_client_id(job.get("client_id"))
        if job_id is not None and job_id == client_record["client_id"]:
            return True
    names = {
        _normalise_commvault_client_name(job.get("client_name")).casefold(),
        _normalise_commvault_client_name(job.get("destination_client_name")).casefold(),
    }
    names.discard("")
    if not names:
        return False
    variants = client_record.get("name_variants") or set()
    return any(name in variants for name in names)


def _build_cached_client_summary(record: dict[str, Any]) -> dict[str, Any]:
    """Build client summary from cached record."""
    return {
        "client_id": record.get("client_id"),
        "name": record.get("name"),
        "display_name": record.get("display_name"),
        "host_name": None,
        "os_name": None,
        "os_type": None,
        "os_subtype": None,
        "processor_type": None,
        "cpu_count": None,
        "is_media_agent": None,
        "is_virtual": None,
        "is_infrastructure": None,
        "is_commserve": None,
        "readiness_status": None,
        "last_ready_time": None,
        "sla_status_code": None,
        "sla_description": None,
        "agent_applications": [],
        "client_groups": [],
    }


def _build_cached_job_metrics(
    jobs: list[dict[str, Any]],
    *,
    since_hours: int,
    retained_only: bool,
    cache_generated_at: str | None,
) -> dict[str, Any]:
    """Build job metrics from cached data."""
    def _to_int(value: Any) -> int:
        try:
            return int(value)
        except (TypeError, ValueError):
            return 0

    total_app = sum(max(0, _to_int(job.get("size_of_application_bytes"))) for job in jobs)
    total_media = sum(max(0, _to_int(job.get("size_on_media_bytes"))) for job in jobs)
    last_start: datetime | None = None
    for job in jobs:
        start_dt = _parse_job_datetime(job.get("start_time"))
        if start_dt and (last_start is None or start_dt > last_start):
            last_start = start_dt

    fetched_at = _parse_job_datetime(cache_generated_at) if cache_generated_at else None
    if fetched_at is None:
        fetched_at = datetime.now(tz=UTC)

    retain_cutoff = fetched_at if retained_only else None

    return {
        "window_hours": since_hours,
        "job_count": len(jobs),
        "total_application_bytes": total_app,
        "total_media_bytes": total_media,
        "last_job_start": last_start.isoformat() if last_start else None,
        "within_window": bool(jobs),
        "descending": True,
        "retain_cutoff": retain_cutoff.isoformat() if retain_cutoff else None,
        "retain_required": retained_only,
        "fetched_at": fetched_at.isoformat(),
    }


def _slugify_commvault_name(value: str, default: str = "commvault") -> str:
    """Slugify name for filename."""
    text = (value or "").strip().lower()
    if not text:
        return default
    safe = re.sub(r"[^a-z0-9]+", "-", text)
    safe = re.sub(r"-+", "-", safe).strip("-")
    return safe or default


def _format_bytes_for_report(value: int | float | None) -> str:
    """Format bytes for human-readable report."""
    if not value:
        return "0 B"
    size = float(value)
    if size <= 0:
        return "0 B"
    units = ["B", "KB", "MB", "GB", "TB", "PB"]
    idx = 0
    while size >= 1024 and idx < len(units) - 1:
        size /= 1024
        idx += 1
    precision = 0 if size >= 100 else 1 if size >= 10 else 2
    return f"{size:.{precision}f} {units[idx]}"


def _format_minutes_label(value: int | None) -> str | None:
    """Format minutes as human-readable label."""
    if value is None:
        return None
    minutes = int(value)
    if minutes < 0:
        minutes = abs(minutes)
    if minutes == 0:
        return "0 minutes"
    if minutes % 1440 == 0:
        days = minutes // 1440
        return f"{days} day" if days == 1 else f"{days} days"
    if minutes % 60 == 0:
        hours = minutes // 60
        return f"{hours} hour" if hours == 1 else f"{hours} hours"
    return f"{minutes} minute" if minutes == 1 else f"{minutes} minutes"


def _plan_status_from_flag(flag: int | None) -> str | None:
    """Convert plan status flag to label."""
    if flag is None:
        return None
    if flag == 0:
        return "Enabled"
    if flag == 1:
        return "Disabled"
    if flag == 2:
        return "Retired"
    return f"Flag {flag}"


def _label_commvault_plan_type(code: int | None, fallback: str | None = None) -> str | None:
    """Label Commvault plan type."""
    if code is None:
        return fallback
    label = COMMVAULT_PLAN_TYPE_LABELS.get(code)
    if label:
        return label
    if fallback:
        return fallback
    return str(code)


def _compute_commvault_server_metrics(jobs: list[dict[str, Any]]) -> dict[str, Any]:
    """Compute server metrics from jobs."""
    if not jobs:
        return {
            "job_count": 0,
            "plan_breakdown": [],
            "subclient_breakdown": [],
            "policy_breakdown": [],
            "total_application_bytes": 0,
            "total_media_bytes": 0,
            "retained_jobs": 0,
            "average_savings_percent": None,
            "average_reduction_ratio": None,
            "average_reduction_ratio_text": None,
            "savings_bytes": 0,
            "series": {"timeline": []},
            "latest_job": None,
            "latest_job_started_at": None,
            "next_retain_expiry": None,
        }

    plan_counts: Counter[str] = Counter()
    subclient_counts: Counter[str] = Counter()
    policy_counts: Counter[str] = Counter()
    total_application = 0
    total_media = 0
    savings_samples: list[float] = []
    retained_jobs = 0
    next_expiry: datetime | None = None
    latest_start: datetime | None = None
    latest_job: dict[str, Any] | None = None
    timeline: list[dict[str, Any]] = []

    def _as_int(value: Any) -> int:
        try:
            return int(value)
        except (TypeError, ValueError):
            return 0

    def _as_float(value: Any) -> float | None:
        try:
            return float(value)
        except (TypeError, ValueError):
            return None

    for job in jobs:
        plan = job.get("plan_name") or "Unassigned"
        subclient = job.get("subclient_name") or "Unspecified"
        policy = job.get("storage_policy_name") or "Unspecified"
        plan_counts[plan] += 1
        subclient_counts[subclient] += 1
        policy_counts[policy] += 1

        app_size = max(0, _as_int(job.get("size_of_application_bytes")))
        media_size = max(0, _as_int(job.get("size_on_media_bytes")))
        total_application += app_size
        total_media += media_size

        savings = _as_float(job.get("percent_savings"))
        if savings is not None:
            savings_samples.append(savings)

        retain_until = _parse_job_datetime(job.get("retain_until"))
        if retain_until:
            retained_jobs += 1
            if next_expiry is None or retain_until < next_expiry:
                next_expiry = retain_until

        start_dt = _parse_job_datetime(job.get("start_time"))
        if start_dt and (latest_start is None or start_dt > latest_start):
            latest_start = start_dt
            latest_job = job

        timeline.append(
            {
                "start_time": job.get("start_time"),
                "start_timestamp": start_dt.timestamp() if start_dt else None,
                "size_of_application_bytes": app_size,
                "size_on_media_bytes": media_size,
                "percent_savings": savings,
                "plan_name": job.get("plan_name"),
                "status": job.get("localized_status") or job.get("status"),
            }
        )

    timeline.sort(key=lambda item: item.get("start_timestamp") or 0)
    for entry in timeline:
        entry.pop("start_timestamp", None)

    average_savings = None
    if savings_samples:
        average_savings = sum(savings_samples) / len(savings_samples)

    reduction_ratio = None
    reduction_ratio_text = None
    if total_media > 0:
        reduction_ratio = total_application / total_media if total_application else 0
        if reduction_ratio and reduction_ratio > 0:
            reduction_ratio_text = f"{reduction_ratio:.1f}:1"

    savings_bytes = max(0, total_application - total_media)

    def _counts_to_rows(counter: Counter[str]) -> list[dict[str, Any]]:
        return [
            {"label": label, "restore_points": count}
            for label, count in counter.most_common()
        ]

    return {
        "job_count": len(jobs),
        "plan_breakdown": _counts_to_rows(plan_counts),
        "subclient_breakdown": _counts_to_rows(subclient_counts),
        "policy_breakdown": _counts_to_rows(policy_counts),
        "total_application_bytes": total_application,
        "total_media_bytes": total_media,
        "retained_jobs": retained_jobs,
        "average_savings_percent": average_savings,
        "average_reduction_ratio": reduction_ratio,
        "average_reduction_ratio_text": reduction_ratio_text,
        "savings_bytes": savings_bytes,
        "series": {"timeline": timeline},
        "latest_job": latest_job,
        "latest_job_started_at": latest_start.isoformat() if latest_start else None,
        "next_retain_expiry": next_expiry.isoformat() if next_expiry else None,
    }


def _load_commvault_server_data(
    client_identifier: str,
    *,
    job_limit: int,
    since_hours: int,
    retained_only: bool,
    refresh_cache: bool,
) -> tuple[dict[str, Any], dict[str, Any], list[dict[str, Any]], dict[str, Any]]:
    """Load Commvault server data from cache."""
    clients, jobs_payload, cache_generated_at = _cached_commvault_clients()
    if not clients:
        raise HTTPException(status_code=503, detail="Commvault cached backups are not available. Run an export first.")

    client_record = _match_cached_commvault_client(client_identifier, clients)

    cutoff = datetime.now(tz=UTC) - timedelta(hours=since_hours) if since_hours > 0 else None
    now_utc = datetime.now(tz=UTC)
    default_order = datetime.min.replace(tzinfo=UTC)

    selected: list[tuple[datetime | None, dict[str, Any]]] = []
    for raw in jobs_payload:
        if not _job_matches_cached_client(raw, client_record):
            continue

        start_dt = _parse_job_datetime(raw.get("start_time"))
        if cutoff and (start_dt is None or start_dt < cutoff):
            continue

        if retained_only:
            retain_dt = _parse_job_datetime(raw.get("retain_until"))
            if retain_dt is None or retain_dt <= now_utc:
                continue

        selected.append((start_dt, dict(raw)))

    selected.sort(key=lambda item: item[0] or default_order, reverse=True)
    jobs = [job for _, job in selected]
    if job_limit > 0:
        jobs = jobs[:job_limit]

    stats = _compute_commvault_server_metrics(jobs)
    metrics_payload = _build_cached_job_metrics(
        jobs,
        since_hours=max(0, since_hours),
        retained_only=retained_only,
        cache_generated_at=cache_generated_at,
    )
    summary_payload = _build_cached_client_summary(client_record)
    return summary_payload, metrics_payload, jobs, stats


@router.get("/server-info")
def commvault_server_info(
    hostname: str = Query(..., description="Client hostname or ID"),
    limit: int = Query(25, ge=1, le=500),
    hours: int = Query(24, ge=1, le=720),
    retained_only: int = Query(0, ge=0, le=1),
):
    """Get summarized Commvault job info for a specific client/hostname."""
    try:
        summary, metrics, jobs, stats = _load_commvault_server_data(
            hostname,
            job_limit=limit,
            since_hours=hours,
            retained_only=bool(retained_only),
            refresh_cache=False,
        )
        return {
            "summary": summary,
            "metrics": metrics,
            "stats": stats,
            "jobs": [
                {
                    "job_id": j.get("job_id"),
                    "status": j.get("status"),
                    "start_time": j.get("start_time"),
                    "type": j.get("job_type"),
                    "plan": j.get("plan_name"),
                    "app_size": j.get("size_of_application_bytes"),
                }
                for j in jobs
            ],
        }
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc))


def _format_iso_human(value: str | None) -> str:
    """Format ISO datetime as human-readable."""
    if not value:
        return "-"
    dt = _parse_job_datetime(value)
    if not dt:
        return "-"
    local_dt = dt.astimezone(AMS_TZ)
    tz_label = local_dt.tzname() or "CET"
    return f"{local_dt.strftime('%Y-%m-%d %H:%M')} {tz_label}"


def _format_percent(value: float | None) -> str:
    """Format percentage value."""
    if value is None:
        return "-"
    return f"{value:.1f}%"


def _commvault_job_table_rows(jobs: list[dict[str, Any]]) -> list[dict[str, str]]:
    """Format job rows for table display."""
    rows: list[dict[str, str]] = []
    for job in jobs:
        job_id = str(job.get("job_id") or "-")
        start_label = _format_iso_human(job.get("start_time"))
        status_label = job.get("localized_status") or job.get("status") or "-"
        plan_label = job.get("plan_name") or "-"
        app_size = _format_bytes_for_report(job.get("size_of_application_bytes"))
        media_size = _format_bytes_for_report(job.get("size_on_media_bytes"))
        savings_value = job.get("percent_savings")
        try:
            savings_float = float(savings_value) if savings_value is not None else None
        except (TypeError, ValueError):
            savings_float = None
        savings_label = _format_percent(savings_float)
        rows.append(
            {
                "job_id": job_id,
                "start": start_label,
                "status": status_label,
                "plan": plan_label,
                "app_size": app_size,
                "media_size": media_size,
                "savings": savings_label,
            }
        )
    return rows


def _render_commvault_server_text_report(
    summary: dict[str, Any],
    stats: dict[str, Any],
    jobs: list[dict[str, Any]],
    metrics_payload: dict[str, Any] | None,
) -> str:
    """Render Commvault server report as plain text."""
    display_name = summary.get("display_name") or summary.get("name") or f"#{summary.get('client_id')}"
    generated = metrics_payload.get("fetched_at") if metrics_payload else None
    generated_label = _format_iso_human(generated)
    window_hours = metrics_payload.get("window_hours") if metrics_payload else None
    job_rows = _commvault_job_table_rows(jobs)

    lines = [
        f"Commvault server report for {display_name} (ID {summary.get('client_id')})",
        f"Generated: {generated_label}",
    ]
    if window_hours is not None:
        lines.append(f"Window: {window_hours}h | Restore points: {stats.get('job_count', 0)}")

    lines.append(
        f"Application data: {_format_bytes_for_report(stats.get('total_application_bytes'))}"
        f" | Media: {_format_bytes_for_report(stats.get('total_media_bytes'))}"
        f" | Savings: {_format_bytes_for_report(stats.get('savings_bytes'))}"
    )

    reduction_text = stats.get("average_reduction_ratio_text") or "-"
    savings_percent = _format_percent(stats.get("average_savings_percent"))
    lines.append(f"Average data reduction: {savings_percent} (≈ {reduction_text})")

    plan_rows = stats.get("plan_breakdown") or []
    if plan_rows:
        lines.append("Plan restore points:")
        for row in plan_rows:
            label = row.get("label") or row.get("plan") or "Unnamed"
            count = row.get("restore_points") or 0
            lines.append(f"  - {label}: {count}")

    subclient_rows = stats.get("subclient_breakdown") or []
    if subclient_rows:
        lines.append("Subclients:")
        for row in subclient_rows:
            label = row.get("label") or "Unnamed"
            count = row.get("restore_points") or 0
            lines.append(f"  - {label}: {count}")

    lines.append("")
    header = "  ".join(
        [
            f"{title:<18}" if title not in {"Job ID", "Plan"} else f"{title:<10}"
            for _, title in COMMVAULT_SERVER_EXPORT_COLUMNS
        ]
    )
    lines.append(header)
    lines.append("-" * len(header))
    for row in job_rows:
        lines.append(
            "  ".join(
                [
                    f"{row['job_id']:<10}",
                    f"{row['start']:<20}",
                    f"{row['status']:<18}",
                    f"{row['plan']:<18}",
                    f"{row['app_size']:>12}",
                    f"{row['media_size']:>12}",
                    f"{row['savings']:>8}",
                ]
            )
        )

    if not job_rows:
        lines.append("No jobs found for the selected parameters.")

    return "\n".join(lines)


def _render_commvault_server_markdown_report(
    summary: dict[str, Any],
    stats: dict[str, Any],
    jobs: list[dict[str, Any]],
    metrics_payload: dict[str, Any] | None,
) -> str:
    """Render Commvault server report as Markdown."""
    display_name = summary.get("display_name") or summary.get("name") or f"#{summary.get('client_id')}"
    generated = metrics_payload.get("fetched_at") if metrics_payload else None
    generated_label = _format_iso_human(generated)
    window_hours = metrics_payload.get("window_hours") if metrics_payload else None
    job_rows = _commvault_job_table_rows(jobs)

    lines = [f"# Commvault server report — {display_name}", ""]
    lines.append(f"*ID*: `{summary.get('client_id')}`")
    lines.append(f"*Generated*: {generated_label}")
    if window_hours is not None:
        lines.append(f"*Window*: {window_hours}h")
    lines.append(
        f"*Application data*: {_format_bytes_for_report(stats.get('total_application_bytes'))}"
        f" — *Media*: {_format_bytes_for_report(stats.get('total_media_bytes'))}"
        f" — *Savings*: {_format_bytes_for_report(stats.get('savings_bytes'))}"
    )
    reduction_text = stats.get("average_reduction_ratio_text") or "-"
    savings_percent = _format_percent(stats.get("average_savings_percent"))
    lines.append(f"*Average data reduction*: {savings_percent} (≈ {reduction_text})")
    lines.append("")

    plan_rows = stats.get("plan_breakdown") or []
    if plan_rows:
        lines.append("## Restore points by plan")
        for row in plan_rows:
            label = row.get("label") or row.get("plan") or "Unnamed"
            count = row.get("restore_points") or 0
            lines.append(f"- **{label}** — {count} restore point(s)")
        lines.append("")

    lines.append("## Jobs")
    header = " | ".join(title for _, title in COMMVAULT_SERVER_EXPORT_COLUMNS)
    lines.append(f"{header}")
    lines.append(" | ".join(["---"] * len(COMMVAULT_SERVER_EXPORT_COLUMNS)))
    for row in job_rows:
        lines.append(
            " | ".join(
                [
                    row["job_id"],
                    row["start"],
                    row["status"],
                    row["plan"],
                    row["app_size"],
                    row["media_size"],
                    row["savings"],
                ]
            )
        )
    if not job_rows:
        lines.append("No jobs found for the selected parameters.")

    return "\n".join(lines)


def _render_commvault_server_html_report(
    summary: dict[str, Any],
    stats: dict[str, Any],
    jobs: list[dict[str, Any]],
    metrics_payload: dict[str, Any] | None,
) -> str:
    """Render Commvault server report as HTML."""
    display_name = summary.get("display_name") or summary.get("name") or f"#{summary.get('client_id')}"
    generated = metrics_payload.get("fetched_at") if metrics_payload else None
    generated_label = _format_iso_human(generated)
    window_hours = metrics_payload.get("window_hours") if metrics_payload else None
    job_rows = _commvault_job_table_rows(jobs)

    plan_rows = stats.get("plan_breakdown") or []
    plan_html = "".join(
        f"<li><strong>{html.escape(str(row.get('label') or row.get('plan') or 'Unnamed'))}</strong>: {row.get('restore_points') or 0}</li>"
        for row in plan_rows
    )

    job_rows_html = "".join(
        "<tr>" + "".join(
            f"<td>{html.escape(row[key])}</td>"
            for key, _ in COMMVAULT_SERVER_EXPORT_COLUMNS
        ) + "</tr>"
        for row in job_rows
    )

    if not job_rows_html:
        job_rows_html = "<tr><td colspan=\"7\">No jobs found for the selected parameters.</td></tr>"

    headings = "".join(f"<th>{html.escape(title)}</th>" for _, title in COMMVAULT_SERVER_EXPORT_COLUMNS)

    return f"""
<!doctype html>
<html>
  <head>
    <meta charset="utf-8" />
    <title>Commvault report — {html.escape(display_name)}</title>
    <style>
      body {{ font-family: system-ui, sans-serif; margin: 24px; color: #0f172a; }}
      h1 {{ margin-top: 0; }}
      table {{ border-collapse: collapse; width: 100%; margin-top: 16px; }}
      th, td {{ border: 1px solid #cbd5f5; padding: 8px 10px; text-align: left; font-size: 13px; }}
      th {{ background: #eef2ff; }}
      caption {{ text-align: left; font-weight: 600; margin-bottom: 8px; }}
    </style>
  </head>
  <body>
    <h1>Commvault server report — {html.escape(display_name)}</h1>
    <p><strong>ID:</strong> {html.escape(str(summary.get('client_id')))}<br/>
       <strong>Generated:</strong> {html.escape(generated_label)}<br/>
       <strong>Window:</strong> {html.escape(str(window_hours) + 'h' if window_hours is not None else 'n/a')}</p>
    <p><strong>Application data:</strong> {_format_bytes_for_report(stats.get('total_application_bytes'))}<br/>
       <strong>Media:</strong> {_format_bytes_for_report(stats.get('total_media_bytes'))}<br/>
       <strong>Savings:</strong> {_format_bytes_for_report(stats.get('savings_bytes'))}<br/>
       <strong>Average reduction:</strong> {_format_percent(stats.get('average_savings_percent'))} (≈ {stats.get('average_reduction_ratio_text') or '-'})
    </p>
    {'<h2>Restore points by plan</h2><ul>' + plan_html + '</ul>' if plan_html else ''}
    <table>
      <caption>Commvault jobs</caption>
      <thead><tr>{headings}</tr></thead>
      <tbody>{job_rows_html}</tbody>
    </table>
  </body>
</html>
"""


def _render_commvault_server_docx(
    summary: dict[str, Any],
    stats: dict[str, Any],
    jobs: list[dict[str, Any]],
    metrics_payload: dict[str, Any] | None,
) -> BytesIO:
    """Render Commvault server report as DOCX."""
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise HTTPException(status_code=500, detail="python-docx is required for DOCX export") from exc

    display_name = summary.get("display_name") or summary.get("name") or f"#{summary.get('client_id')}"
    generated = metrics_payload.get("fetched_at") if metrics_payload else None
    generated_label = _format_iso_human(generated)
    window_hours = metrics_payload.get("window_hours") if metrics_payload else None
    job_rows = _commvault_job_table_rows(jobs)

    doc = Document()
    doc.add_heading(f"Commvault server report — {display_name}", level=0)
    doc.add_paragraph(f"Client ID: {summary.get('client_id')}")
    doc.add_paragraph(f"Generated: {generated_label}")
    if window_hours is not None:
        doc.add_paragraph(f"Window: {window_hours}h")

    doc.add_paragraph(
        "Application data: "
        + _format_bytes_for_report(stats.get("total_application_bytes"))
        + "; Media: "
        + _format_bytes_for_report(stats.get("total_media_bytes"))
        + "; Savings: "
        + _format_bytes_for_report(stats.get("savings_bytes"))
    )
    doc.add_paragraph(
        "Average data reduction: "
        + _format_percent(stats.get("average_savings_percent"))
        + f" (≈ {stats.get('average_reduction_ratio_text') or '-'})"
    )

    plan_rows = stats.get("plan_breakdown") or []
    if plan_rows:
        doc.add_heading("Restore points by plan", level=1)
        for row in plan_rows:
            label = row.get("label") or row.get("plan") or "Unnamed"
            count = row.get("restore_points") or 0
            doc.add_paragraph(f"{label}: {count}", style="List Bullet")

    doc.add_heading("Jobs", level=1)
    table = doc.add_table(rows=1, cols=len(COMMVAULT_SERVER_EXPORT_COLUMNS))
    header_cells = table.rows[0].cells
    for idx, (_, title) in enumerate(COMMVAULT_SERVER_EXPORT_COLUMNS):
        header_cells[idx].text = title
    for row in job_rows:
        cells = table.add_row().cells
        for idx, (key, _) in enumerate(COMMVAULT_SERVER_EXPORT_COLUMNS):
            cells[idx].text = row.get(key, "")
    if not job_rows:
        row = table.add_row().cells
        row[0].text = "No jobs"

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _render_commvault_server_xlsx(
    summary: dict[str, Any],
    jobs: list[dict[str, Any]],
) -> BytesIO:
    """Render Commvault server jobs as XLSX."""
    try:
        from openpyxl import Workbook  # type: ignore
    except ImportError as exc:  # pragma: no cover - dependency managed elsewhere
        raise HTTPException(status_code=500, detail="openpyxl is required for XLSX export") from exc

    workbook = Workbook()
    sheet = workbook.active
    display_name = summary.get("display_name") or summary.get("name") or f"#{summary.get('client_id')}"
    sheet.title = display_name[:31]

    headers = [
        "job_id",
        "job_type",
        "status",
        "localized_status",
        "localized_operation",
        "client_name",
        "client_id",
        "destination_client_name",
        "subclient_name",
        "backup_set_name",
        "application_name",
        "backup_level_name",
        "plan_name",
        "client_groups",
        "storage_policy_name",
        "start_time",
        "end_time",
        "elapsed_seconds",
        "size_of_application_bytes",
        "size_on_media_bytes",
        "total_num_files",
        "percent_complete",
        "percent_savings",
        "average_throughput_gb_per_hr",
        "retain_until",
    ]
    sheet.append(headers)

    for job in jobs:
        row = dict(job)
        groups = row.get("client_groups")
        if isinstance(groups, list):
            row["client_groups"] = ";".join(groups)
        sheet.append([row.get(header, "") for header in headers])

    buffer = BytesIO()
    workbook.save(buffer)
    buffer.seek(0)
    return buffer


def _render_commvault_server_csv(jobs: list[dict[str, Any]]) -> StringIO:
    """Render Commvault server jobs as CSV."""
    fieldnames = [
        "job_id",
        "job_type",
        "status",
        "localized_status",
        "localized_operation",
        "client_name",
        "client_id",
        "destination_client_name",
        "subclient_name",
        "backup_set_name",
        "application_name",
        "backup_level_name",
        "plan_name",
        "client_groups",
        "storage_policy_name",
        "start_time",
        "end_time",
        "elapsed_seconds",
        "size_of_application_bytes",
        "size_on_media_bytes",
        "total_num_files",
        "percent_complete",
        "percent_savings",
        "average_throughput_gb_per_hr",
        "retain_until",
    ]

    buffer = StringIO()
    writer = csv.DictWriter(buffer, fieldnames=fieldnames)
    writer.writeheader()
    for job in jobs:
        row = dict(job)
        groups = row.get("client_groups")
        if isinstance(groups, list):
            row["client_groups"] = ";".join(groups)
        writer.writerow({key: row.get(key, "") for key in fieldnames})
    buffer.seek(0)
    return buffer


def _mb_to_bytes(value: Any) -> int | None:
    """Convert megabytes to bytes."""
    try:
        number = int(value)
    except (TypeError, ValueError):
        return None
    return number * 1024 * 1024


def _normalise_dedupe_ratio(value: Any) -> float | None:
    """Normalize deduplication ratio."""
    try:
        ratio = float(value)
    except (TypeError, ValueError):
        return None
    if ratio <= 0:
        return None
    if ratio > 0 and ratio < 1:
        return ratio * 100.0
    return ratio


def _safe_int(value: Any) -> int | None:
    """Safely parse integer."""
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def _safe_float(value: Any) -> float | None:
    """Safely parse float."""
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _load_commvault_storage() -> dict[str, Any]:
    """Load cached Commvault storage from disk."""
    path = _data_dir() / COMMVAULT_STORAGE_JSON
    if not path.exists():
        return {
            "pools": [],
            "generated_at": None,
            "total_cached": 0,
            "version": 1,
        }
    try:
        with path.open("r", encoding="utf-8") as handle:
            data = json.load(handle)
            if isinstance(data, Mapping):
                pools = list(data.get("pools") or [])
                generated = data.get("generated_at")
                total_cached = data.get("total_cached")
                try:
                    total_cached = int(total_cached)
                except (TypeError, ValueError):
                    total_cached = len(pools)
                return {
                    "pools": pools,
                    "generated_at": generated,
                    "total_cached": total_cached,
                    "version": data.get("version", 1),
                }
    except Exception:
        logger.exception("Failed to load cached Commvault storage pools", extra={"event": "commvault_storage_cache_error"})
    return {
        "pools": [],
        "generated_at": None,
        "total_cached": 0,
        "version": 1,
    }


def _summarise_storage_pool(
    pool: CommvaultStoragePool, details: Mapping[str, Any] | None
) -> dict[str, Any]:
    """Summarize storage pool information."""
    details = details or {}
    total_capacity_bytes = _mb_to_bytes(pool.total_capacity_mb)
    used_bytes = _mb_to_bytes(pool.size_on_disk_mb)
    free_bytes = _mb_to_bytes(pool.total_free_space_mb)
    usage_pct: float | None = None
    if total_capacity_bytes and used_bytes is not None:
        try:
            usage_pct = (used_bytes / total_capacity_bytes) * 100.0
        except ZeroDivisionError:
            usage_pct = None

    data_reduction = details.get("dataReduction") or details.get("dataReductionInfo")
    dedupe_ratio = None
    dedupe_savings_bytes = None
    if isinstance(data_reduction, Mapping):
        ratios = [
            data_reduction.get("dedupeRatio"),
            data_reduction.get("reductionRatio"),
            data_reduction.get("globalReductionPercent"),
        ]
        for candidate in ratios:
            dedupe_ratio = _normalise_dedupe_ratio(candidate)
            if dedupe_ratio is not None:
                break
        savings_keys = [
            data_reduction.get("dedupeSavingsInBytes"),
            data_reduction.get("totalSavingsInBytes"),
        ]
        for candidate in savings_keys:
            dedupe_savings_bytes = _safe_int(candidate)
            if dedupe_savings_bytes is not None:
                break

    logical_capacity_bytes = None
    if isinstance(details.get("usage"), Mapping):
        logical_capacity_bytes = _safe_int(details["usage"].get("logicalSpaceUsedInBytes"))

    return {
        "pool_id": pool.pool_id,
        "name": pool.name,
        "status": pool.status,
        "storage_type_code": pool.storage_type_code,
        "storage_pool_type_code": pool.storage_pool_type_code,
        "storage_sub_type_code": pool.storage_sub_type_code,
        "storage_policy_id": pool.storage_policy_id,
        "storage_policy_name": pool.storage_policy_name,
        "region_name": pool.region_display_name or pool.region_name,
        "total_capacity_bytes": total_capacity_bytes,
        "used_bytes": used_bytes,
        "free_bytes": free_bytes,
        "logical_capacity_bytes": logical_capacity_bytes,
        "usage_percent": usage_pct,
        "dedupe_ratio": dedupe_ratio,
        "dedupe_savings_bytes": dedupe_savings_bytes,
        "number_of_nodes": pool.number_of_nodes,
        "is_archive_storage": pool.is_archive_storage,
        "cloud_storage_class_name": pool.cloud_storage_class_name,
        "library_ids": list(pool.library_ids),
        "raw": pool.raw,
        "details": details,
    }


def _enrich_commvault_plan_row(row: dict[str, Any], detail: Mapping[str, Any]) -> None:
    """Enrich plan row with detail information."""
    plan_section = detail.get("plan") if isinstance(detail.get("plan"), Mapping) else detail
    summary = plan_section.get("summary") if isinstance(plan_section.get("summary"), Mapping) else {}
    plan_info = summary.get("plan") if isinstance(summary.get("plan"), Mapping) else {}

    plan_type_code = _safe_int(plan_info.get("planType") or summary.get("type"))
    if plan_type_code is not None:
        row["plan_type"] = _label_commvault_plan_type(plan_type_code, row.get("plan_type"))
    elif isinstance(row.get("plan_type"), str) and row["plan_type"].isdigit():
        row["plan_type"] = _label_commvault_plan_type(_safe_int(row["plan_type"]), row["plan_type"])

    status_flag = _safe_int(summary.get("planStatusFlag"))
    status_label = _plan_status_from_flag(status_flag)
    if status_label:
        row["status"] = status_label

    rpo_minutes = _safe_int(summary.get("rpoInMinutes"))
    if rpo_minutes is None:
        rpo_minutes = _safe_int(summary.get("slaInMinutes"))
    rpo_label = _format_minutes_label(rpo_minutes)
    if rpo_label:
        row["rpo"] = rpo_label

    if not row.get("associated_entities"):
        assoc = plan_section.get("associatedEntitiesCount")
        total = 0
        if isinstance(assoc, Mapping):
            all_clients = assoc.get("allClients")
            if isinstance(all_clients, Sequence):
                for item in all_clients:
                    if not isinstance(item, Mapping):
                        continue
                    count = _safe_int(item.get("count"))
                    if count:
                        total += count
        if total:
            row["associated_entities"] = total

    if not row.get("copy_count"):
        num_copies = _safe_int(plan_section.get("numCopies")) or _safe_int(summary.get("numCopies"))
        if num_copies:
            row["copy_count"] = num_copies


def _refresh_commvault_plans_sync(limit: int | None = None, plan_type: str | None = None) -> dict[str, Any]:
    """Refresh Commvault plans cache."""
    fetch_limit = 500 if limit is None else max(1, limit)

    with _commvault_plans_lock:
        with task_logging("commvault.plans.refresh", limit=fetch_limit, plan_type=plan_type) as task_log:
            try:
                client = _commvault_client_from_env()
            except RuntimeError as exc:
                raise HTTPException(status_code=500, detail=str(exc)) from exc

            try:
                plans = client.list_plans(limit=fetch_limit, plan_type=plan_type)
            except CommvaultError as exc:
                raise HTTPException(status_code=502, detail=f"Commvault error: {exc}") from exc

            plan_details: dict[int, Mapping[str, Any]] = {}
            for plan in plans:
                if plan.plan_id is None:
                    continue
                try:
                    plan_details[plan.plan_id] = client.get_plan_details(plan.plan_id)
                except CommvaultError as exc:
                    logger.warning(
                        "Failed to load plan details",
                        extra={"event": "commvault_plan_detail_error", "plan_id": plan.plan_id, "error": str(exc)},
                    )

            rows: list[dict[str, Any]] = []
            for plan in plans:
                row = _serialise_commvault_plan(plan)
                detail = plan_details.get(plan.plan_id or -1)
                if detail:
                    _enrich_commvault_plan_row(row, detail)
                rows.append(row)

            plan_types = sorted(
                {
                    str(row.get("plan_type")).strip()
                    for row in rows
                    if isinstance(row.get("plan_type"), str) and str(row.get("plan_type")).strip()
                },
                key=str.casefold,
            )

            payload: dict[str, Any] = {
                "plans": rows,
                "generated_at": datetime.now(tz=UTC).isoformat(),
                "total_cached": len(rows),
                "plan_types": plan_types,
                "version": 1,
            }
            _write_commvault_plans_json(payload)
            task_log.add_success(plans=len(rows), plan_types=len(plan_types))
            return payload


def _refresh_commvault_storage_sync() -> dict[str, Any]:
    """Refresh Commvault storage cache."""
    with _commvault_storage_lock:
        with task_logging("commvault.storage.refresh") as task_log:
            try:
                client = _commvault_client_from_env()
            except RuntimeError as exc:
                raise HTTPException(status_code=500, detail=str(exc)) from exc

            try:
                pools = list(client.list_storage_pools())
            except CommvaultError as exc:
                raise HTTPException(status_code=502, detail=f"Commvault error: {exc}") from exc

            summaries: list[dict[str, Any]] = []
            for pool in pools:
                details_payload: Mapping[str, Any] | None = None
                try:
                    details = client.get_storage_pool_details(pool.pool_id, summary=pool)
                    details_payload = details.details
                except CommvaultError as exc:
                    logger.warning(
                        "Failed to fetch storage pool details",
                        extra={"event": "commvault_storage_pool_details_error", "pool_id": pool.pool_id, "error": str(exc)},
                    )
                summaries.append(_summarise_storage_pool(pool, details_payload))

            payload: dict[str, Any] = {
                "pools": summaries,
                "generated_at": datetime.now(tz=UTC).isoformat(),
                "total_cached": len(summaries),
                "version": 1,
            }
            _write_commvault_storage_json(payload)
            task_log.add_success(pools=len(summaries))
            return payload


def _refresh_commvault_backups_sync(limit: int, since_hours: int) -> dict[str, Any]:
    """Refresh Commvault backups cache."""
    lookback = max(since_hours, 0)
    since = datetime.now(tz=UTC) - timedelta(hours=lookback) if lookback > 0 else None

    with _commvault_backups_lock:
        with task_logging("commvault.backups.refresh", limit=limit, since_hours=since_hours) as task_log:
            try:
                client = _commvault_client_from_env()
            except RuntimeError as exc:
                raise HTTPException(status_code=500, detail=str(exc)) from exc

            cache_snapshot = _load_commvault_backups()
            jobs_by_id: dict[int, dict[str, Any]] = {}
            prev_digests: dict[int, str] = {}
            for job in cache_snapshot.get("jobs", []):
                job_id = job.get("job_id")
                if job_id is None:
                    continue
                try:
                    key = int(job_id)
                except (TypeError, ValueError):
                    continue
                payload = dict(job)
                jobs_by_id[key] = payload
                prev_digests[key] = _commvault_job_digest(payload)

            tracked_job_ids = {
                job_id for job_id, payload in jobs_by_id.items() if _commvault_job_is_active(payload)
            }
            latest_cached_start = _latest_commvault_start_time(jobs_by_id.values())

            fetch_cutoff = since
            if latest_cached_start and (fetch_cutoff is None or latest_cached_start < fetch_cutoff):
                fetch_cutoff = latest_cached_start

            tracked_min_time: datetime | None = None
            if tracked_job_ids:
                for job_id in tracked_job_ids:
                    payload = jobs_by_id.get(job_id)
                    if not isinstance(payload, Mapping):
                        continue
                    start_dt = _parse_job_datetime(payload.get("start_time"))
                    if start_dt and (tracked_min_time is None or start_dt < tracked_min_time):
                        tracked_min_time = start_dt
            if tracked_min_time and (fetch_cutoff is None or tracked_min_time < fetch_cutoff):
                fetch_cutoff = tracked_min_time

            cached_recent_count: int | None = None
            if fetch_cutoff is not None:
                recent_counter = 0
                for payload in jobs_by_id.values():
                    if not isinstance(payload, Mapping):
                        continue
                    start_dt = _parse_job_datetime(payload.get("start_time"))
                    if start_dt and start_dt >= fetch_cutoff:
                        recent_counter += 1
                cached_recent_count = recent_counter

            try:
                job_list = _collect_commvault_jobs_for_ui(
                    client,
                    limit=limit,
                    offset=0,
                    since=since,
                    tracked_job_ids=tracked_job_ids,
                    latest_start=latest_cached_start,
                    cached_recent_count=cached_recent_count,
                    cutoff_start=fetch_cutoff,
                )
            except CommvaultError as exc:
                raise HTTPException(status_code=502, detail=f"Commvault error: {exc}") from exc

            jobs = list(job_list.jobs)
            if since:
                jobs = [job for job in jobs if job.start_time and job.start_time >= since]

            rows = [_serialise_commvault_job(job) for job in jobs]
            new_jobs_count = 0
            updated_jobs_count = 0
            for row in rows:
                job_id = row.get("job_id")
                if job_id is None:
                    continue
                try:
                    key = int(job_id)
                except (TypeError, ValueError):
                    continue
                digest = _commvault_job_digest(row)
                previous_digest = prev_digests.get(key)
                if previous_digest is None:
                    new_jobs_count += 1
                elif previous_digest != digest:
                    updated_jobs_count += 1
                jobs_by_id[key] = row
                prev_digests[key] = digest

            merged_jobs = sorted(jobs_by_id.values(), key=lambda job: _parse_job_datetime(job.get("start_time")) or datetime.min.replace(tzinfo=UTC), reverse=True)

            cache_payload: dict[str, Any] = {
                "jobs": merged_jobs,
                "generated_at": datetime.now(tz=UTC).isoformat(),
                "total_cached": len(merged_jobs),
                "version": 2,
                "last_refresh_since_hours": since_hours,
                "last_refresh_limit": limit,
                "last_refresh_stats": {
                    "new_jobs": new_jobs_count,
                    "updated_jobs": updated_jobs_count,
                },
            }
            _write_commvault_backups_json(cache_payload)

            filtered_jobs = _filter_commvault_jobs(merged_jobs, since_hours)
            task_log.add_success(jobs=len(filtered_jobs), cached=len(merged_jobs))
            return {
                "jobs": filtered_jobs,
                "generated_at": cache_payload["generated_at"],
                "total_cached": cache_payload["total_cached"],
                "returned": len(filtered_jobs),
                "since_hours": since_hours,
                "limit": limit,
                "new_jobs": new_jobs_count,
                "updated_jobs": updated_jobs_count,
            }


# API Routes

@router.get("/backups")
def commvault_backups_data(
    since_hours: int = Query(24, ge=0, le=24 * 90, description="Only include jobs newer than this window (hours)."),
) -> dict[str, Any]:
    """Return cached Commvault backup jobs for the dashboard."""
    cache = _load_commvault_backups()
    jobs = cache.get("jobs") or []
    filtered = _filter_commvault_jobs(jobs, since_hours)
    return {
        "jobs": filtered,
        "generated_at": cache.get("generated_at"),
        "total_cached": cache.get("total_cached", len(jobs)),
        "returned": len(filtered),
        "since_hours": since_hours,
    }


@router.get("/plans")
def commvault_plans_data(
    refresh: int = Query(0, ge=0, le=1, description="Force refresh of the Commvault plan cache."),
    plan_type: str | None = Query(None, description="Optional case-insensitive plan type filter."),
    limit: int | None = Query(None, ge=1, le=1000, description="Maximum number of plans to return."),
) -> dict[str, Any]:
    """Return cached Commvault plan definitions for the dashboard."""
    cache = _load_commvault_plans()
    if refresh == 1 or not cache.get("plans"):
        try:
            cache = _refresh_commvault_plans_sync(limit=None, plan_type=None)
        except HTTPException as exc:
            if cache.get("plans"):
                logger.warning(
                    "Commvault plan refresh failed; serving cached data",
                    extra={"event": "commvault_plan_refresh_failed", "error": exc.detail if hasattr(exc, "detail") else str(exc)},
                )
            else:
                raise

    plans_payload = list(cache.get("plans") or [])
    plan_type_norm = plan_type.strip().casefold() if plan_type else None
    filtered: list[Mapping[str, Any]] = []
    if plan_type_norm:
        for plan in plans_payload:
            plan_type_value = str(plan.get("plan_type") or "").strip().casefold()
            if plan_type_value == plan_type_norm:
                filtered.append(plan)
    else:
        filtered = plans_payload

    if limit is not None:
        filtered = filtered[:limit]

    return {
        "plans": filtered,
        "total_cached": cache.get("total_cached", len(plans_payload)),
        "generated_at": cache.get("generated_at"),
        "plan_types": cache.get("plan_types") or [],
        "requested_plan_type": plan_type,
        "returned": len(filtered),
        "version": cache.get("version", 1),
    }


@router.post("/plans/refresh")
async def commvault_plans_refresh(
    limit: int = Query(0, ge=0, le=1000, description="Maximum number of plans to fetch (0 = default)."),
    plan_type: str | None = Query(None, description="Restrict refresh to a specific plan type."),
) -> dict[str, Any]:
    """Refresh Commvault plans cache."""
    fetch_limit = None if limit == 0 else limit
    return await asyncio.to_thread(_refresh_commvault_plans_sync, fetch_limit, plan_type)


@router.post("/backups/refresh")
async def commvault_backups_refresh(
    limit: int = Query(0, ge=0, le=500, description="Maximum number of jobs to retain (0 = all)"),
    since_hours: int = Query(24, ge=0, le=24 * 90, description="Only include jobs newer than this window (hours)."),
) -> dict[str, Any]:
    """Fetch fresh Commvault backup jobs and cache them to disk."""
    return await asyncio.to_thread(_refresh_commvault_backups_sync, limit, since_hours)


@router.get("/backups/file", response_model=None)
def commvault_backups_file(
    file_format: Literal["json", "csv"] = Query("json", description="File format to download"),
) -> FileResponse | StreamingResponse:
    """Download Commvault backups as file."""
    if file_format == "json":
        path = _data_dir() / COMMVAULT_BACKUPS_JSON
        if not path.exists():
            raise HTTPException(status_code=404, detail="Commvault export not available")
        return FileResponse(path, media_type="application/json", filename=path.name)

    # Generate CSV on demand from cached dataset
    cache = _load_commvault_backups()
    jobs = cache.get("jobs") or []
    if not jobs:
        raise HTTPException(status_code=404, detail="Commvault export not available")

    fieldnames = [
        "job_id",
        "job_type",
        "status",
        "localized_status",
        "localized_operation",
        "client_name",
        "client_id",
        "destination_client_name",
        "subclient_name",
        "backup_set_name",
        "application_name",
        "backup_level_name",
        "plan_name",
        "client_groups",
        "storage_policy_name",
        "start_time",
        "end_time",
        "elapsed_seconds",
        "size_of_application_bytes",
        "size_on_media_bytes",
        "total_num_files",
        "percent_complete",
        "percent_savings",
        "average_throughput_gb_per_hr",
        "retain_until",
    ]

    def _iter_rows():
        import csv as _csv
        from io import StringIO

        buffer = StringIO()
        writer = _csv.DictWriter(buffer, fieldnames=fieldnames)
        writer.writeheader()
        yield buffer.getvalue()
        buffer.seek(0)
        buffer.truncate(0)

        writer = _csv.DictWriter(buffer, fieldnames=fieldnames)
        for job in jobs:
            row = dict(job)
            groups = row.get("client_groups")
            if isinstance(groups, list):
                row["client_groups"] = ";".join(groups)
            writer.writerow({k: row.get(k, "") for k in fieldnames})
            yield buffer.getvalue()
            buffer.seek(0)
            buffer.truncate(0)

    return StreamingResponse(_iter_rows(), media_type="text/csv", headers={"Content-Disposition": "attachment; filename=commvault_backups.csv"})


@router.get("/storage")
def commvault_storage_data(
    refresh: int = Query(0, ge=0, le=1, description="Force refresh of storage pool cache before returning data."),
) -> dict[str, Any]:
    """Return cached Commvault storage pool data."""
    if refresh == 1:
        return _refresh_commvault_storage_sync()

    cache = _load_commvault_storage()
    if cache.get("pools"):
        return cache

    return _refresh_commvault_storage_sync()


@router.get("/servers/search")
def commvault_servers_search(
    q: str = Query(..., min_length=1, description="Client name fragment or ID."),
    limit: int = Query(10, ge=1, le=200, description="Maximum number of matches to return."),
) -> dict[str, Any]:
    """Search Commvault servers by name or ID."""
    needle = q.strip().casefold()
    if not needle:
        raise HTTPException(status_code=400, detail="Query must not be empty")

    clients, _, _ = _cached_commvault_clients()
    if not clients:
        return {"results": []}

    matches: list[dict[str, Any]] = []
    seen_names: set[str] = set()
    
    for record in clients:
        variants = record.get("name_variants") or set()
        client_id = record.get("client_id")
        name = record.get("name")
        display_name = record.get("display_name")
        
        # Determine the "primary" name we want to show
        primary_name = display_name or name or ""
        primary_key = primary_name.lower().strip()
        
        should_add = False
        
        if needle.isdigit() and client_id == int(needle):
            should_add = True
            # For ID match, we might want to show it even if name is seen?
            # But the goal is "clean hostnames". If I search by ID, I probably want that specific one.
            # Let's assume ID search overrides dedup or we just dedup by name anyway.
            # If I search 1478, I get "server". If I search 1597, I get "server".
            # If they map to the same name, we might only want one. 
            # But if I search by specific unique ID, I expect that specific result.
            
            # Use specific logic for ID match: exact return
            matches = [
                {
                    "client_id": client_id,
                    "name": name,
                    "display_name": display_name,
                    "job_count": record.get("job_count", 0),
                }
            ]
            break 
             
        elif any(needle in variant for variant in variants if variant):
            should_add = True
            
        if should_add:
            if primary_key not in seen_names:
                seen_names.add(primary_key)
                matches.append(
                    {
                        "client_id": client_id,
                        "name": name,
                        "display_name": display_name,
                        "job_count": record.get("job_count", 0),
                    }
                )
        
        if len(matches) >= limit:
            break

    return {"results": matches[:limit]}


@router.get("/servers/summary")
def commvault_servers_summary(
    client: str = Query(..., description="Commvault client name or ID."),
    job_limit: int = Query(500, ge=0, le=2000, description="Maximum jobs to include (0 = all)."),
    since_hours: int = Query(0, ge=0, le=24 * 365, description="Lookback window in hours (0 = all)."),
    retained_only: bool = Query(True, description="Only include jobs with retention metadata."),
    refresh_cache: bool = Query(False, description="Force bypass of cached job metrics."),
) -> dict[str, Any]:
    """Get Commvault server summary and job metrics."""
    summary, metrics, jobs, stats = _load_commvault_server_data(
        client,
        job_limit=job_limit,
        since_hours=since_hours,
        retained_only=retained_only,
        refresh_cache=refresh_cache,
    )
    return {
        "client": summary,
        "job_metrics": metrics,
        "stats": stats,
        "jobs": jobs,
    }


@router.get("/servers/export")
def commvault_servers_export(
    client: str = Query(..., description="Commvault client name or ID."),
    file_format: Literal["text", "markdown", "html", "docx", "xlsx", "csv"] = Query(
        "xlsx", description="Export format"
    ),
    job_limit: int = Query(500, ge=0, le=2000, description="Maximum jobs to include (0 = all)."),
    since_hours: int = Query(0, ge=0, le=24 * 365, description="Lookback window in hours (0 = all)."),
    retained_only: bool = Query(True, description="Only include jobs with retention metadata."),
    refresh_cache: bool = Query(False, description="Force bypass of cached job metrics."),
):
    """Export Commvault server report in various formats."""
    summary, metrics, jobs, stats = _load_commvault_server_data(
        client,
        job_limit=job_limit,
        since_hours=since_hours,
        retained_only=retained_only,
        refresh_cache=refresh_cache,
    )
    basename_source = summary.get("display_name") or summary.get("name") or f"client-{summary.get('client_id')}"
    basename = _slugify_commvault_name(str(basename_source))

    extensions = {
        "text": "txt",
        "markdown": "md",
        "html": "html",
        "docx": "docx",
        "xlsx": "xlsx",
        "csv": "csv",
    }
    ext = extensions[file_format]
    filename = f"{basename}.{ext}"
    disposition = {"Content-Disposition": f"attachment; filename={filename}"}

    if file_format == "text":
        content = _render_commvault_server_text_report(summary, stats, jobs, metrics)
        return PlainTextResponse(content, headers=disposition)

    if file_format == "markdown":
        content = _render_commvault_server_markdown_report(summary, stats, jobs, metrics)
        return PlainTextResponse(content, headers=disposition, media_type="text/markdown")

    if file_format == "html":
        content = _render_commvault_server_html_report(summary, stats, jobs, metrics)
        return HTMLResponse(content, headers=disposition)

    if file_format == "docx":
        buffer = _render_commvault_server_docx(summary, stats, jobs, metrics)
        return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers=disposition)

    if file_format == "xlsx":
        buffer = _render_commvault_server_xlsx(summary, jobs)
        return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers=disposition)

    buffer = _render_commvault_server_csv(jobs)
    csv_bytes = BytesIO(buffer.getvalue().encode("utf-8"))
    csv_bytes.seek(0)
    return StreamingResponse(csv_bytes, media_type="text/csv", headers=disposition)


@router.get("/backup-status")
def commvault_backup_status(
    hostname: str = Query(..., min_length=1, description="Hostname, client name, or VM name to search."),
    hours: int = Query(24, ge=1, le=24 * 365, description="Hours of history to look back."),
    limit: int = Query(10, ge=1, le=100, description="Maximum number of jobs to return."),
) -> dict[str, Any]:
    """Get Commvault backup status and recent job history for a hostname.

    This endpoint is optimized for AI tools to quickly check backup health.
    Returns:
    - Client information
    - Recent backup jobs with status
    - Summary statistics (success rate, last successful backup)
    """
    # First search for the client
    needle = hostname.strip().casefold()
    if not needle:
        raise HTTPException(status_code=400, detail="Hostname must not be empty")

    clients, _, _ = _cached_commvault_clients()
    if not clients:
        return {
            "found": False,
            "hostname": hostname,
            "error": "Commvault cache not available - run /commvault/backups/refresh first",
        }

    # Find matching client
    match = None
    for record in clients:
        variants = record.get("name_variants") or set()
        name = record.get("name", "").lower()
        display_name = record.get("display_name", "").lower()

        if needle in name or needle in display_name:
            match = record
            break
        if any(needle in variant for variant in variants if variant):
            match = record
            break

    if not match:
        return {
            "found": False,
            "hostname": hostname,
            "message": f"No Commvault client found matching '{hostname}'",
        }

    # Get client details and jobs
    client_id = match.get("client_id")
    client_name = match.get("display_name") or match.get("name")

    try:
        summary, metrics, jobs, stats = _load_commvault_server_data(
            str(client_name),
            job_limit=limit,
            since_hours=hours,
            retained_only=False,
            refresh_cache=False,
        )
    except Exception as e:
        return {
            "found": True,
            "hostname": hostname,
            "client_id": client_id,
            "client_name": client_name,
            "error": f"Failed to load job data: {e}",
        }

    # Build response
    recent_jobs = []
    for job in jobs[:limit]:
        recent_jobs.append({
            "job_id": job.get("job_id"),
            "start": job.get("start"),
            "end": job.get("end"),
            "status": job.get("status"),
            "plan": job.get("plan"),
            "app_size": job.get("app_size"),
            "media_size": job.get("media_size"),
        })

    # Calculate success rate
    total = stats.get("job_count", 0)
    successful = stats.get("completed", 0)
    success_rate = (successful / total * 100) if total > 0 else 0

    return {
        "found": True,
        "hostname": hostname,
        "client_id": client_id,
        "client_name": client_name,
        "lookback_hours": hours,
        "summary": {
            "total_jobs": total,
            "successful": successful,
            "failed": stats.get("failed", 0),
            "running": stats.get("running", 0),
            "success_rate_pct": round(success_rate, 1),
            "last_successful": stats.get("last_success"),
            "last_job": stats.get("last_job"),
        },
        "recent_jobs": recent_jobs,
    }
